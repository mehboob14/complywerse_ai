# -*- coding: utf-8 -*-
"""Parse framework template folders into per-framework tab definitions (JSON).

Each definition mirrors the frontend register/document config so the backend can
serve it and the generic UI can render it. xlsx -> register tab, docx -> document tab.
"""
import os, re, json, sys
import openpyxl
import docx

SRC = r"C:\Users\Admin\Downloads\Compliance-Templates-Organized-20260706T193148Z-3-001\Compliance-Templates-Organized"
OUT = r"C:\Users\Admin\AppData\Local\Temp\claude\C--Users-Admin-Documents-GRC-Tenant\69b1152d-3512-4d95-a33b-6098e93d2ab3\scratchpad\defs"

# folder -> (framework_key, display_name, [name-match patterns for frontend gating])
FRAMEWORKS = {
    "ISO-27001":       ("iso_27001",   "ISO/IEC 27001", ["27001"]),
    "CIS-Controls-v8": ("cis_controls","CIS Controls v8", ["ciscontrol", "cisv8", "cis8", "ciscontrolsv8"]),
    "DORA":            ("dora",        "DORA", ["dora", "digitaloperationalresilience"]),
    "GDPR":            ("gdpr",        "GDPR", ["gdpr", "generaldataprotection"]),
    "HIPAA":           ("hipaa",       "HIPAA", ["hipaa"]),
    "HITRUST":         ("hitrust_csf", "HITRUST CSF", ["hitrust"]),
    "ISO-22301":       ("iso_22301",   "ISO 22301", ["22301"]),
    "ISO-42001":       ("iso_42001",   "ISO 42001", ["42001"]),
    "NIS2":            ("nis2",        "NIS2", ["nis2"]),
    "NIST-800-Series": ("nist_800_53", "NIST 800-53 / 800-171", ["80053", "800171", "nist800"]),
    "NIST-AI-RMF":     ("nist_airmf",  "NIST AI RMF", ["airmf", "airiskmanagement", "nistai"]),
    "NIST-CSF-2.0":    ("nist_csf",    "NIST CSF 2.0", ["nistcsf", "cybersecurityframework", "csf2"]),
    "PCI-DSS":         ("pci_dss",     "PCI DSS", ["pci"]),
    "SOC-2":           ("soc2",        "SOC 2", ["soc2"]),
    "SOX-ITGC":        ("sox",         "SOX ITGC", ["sox", "sarbanes", "itgc"]),
}

# ── helpers ──────────────────────────────────────────────────────────────────
def slug(s):
    s = re.sub(r"[^a-z0-9]+", "_", (s or "").lower()).strip("_")
    return re.sub(r"_+", "_", s)[:48] or "col"

def humanize_file(fn):
    base = os.path.splitext(fn)[0]
    base = re.sub(r"[-_]template.*$", "", base, flags=re.I)
    base = re.sub(r"\s*\(\d+\)\s*$", "", base)
    base = base.replace("-", " ").replace("_", " ")
    # strip framework prefixes
    base = re.sub(r"^(iso.?27001|cis.?controls|dora|gdpr|hipaa|hitrust|iso.?22301|iso.?42001|nis2|nist.?800.?\d*|nist.?ai.?rmf|nist.?csf|pci.?dss|soc.?2|sox)\b[\s-]*", "", base, flags=re.I)
    return re.sub(r"\s+", " ", base).strip().title() or "Register"

DATE_RE = re.compile(r"\b(date|due|target|deadline|review|expiry|effective|when|timeline)\b", re.I)
USER_RE = re.compile(r"\b(owner|responsible|assigned|assignee|reviewer|approver|approved by|accountable|dpo|contact|lead|sponsor)\b", re.I)
EVID_RE = re.compile(r"\b(evidence|attachment|link to)\b", re.I)
SELECT_RE = re.compile(r"\b(status|result|rating|level|priority|severity|likelihood|impact|applicab|maturity|conform|compliance|y/?n|yes/no|band|tier|score|state|decision|classification|category|frequency|disposition|outcome)\b", re.I)
LONG_RE = re.compile(r"\b(description|justification|gap|action|notes|remediation|requirement|comment|detail|analysis|scope|plan|recommendation|finding|purpose|guidance|summary|question|objective|control|activity|measure|procedure|mitigation|treatment|response|note)\b", re.I)

TONE_MAP = [
    (re.compile(r"^(conform|compliant|yes|covered|implemented|complete|pass|closed|low|effective|met|approved)$", re.I), "emerald"),
    (re.compile(r"^(nonconform|non-conform|no|not implemented|fail|open|high|critical|rejected|overdue|gap)$", re.I), "rose"),
    (re.compile(r"^(partial|in progress|in-progress|medium|moderate|pending|planned|ofi|review)$", re.I), "amber"),
    (re.compile(r"^(n/?a|not applicable|na|deferred|informational|info)$", re.I), "slate"),
    (re.compile(r"^(very high|severe)$", re.I), "orange"),
]
def tone_for(v):
    for rx, t in TONE_MAP:
        if rx.match((v or "").strip()):
            return t
    return "slate"

def cell_str(v):
    if v is None: return ""
    s = str(v).replace("\r", " ").strip()
    return re.sub(r"\s+\n", "\n", s)

# ── xlsx -> register def ─────────────────────────────────────────────────────
GUIDANCE = re.compile(r"how to use|guidance|instruction|read me|readme|cover|intro|legend|key\b|about", re.I)

def parse_xlsx(path):
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        return None
    def _hscore(cells):
        # header-likeness: short label-ish cells score +1, long sentences (data) -1
        s = 0
        for c in cells:
            c = (c or "").strip()
            if not c:
                continue
            if len(c) > 55:
                s -= 1
            elif re.match(r"^\d+([.,]\d+)?$", c):
                pass  # pure number = data
            elif re.match(r"^[A-Za-z(#].{0,54}$", c):
                s += 1
        return s

    best = None  # (score, sheet, header_row_idx, header_cells, rows)
    for ws in wb.worksheets:
        if GUIDANCE.search(ws.title or ""):
            continue
        rows = list(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 250), values_only=True))
        hdr_idx, hdr_cells, hdr_sc = None, None, 1
        for i, r in enumerate(rows[:15]):
            cells = [cell_str(c) for c in r]
            n = sum(1 for c in cells if c)
            sc = _hscore(cells)
            has_below = any(any(cell_str(c) for c in rr) for rr in rows[i + 1:i + 4])
            if n >= 2 and sc >= 2 and has_below and sc > hdr_sc:
                hdr_idx, hdr_cells, hdr_sc = i, cells, sc
        if hdr_idx is None:
            continue
        if best is None or hdr_sc > best[0]:
            best = (hdr_sc, ws, hdr_idx, hdr_cells, rows)
    if not best:
        # fall back: first non-guidance sheet, first row as header
        for ws in wb.worksheets:
            if GUIDANCE.search(ws.title or ""): continue
            rows = list(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 40), values_only=True))
            if rows:
                best = (0, ws, 0, [cell_str(c) for c in rows[0]], rows); break
    if not best:
        return None
    _, ws, hdr_idx, hdr_cells, rows = best

    # build columns from header cells (non-empty)
    cols, seen = [], set()
    col_positions = []
    for pos, h in enumerate(hdr_cells):
        if not h: continue
        key = slug(h)
        base = key; k = 1
        while key in seen:
            k += 1; key = f"{base}_{k}"
        seen.add(key)
        col_positions.append((pos, key, h))

    # skip registers whose header row didn't parse into real column labels
    header_like = sum(1 for _, _, h in col_positions
                      if re.match(r"^[A-Za-z(#].{0,54}$", (h or "").strip()) and not re.match(r"^\d+$", (h or "").strip()))
    if len(col_positions) < 2 or header_like < 2:
        return None

    # collect data rows (after header), map to {key: value}
    data_rows = []
    for r in rows[hdr_idx + 1:]:
        vals = [cell_str(c) for c in r]
        if not any(vals): continue
        row = {}
        for pos, key, _ in col_positions:
            row[key] = vals[pos] if pos < len(vals) else ""
        # skip trailing summary rows (single cell like "Coverage %:")
        nonempty = [v for v in row.values() if v]
        if len(nonempty) <= 1 and any(re.search(r"%|total|coverage|count", (v or ""), re.I) for v in nonempty):
            continue
        data_rows.append(row)
    data_rows = data_rows[:150]

    # infer column type per header
    columns = []
    for pos, key, label in col_positions:
        vals = [dr[key] for dr in data_rows if dr.get(key)]
        avglen = (sum(len(v) for v in vals) / len(vals)) if vals else 0
        distinct = list(dict.fromkeys(vals))
        ctype, options, picker = "text", None, None
        disp_label = label
        # header-embedded options, e.g. "Response (Yes/Partial/No/N/A)" or "Applicable? (Y/N)"
        hdr_opts = None
        hm = re.search(r"\(([^)]{2,44})\)", label)
        if hm and ("/" in hm.group(1) or "," in hm.group(1)):
            inner = re.sub(r"\bN\s*/\s*A\b", "N⁄A", hm.group(1), flags=re.I)  # protect N/A
            toks = [t.strip().replace("⁄", "/") for t in re.split(r"[/,]", inner) if t.strip()]
            toks = [t for t in toks if 1 <= len(t) <= 22]
            if 2 <= len(toks) <= 6:
                hdr_opts = toks
                disp_label = re.sub(r"\s*\([^)]*\)\s*$", "", label).strip() or label
        if USER_RE.search(label):
            picker = "users"; ctype = "text"
        elif EVID_RE.search(label):
            picker = "evidence"; ctype = "text"
        elif hdr_opts:
            ctype = "select"
            options = [{"value": v, "label": v, "tone": tone_for(v)} for v in hdr_opts]
        elif DATE_RE.search(label) and avglen <= 24:
            ctype = "date"
        elif SELECT_RE.search(label) and 0 < len(distinct) <= 12 and all(len(v) <= 32 for v in distinct):
            ctype = "select"
            options = [{"value": v, "label": v, "tone": tone_for(v)} for v in distinct]
        elif LONG_RE.search(label) or avglen > 55:
            ctype = "textarea"
        columns.append({"key": key, "label": disp_label[:60], "type": ctype,
                        **({"options": options} if options else {}),
                        **({"picker": picker} if picker else {})})

    # seed rows keep only their data values
    seed = [{"data": dr} for dr in data_rows]

    # auto form sections: item (first 2 non-long), ownership (date/user), rest details
    first_keys = [c["key"] for c in columns[:2]]
    own_keys = [c["key"] for c in columns if c["type"] == "date" or c.get("picker") == "users"]
    mid_keys = [c["key"] for c in columns if c["key"] not in first_keys and c["key"] not in own_keys]
    sections = []
    if first_keys: sections.append({"title": "Item", "keys": first_keys})
    if mid_keys: sections.append({"title": "Assessment", "keys": mid_keys})
    if own_keys: sections.append({"title": "Ownership & dates", "keys": own_keys})

    return {"columns": columns, "seed": seed, "formSections": sections, "sheet": ws.title}

# ── docx -> document def ─────────────────────────────────────────────────────
BOILER = re.compile(r"security scientist|before you use this template|free template|works with|part of the free", re.I)

def parse_docx(path):
    try:
        d = docx.Document(path)
    except Exception:
        return None
    # iterate body in order (paragraphs + tables)
    from docx.oxml.ns import qn
    body = d.element.body
    para_map = {p._p: p for p in d.paragraphs}
    tbl_map = {t._tbl: t for t in d.tables}
    sections = []
    cur = None
    started = False
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = para_map.get(child)
            if p is None: continue
            txt = p.text.strip()
            if not txt: continue
            style = (p.style.name if p.style else "") or ""
            is_head = style.startswith("Heading") or style == "Title"
            if is_head:
                # skip the "Before you use..." boilerplate heading
                if BOILER.search(txt):
                    cur = None; continue
                started = True
                cur = {"heading": txt[:120], "body": ""}
                sections.append(cur)
            else:
                if not started or cur is None:
                    continue
                if BOILER.search(txt):
                    continue
                prefix = "- " if style.startswith("List") else ""
                cur["body"] = (cur["body"] + "\n" + prefix + txt).strip() if cur["body"] else (prefix + txt)
        elif child.tag == qn('w:tbl'):
            t = tbl_map.get(child)
            if t is None or cur is None or not started:
                continue
            rows = []
            for row in t.rows[:30]:
                rows.append([c.text.replace("\n", " ").strip()[:200] for c in row.cells])
            if rows:
                cur["table"] = {"columns": rows[0], "rows": rows[1:]}
    # trim empties, cap
    sections = [s for s in sections if s.get("body") or s.get("table")][:30]
    # control_ref: look for "Cl. X" in first ~30 paras
    cref = None
    for p in d.paragraphs[:40]:
        m = re.search(r"\bCl\.?\s*([0-9A-Z][0-9A-Za-z.\-]*)", p.text)
        if m: cref = "Cl. " + m.group(1); break
    return {"sections": sections, "control_ref": cref}

# ── main ─────────────────────────────────────────────────────────────────────
def main():
    os.makedirs(OUT, exist_ok=True)
    summary = []
    for folder, (key, display, patterns) in FRAMEWORKS.items():
        fdir = os.path.join(SRC, folder)
        if not os.path.isdir(fdir):
            print("MISSING", folder); continue
        registers, documents = [], []
        for fn in sorted(os.listdir(fdir)):
            full = os.path.join(fdir, fn)
            low = fn.lower()
            if "(1)" in low or "(2)" in low:  # skip duplicate copies
                continue
            if low.endswith(".xlsx"):
                r = parse_xlsx(full)
                if r and r["columns"]:
                    registers.append({
                        "type": f"{key}__{slug(os.path.splitext(fn)[0])}",
                        "label": humanize_file(fn),
                        "description": f"{display} — {humanize_file(fn)} (from the template).",
                        "columns": r["columns"], "seed": r["seed"],
                        "formSections": r["formSections"],
                    })
            elif low.endswith(".docx"):
                dd = parse_docx(full)
                if dd and dd["sections"]:
                    documents.append({
                        "type": f"{key}__{slug(os.path.splitext(fn)[0])}",
                        "label": humanize_file(fn),
                        "control_ref": dd["control_ref"],
                        "sections": dd["sections"],
                    })
        definition = {
            "framework_key": key, "display_name": display, "name_patterns": patterns,
            "registers": registers, "documents": documents,
        }
        with open(os.path.join(OUT, f"{key}.json"), "w", encoding="utf-8") as f:
            json.dump(definition, f, ensure_ascii=False, indent=1)
        summary.append((folder, key, len(registers), sum(len(r['columns']) for r in registers), len(documents)))
    print(f"{'FOLDER':<18} {'KEY':<14} {'REG':<4} {'COLS':<5} {'DOC':<4}")
    for folder, key, nr, nc, nd in summary:
        print(f"{folder:<18} {key:<14} {nr:<4} {nc:<5} {nd:<4}")
    print("\nTotal registers:", sum(s[2] for s in summary), "| documents:", sum(s[4] for s in summary))

if __name__ == "__main__":
    main()
