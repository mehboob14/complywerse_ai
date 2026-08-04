"""Render an artifact's stored content into a downloadable native file.

The artifact content is GitHub-flavoured Markdown (documents/guides) plus, for
tabular artifacts, a structured `table` ({columns, example_rows, field_guidance,
maintenance}). This module turns either of those into:

  md   → the markdown source as-is
  csv  → the table (structured if present, else the first markdown table)
  xlsx → a real spreadsheet (template sheet + a Column Guidance + Notes sheet)
  docx → a Word document (headings / paragraphs / lists / tables)
  pdf  → a PDF (reportlab platypus)

openpyxl / python-docx (`docx`) / reportlab are all already installed. The
markdown parser is intentionally small — it handles the constructs our generator
emits (headings, paragraphs, bullet/numbered lists, pipe tables) and degrades
gracefully on anything else.
"""
from __future__ import annotations

import csv as _csv
import html as _htmlmod
import io
import re
from typing import Any, Dict, List, Optional, Tuple

_MEDIA: Dict[str, Tuple[str, str]] = {
    "md": ("text/markdown; charset=utf-8", "md"),
    "csv": ("text/csv; charset=utf-8", "csv"),
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "pdf": ("application/pdf", "pdf"),
}
SUPPORTED_FORMATS = list(_MEDIA.keys())


def formats_for(content_format: Optional[str]) -> List[str]:
    """The download formats that make sense for a given content mode."""
    if (content_format or "").lower() == "table":
        return ["xlsx", "csv", "pdf", "md"]
    return ["docx", "pdf", "md"]   # documents + guides


# ─── tiny markdown parser ────────────────────────────────────────────

def parse_markdown(md: str) -> List[dict]:
    """Parse markdown into a flat block list: heading/para/li/table."""
    blocks: List[dict] = []
    lines = (md or "").replace("\r\n", "\n").split("\n")
    para: List[str] = []

    def flush() -> None:
        if para:
            blocks.append({"type": "para", "text": " ".join(para).strip()})
            para.clear()

    i = 0
    n = len(lines)
    while i < n:
        s = lines[i].strip()
        # pipe table — collect consecutive '|' lines
        if s.startswith("|") and s.count("|") >= 2:
            flush()
            raw: List[str] = []
            while i < n and lines[i].strip().startswith("|"):
                raw.append(lines[i].strip())
                i += 1
            parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in raw]
            # drop separator rows (all cells are dashes/colons)
            data = [r for r in parsed if not all(c and set(c) <= set("-: ") for c in r)]
            headers = data[0] if data else []
            body = data[1:] if len(data) > 1 else []
            blocks.append({"type": "table", "headers": headers, "rows": body})
            continue
        if not s:
            flush()
            i += 1
            continue
        if s.startswith("#"):
            flush()
            level = len(s) - len(s.lstrip("#"))
            blocks.append({"type": "heading", "level": min(max(level, 1), 6), "text": s.lstrip("#").strip()})
            i += 1
            continue
        if s.startswith(("- ", "* ", "+ ")) or re.match(r"^\d+\.\s", s):
            flush()
            blocks.append({"type": "li", "text": re.sub(r"^(\s*[-*+]\s+|\s*\d+\.\s+)", "", s)})
            i += 1
            continue
        para.append(s)
        i += 1
    flush()
    return blocks


def _plain(text: Any) -> str:
    """Strip markdown emphasis for plain renderers (docx)."""
    t = str(text or "")
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"`(.+?)`", r"\1", t)
    t = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"\1", t)
    return t


def _rl(text: Any) -> str:
    """Escape + minimal inline markup for reportlab Paragraph."""
    t = _htmlmod.escape(str(text or ""))
    t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
    t = re.sub(r"`(.+?)`", r"\1", t)
    t = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)", r"<i>\1</i>", t)
    return t


def _first_table(blocks: List[dict]) -> Optional[dict]:
    return next((b for b in blocks if b["type"] == "table"), None)


def _table_data(blocks: List[dict], table: Optional[dict]) -> Tuple[List[str], List[list], Dict[str, str], str]:
    """Resolve (headers, rows, field_guidance, maintenance) preferring structured data."""
    if table and table.get("columns"):
        return (list(table["columns"]), list(table.get("example_rows") or []),
                dict(table.get("field_guidance") or {}), str(table.get("maintenance") or ""))
    tb = _first_table(blocks)
    if tb:
        return list(tb["headers"]), list(tb["rows"]), {}, ""
    return [], [], {}, ""


# ─── builders ────────────────────────────────────────────────────────

def _to_docx(title: str, blocks: List[dict]) -> bytes:
    from docx import Document
    doc = Document()
    if title:
        doc.add_heading(_plain(title), 0)
    for b in blocks:
        t = b["type"]
        if t == "heading":
            doc.add_heading(_plain(b["text"]), min(max(b["level"], 1), 4))
        elif t == "para":
            doc.add_paragraph(_plain(b["text"]))
        elif t == "li":
            doc.add_paragraph(_plain(b["text"]), style="List Bullet")
        elif t == "table" and b["headers"]:
            cols = len(b["headers"])
            tbl = doc.add_table(rows=1, cols=cols)
            tbl.style = "Table Grid"
            for j, h in enumerate(b["headers"]):
                tbl.rows[0].cells[j].text = _plain(h)
            for r in b["rows"]:
                cells = tbl.add_row().cells
                for j in range(cols):
                    cells[j].text = _plain(r[j]) if j < len(r) else ""
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _to_pdf(title: str, blocks: List[dict]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle)

    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, title=_plain(title),
                            topMargin=1.4 * cm, bottomMargin=1.4 * cm,
                            leftMargin=1.4 * cm, rightMargin=1.4 * cm)
    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("cell", parent=styles["BodyText"], fontSize=7, leading=9)
    head_style = ParagraphStyle("cellh", parent=cell_style, textColor=colors.white)
    story: List[Any] = []
    if title:
        story.append(Paragraph(_rl(title), styles["Title"]))
        story.append(Spacer(1, 8))
    for b in blocks:
        t = b["type"]
        if t == "heading":
            story.append(Paragraph(_rl(b["text"]), styles[f"Heading{min(max(b['level'], 1), 4)}"]))
        elif t == "para":
            story.append(Paragraph(_rl(b["text"]), styles["BodyText"]))
        elif t == "li":
            story.append(Paragraph("• " + _rl(b["text"]), styles["BodyText"]))
        elif t == "table" and (b["headers"] or b["rows"]):
            ncol = len(b["headers"]) or max((len(r) for r in b["rows"]), default=1)
            header = [Paragraph(_rl(c), head_style) for c in b["headers"]] if b["headers"] else None
            data = []
            if header:
                data.append(header)
            for r in b["rows"]:
                row = [Paragraph(_rl(r[j]) if j < len(r) else "", cell_style) for j in range(ncol)]
                data.append(row)
            if data:
                tb = Table(data, repeatRows=1 if header else 0, hAlign="LEFT",
                           colWidths=[doc.width / ncol] * ncol)
                tb.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]))
                story.append(tb)
        story.append(Spacer(1, 6))
    doc.build(story or [Paragraph(_rl(title or "Document"), styles["BodyText"])])
    return bio.getvalue()


def _to_xlsx(title: str, blocks: List[dict], table: Optional[dict]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    headers, rows, guidance, maintenance = _table_data(blocks, table)
    wb = Workbook()
    ws = wb.active
    ws.title = "Template"
    if headers:
        ws.append([str(h) for h in headers])
        for c in ws[1]:
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor="1F2937")
            c.alignment = Alignment(wrap_text=True, vertical="top")
        for r in rows:
            ws.append([(str(r[j]) if j < len(r) else "") for j in range(len(headers))])
        for j, h in enumerate(headers, 1):
            ws.column_dimensions[get_column_letter(j)].width = min(max(len(str(h)) + 4, 16), 42)
        ws.freeze_panes = "A2"
    else:
        # No table → dump the document text so the file is never empty.
        ws.append([title or "Document"])
        for b in blocks:
            if b.get("text"):
                ws.append([_plain(b["text"])])
        ws.column_dimensions["A"].width = 100
    if guidance:
        gs = wb.create_sheet("Column Guidance")
        gs.append(["Column", "What to enter"])
        for c in gs[1]:
            c.font = Font(bold=True)
        for k, v in guidance.items():
            gs.append([str(k), str(v)])
        gs.column_dimensions["A"].width = 30
        gs.column_dimensions["B"].width = 80
    if maintenance:
        ns = wb.create_sheet("Notes")
        ns.append(["Maintenance"])
        ns["A1"].font = Font(bold=True)
        ns.append([str(maintenance)])
        ns.column_dimensions["A"].width = 100
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def _to_csv(blocks: List[dict], table: Optional[dict]) -> bytes:
    headers, rows, _g, _m = _table_data(blocks, table)
    sio = io.StringIO()
    w = _csv.writer(sio)
    if headers:
        w.writerow([str(h) for h in headers])
        for r in rows:
            w.writerow([(str(r[j]) if j < len(r) else "") for j in range(len(headers))])
    else:
        for b in blocks:
            if b.get("text"):
                w.writerow([_plain(b["text"])])
    return sio.getvalue().encode("utf-8-sig")  # BOM so Excel opens UTF-8 cleanly


def build_export(
    fmt: str,
    *,
    title: str,
    content: str,
    content_format: Optional[str] = None,
    table: Optional[dict] = None,
) -> Tuple[bytes, str, str]:
    """Return (file_bytes, media_type, extension) for the requested format."""
    fmt = (fmt or "").lower().strip()
    if fmt not in _MEDIA:
        raise ValueError(f"Unsupported format '{fmt}'. Supported: {', '.join(SUPPORTED_FORMATS)}")
    blocks = parse_markdown(content or "")
    if fmt == "md":
        data = (content or "").encode("utf-8")
    elif fmt == "csv":
        data = _to_csv(blocks, table)
    elif fmt == "xlsx":
        data = _to_xlsx(title, blocks, table)
    elif fmt == "docx":
        data = _to_docx(title, blocks)
    else:  # pdf
        data = _to_pdf(title, blocks)
    media, ext = _MEDIA[fmt]
    return data, media, ext
