"""NCA Cybersecurity Document Templates router.

Exposes the NCA Saudi reference document catalogue (~80 .docx templates) so
users can:
  1. List the available templates by category
  2. Preview their contents (extracted text)
  3. Create a new GovernanceDocument seeded from a template (as-is or modified)
  4. Generate an AI-drafted document using a template as the reference
  5. Run a side-by-side comparison between an existing document and a template

Templates are stored on disk under `NCA_Documents/` at the repo root and are
auto-discovered at module import time. No user state is mutated by listing
or preview endpoints — template files themselves are read-only.
"""
from ..config import get_openai_api_key, get_openai_model

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..models import GovernanceDocument, GRCUser, get_db
from .auth_router import require_auth, get_user_primary_tenant

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/governance/nca-templates", tags=["NCA Document Templates"])

OPENAI_API_KEY = get_openai_api_key()
OPENAI_BASE_URL = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL") or os.environ.get("OPENAI_BASE_URL")


# ─── Catalogue discovery ─────────────────────────────────────────────────────

def _resolve_templates_dir() -> Optional[Path]:
    """Locate the NCA templates folder.

    Walks up from this file looking for a sibling directory named
    ``NCA_Templates`` (current name) or ``NCA_Documents`` (legacy name).
    Falls back to the ``NCA_TEMPLATES_DIR`` env override. The folder lives
    at the repo root, above ``backend/``.
    """
    here = Path(__file__).resolve()
    candidate_names = ("NCA_Templates", "NCA_Documents")
    for parent in [here, *here.parents]:
        for name in candidate_names:
            candidate = parent.parent / name
            if candidate.is_dir():
                return candidate
    # Fallback: env override
    override = os.environ.get("NCA_TEMPLATES_DIR")
    if override and Path(override).is_dir():
        return Path(override)
    return None


TEMPLATES_DIR = _resolve_templates_dir()


def _detect_category(filename: str) -> str:
    """Map filename prefix to a friendly category label."""
    name_lower = filename.lower()
    if name_lower.startswith("policy"):
        return "Policy"
    if name_lower.startswith("standard"):
        return "Standard"
    if name_lower.startswith("procedure"):
        return "Procedure"
    if name_lower.startswith("program"):
        return "Program"
    if name_lower.startswith("checklist"):
        return "Checklist"
    if name_lower.startswith("form"):
        return "Form"
    if name_lower.startswith("report"):
        return "Report"
    if name_lower.startswith("cybersecurity"):
        return "Cybersecurity Foundation"
    return "Other"


def _humanize_title(filename: str) -> str:
    """Convert filename to a human-readable title."""
    stem = Path(filename).stem
    # Drop trailing template/version markers
    stem = re.sub(r"[_\- ]?template[_\- ]?en[_\- ]?v?[0-9.]*[_\-]?$", "", stem, flags=re.I)
    stem = re.sub(r"[_\- ]?template[_\- ]?$", "", stem, flags=re.I)
    stem = re.sub(r"[_\- ]?\([0-9]+\)$", "", stem)
    # Drop leading category prefix (we already extract it separately)
    stem = re.sub(r"^(POLICY|STANDARD|Standard|PROCEDURE|Procedure|PROGRAM|Checklist|Form|Report|FORM)[_\- ]+", "", stem)
    # Underscore/dash → space
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    return stem


def _build_catalog() -> List[Dict[str, Any]]:
    """Scan TEMPLATES_DIR and return ordered metadata for every .docx/.xlsx."""
    if not TEMPLATES_DIR or not TEMPLATES_DIR.is_dir():
        logger.warning("NCA templates directory not found")
        return []

    items: List[Dict[str, Any]] = []
    seen_keys: set = set()

    for path in sorted(TEMPLATES_DIR.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".docx", ".xlsx"}:
            continue

        title = _humanize_title(path.name)
        category = _detect_category(path.name)
        # Stable id: hash-free, slugified filename
        template_id = re.sub(r"[^a-z0-9]+", "-", path.stem.lower()).strip("-")

        # De-duplicate near-identical filenames (e.g. " (1).docx" copies)
        dedupe_key = re.sub(r"\s*\(\d+\)$", "", path.stem.lower()).rstrip("-_ ")
        if dedupe_key in seen_keys:
            continue
        seen_keys.add(dedupe_key)

        items.append({
            "id": template_id,
            "filename": path.name,
            "title": title,
            "category": category,
            "size_bytes": path.stat().st_size,
            "format": path.suffix.lower().lstrip("."),
        })
    return items


_CATALOG_CACHE: Optional[List[Dict[str, Any]]] = None


def _get_catalog() -> List[Dict[str, Any]]:
    global _CATALOG_CACHE
    if _CATALOG_CACHE is None:
        _CATALOG_CACHE = _build_catalog()
    return _CATALOG_CACHE


def _find_template(template_id: str) -> Optional[Dict[str, Any]]:
    for item in _get_catalog():
        if item["id"] == template_id:
            return item
    return None


def _template_path(template_id: str) -> Path:
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")
    if not TEMPLATES_DIR:
        raise HTTPException(status_code=500, detail="Templates directory not configured")
    return TEMPLATES_DIR / item["filename"]


# ─── Content extraction ─────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    """Extract a .docx file into markdown.

    Walks the document body in document order so headings, paragraphs, and
    tables stay interleaved the way the author wrote them. Uses
    `doc.iter_inner_content()` (python-docx ≥ 1.0) which yields Paragraph
    and Table objects already wrapped — no fragile lxml lookups.

    Word "Heading N" / "Title" styles become `#`-prefixed markdown
    headings; tables become GFM tables with a header separator so
    ReactMarkdown + remark-gfm renders them as real <table>s. Merged cells
    (`gridSpan` / `vMerge`) are de-duplicated so the same content doesn't
    repeat across columns.
    """
    try:
        from docx import Document  # python-docx
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document(str(path))

    def _escape_md_text(s: str) -> str:
        """Escape angle brackets so ReactMarkdown + remark-gfm doesn't try
        to parse template placeholders like `<organization name>` as HTML
        (which would silently eat the text). Backslash-escaped angle
        brackets render as literal `<` / `>` characters.
        """
        return s.replace("<", "\\<").replace(">", "\\>")

    _HEADING_STYLE_RE = re.compile(r"^heading\s*(\d{1,2})?", re.IGNORECASE)

    def _heading_level_from_style(style_name: str) -> Optional[int]:
        """Return a 1..6 heading level if the style name reads as a heading
        ("Heading 1", "heading 12", "Heading", "HeadingTwo", etc.), else
        None. NCA templates inconsistently capitalise and sometimes use
        compound style names like 'heading 12' so a loose regex is safer
        than a strict prefix-equals check.
        """
        m = _HEADING_STYLE_RE.match(style_name or "")
        if not m:
            return None
        raw = m.group(1)
        if not raw:
            return 2
        try:
            level = int(raw)
        except ValueError:
            return 2
        # 'heading 12' is a Word-internal style number, not a 12-level
        # heading. Anything past markdown's 6 levels collapses to H2 so
        # the structure still reads as a real section instead of as
        # microscopic h6 text.
        if level <= 0:
            return 2
        if level > 6:
            return 2
        return max(1, level)

    def _is_visually_bold_heading(para) -> bool:
        """Heuristic: a short paragraph whose entire content is bold and
        which doesn't end in sentence punctuation is almost certainly a
        section label even when the author skipped Word's Heading styles.

        Checks bold at three layers — the run, the run's font, and the
        paragraph style's font — because NCA templates sometimes inherit
        bold from a custom 'Normal'-derived style instead of declaring it
        on the run itself.
        """
        text = (para.text or "").strip()
        if not text or len(text) > 80:
            return False
        if text.endswith((".", "?", "!", ":", ";")):
            return False
        try:
            runs = [r for r in para.runs if (r.text or "").strip()]
        except Exception:  # noqa: BLE001
            return False
        if not runs:
            return False
        # Style-level bold fallback
        style_bold = False
        try:
            style_font = para.style.font if para.style else None
            if style_font is not None and getattr(style_font, "bold", None):
                style_bold = True
        except Exception:  # noqa: BLE001
            pass

        def _run_is_bold(r) -> bool:
            if getattr(r, "bold", False):
                return True
            font = getattr(r, "font", None)
            if font is not None and getattr(font, "bold", False):
                return True
            return False

        return all(_run_is_bold(r) or style_bold for r in runs)

    def _format_paragraph(para) -> Optional[str]:
        text = (para.text or "").strip()
        if not text:
            return None
        style_name = ""
        try:
            style_name = (para.style.name or "") if para.style else ""
        except Exception:  # noqa: BLE001 — defensive against malformed styles
            style_name = ""
        escaped = _escape_md_text(text)
        # Title-style → top-level heading.
        if style_name.lower() == "title":
            return f"# {escaped}"
        if style_name.lower() == "subtitle":
            return f"## {escaped}"
        # Any case-variant of "Heading N" (incl. odd compounds like "heading 12").
        level = _heading_level_from_style(style_name)
        if level is not None:
            return f"{'#' * level} {escaped}"
        # Bold-only short paragraphs ⇒ promote to H2 so NCA templates that
        # use bold-runs (or style-inherited bold) instead of Heading
        # styles still read as a hierarchical document.
        if _is_visually_bold_heading(para):
            return f"## {escaped}"
        return escaped

    def _format_table(table) -> Optional[str]:
        """Render a Word table as a GFM markdown table.

        Merged cells in python-docx return the SAME cell object for every
        grid position they cover — which would otherwise produce
        `| A | A | A | B |` rows. We de-duplicate by tracking which
        underlying `cell._tc` (the lxml table-cell element) we've already
        emitted within a row, and emit empty strings for the repeats so
        the column count still lines up.
        """
        rows: List[List[str]] = []
        for row in table.rows:
            seen_in_row: set = set()
            cells: List[str] = []
            for cell in row.cells:
                tc = getattr(cell, "_tc", None)
                if tc is not None and id(tc) in seen_in_row:
                    cells.append("")
                    continue
                if tc is not None:
                    seen_in_row.add(id(tc))
                txt = (cell.text or "").strip()
                # Strip embedded newlines + escape any pipes that would
                # break the GFM row delimiter. Also escape angle brackets
                # so `<placeholder>` text in the template isn't silently
                # consumed by remark-gfm's HTML parser.
                txt = (
                    txt.replace("\n", " ")
                    .replace("\r", " ")
                    .replace("|", "\\|")
                    .replace("<", "\\<")
                    .replace(">", "\\>")
                )
                # Collapse runs of internal whitespace introduced by the
                # newline replacement.
                txt = re.sub(r"\s{2,}", " ", txt)
                cells.append(txt)
            if any(c for c in cells):
                rows.append(cells)
        if not rows:
            return None
        # Pad short rows so column count is uniform.
        width = max(len(r) for r in rows)
        rows = [r + [""] * (width - len(r)) for r in rows]
        header = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        for r in body_rows:
            lines.append("| " + " | ".join(r) + " |")
        return "\n".join(lines)

    chunks: List[str] = []

    # Preferred path — python-docx ≥ 1.0 exposes iter_inner_content()
    # which yields Paragraph + Table in document order, no lxml plumbing.
    iter_inner = getattr(doc, "iter_inner_content", None)
    if callable(iter_inner):
        for item in iter_inner():
            if isinstance(item, DocxParagraph):
                formatted = _format_paragraph(item)
                if formatted:
                    chunks.append(formatted)
            elif isinstance(item, DocxTable):
                formatted = _format_table(item)
                if formatted:
                    chunks.append(formatted)
        return "\n\n".join(chunks)

    # Fallback — older python-docx versions. Loses doc-order but keeps
    # headings + tables. Acceptable until the dependency is bumped.
    for para in doc.paragraphs:
        formatted = _format_paragraph(para)
        if formatted:
            chunks.append(formatted)
    for table in doc.tables:
        formatted = _format_table(table)
        if formatted:
            chunks.append(formatted)
    return "\n\n".join(chunks)


def _extract_xlsx_text(path: Path) -> str:
    """Flatten an .xlsx into newline-separated rows."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed")

    wb = load_workbook(str(path), read_only=True, data_only=True)
    chunks: List[str] = []
    for ws in wb.worksheets:
        chunks.append(f"## Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _extract_template_text(template_id: str) -> str:
    path = _template_path(template_id)
    if path.suffix.lower() == ".docx":
        return _extract_docx_text(path)
    if path.suffix.lower() == ".xlsx":
        return _extract_xlsx_text(path)
    raise HTTPException(status_code=415, detail=f"Unsupported template format: {path.suffix}")


def _split_template_into_sections(text: str) -> List[Dict[str, str]]:
    """Split the NCA template markdown into (heading, content) sections.

    Cheap regex — looks for markdown headings (# / ## / ### …) and
    splits the body between them. The preview UI uses these to render
    a per-section accordion. If no markdown headings exist (unlikely
    for NCA templates which are always section-numbered), returns the
    whole body as a single "Document" section.
    """
    import re
    if not text or not text.strip():
        return []
    lines = text.splitlines()
    sections: List[Dict[str, str]] = []
    current_heading: Optional[str] = None
    current_body: list[str] = []
    heading_re = re.compile(r"^\s*(#{1,6})\s+(.+?)\s*$")
    for line in lines:
        m = heading_re.match(line)
        if m:
            if current_heading is not None or current_body:
                sections.append({
                    "heading": current_heading or "Preamble",
                    "content": "\n".join(current_body).strip(),
                })
            current_heading = m.group(2).strip()
            current_body = []
        else:
            current_body.append(line)
    if current_heading is not None or current_body:
        sections.append({
            "heading": current_heading or "Document",
            "content": "\n".join(current_body).strip(),
        })
    return [s for s in sections if s["content"] or s["heading"]]


def _estimate_review_minutes(word_count: int) -> str:
    """Rough reading-rate estimate (200 wpm) for the UI banner."""
    if word_count <= 0:
        return "~1 minute"
    minutes = max(1, round(word_count / 200))
    if minutes == 1:
        return "~1 minute"
    if minutes < 60:
        return f"~{minutes} minutes"
    hours = minutes // 60
    rem = minutes % 60
    return f"~{hours}h{rem:02d}m" if rem else f"~{hours}h"


def _persist_document_from_text(
    *,
    db: Session,
    tenant_id: int,
    user: GRCUser,
    title: str,
    doc_type: str,
    classification: str,
    content: str,
    source_template_id: str,
    source_template_title: str,
):
    """Save the verbatim NCA template text as a new GovernanceDocument.

    Mirrors the post-AI persistence path so the document looks identical
    in the documents list regardless of how it was created. Audit log
    will tag it with the source_template_id so reviewers know the
    content came straight from the NCA library — no AI in the loop.
    """
    from ..models import GovernanceDocument
    doc = GovernanceDocument(
        tenant_id=tenant_id,
        title=title,
        content=content,
        doc_type=doc_type,
        classification=classification,
        status="draft",
        owner_id=getattr(user, "id", None),
        author_id=getattr(user, "id", None),
        tags=[f"nca-template:{source_template_id}", "verbatim"],
    )
    db.add(doc)
    db.flush()
    return doc


# ─── Schemas ────────────────────────────────────────────────────────────────

class CreateFromTemplateBody(BaseModel):
    title: Optional[str] = None
    classification: Optional[str] = "internal"
    doc_type: Optional[str] = None
    description: Optional[str] = None
    customizations: Optional[str] = None  # If set, AI rewrites the template with these tweaks


class AIDraftFromTemplateBody(BaseModel):
    title: str
    organization_context: Optional[str] = None
    additional_requirements: Optional[str] = None
    classification: Optional[str] = "internal"
    doc_type: Optional[str] = None
    save_as_document: bool = False  # If True, persists the result as a GovernanceDocument
    # Default behaviour for NCA templates: use the template text VERBATIM
    # as the draft content. Senior's ask — the NCA documents are official
    # control-text, AI re-drafting actively damages them. The legacy AI
    # path is still available by setting ai_redraft=True.
    ai_redraft: bool = False


class CompareBody(BaseModel):
    document_id: int


# ─── Helpers ────────────────────────────────────────────────────────────────

CATEGORY_TO_DOC_TYPE = {
    "Policy": "policy",
    "Standard": "standard",
    "Procedure": "procedure",
    "Program": "program",
    "Checklist": "checklist",
    "Form": "form",
    "Report": "report",
    "Cybersecurity Foundation": "policy",
    "Other": "policy",
}


def _next_document_code(tenant_id: int, doc_type: str, db: Session) -> str:
    prefix = (doc_type or "doc").upper()[:4]
    count = db.query(GovernanceDocument).filter(
        GovernanceDocument.tenant_id == tenant_id,
        GovernanceDocument.doc_type == doc_type,
    ).count()
    return f"{prefix}-{(count + 1):04d}"


# ─── Endpoints ──────────────────────────────────────────────────────────────

@router.get("")
def list_templates(
    user: GRCUser = Depends(require_auth),
):
    """Return all NCA template metadata grouped by category."""
    catalog = _get_catalog()
    by_category: Dict[str, List[Dict[str, Any]]] = {}
    for item in catalog:
        by_category.setdefault(item["category"], []).append(item)
    return {
        "total": len(catalog),
        "templates": catalog,
        "categories": [{"name": c, "count": len(items)} for c, items in sorted(by_category.items())],
    }


@router.get("/{template_id}/content")
def get_template_content(
    template_id: str,
    user: GRCUser = Depends(require_auth),
):
    """Return the extracted text content of a template for preview / editing."""
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")
    text = _extract_template_text(template_id)
    return {
        "id": template_id,
        "title": item["title"],
        "category": item["category"],
        "filename": item["filename"],
        "content": text,
        "word_count": len(text.split()),
    }


@router.get("/{template_id}/download")
def download_template(
    template_id: str,
    user: GRCUser = Depends(require_auth),
):
    """Stream the original .docx/.xlsx file for the user to download."""
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")
    return FileResponse(
        path=str(_template_path(template_id)),
        filename=item["filename"],
        media_type="application/octet-stream",
    )


@router.post("/{template_id}/create-document")
def create_document_from_template(
    template_id: str,
    body: CreateFromTemplateBody,
    db: Session = Depends(get_db),
    user: GRCUser = Depends(require_auth),
):
    """Create a new GovernanceDocument seeded from a template.

    If `customizations` is provided, AI rewrites the template content with
    those tweaks; otherwise the document is created with the raw template
    text and the user can edit it via the regular document editor.
    """
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")

    tenant_id = get_user_primary_tenant(user, db)
    base_content = _extract_template_text(template_id)
    final_content = base_content
    ai_used = False

    if body.customizations and OPENAI_API_KEY:
        try:
            from openai import OpenAI
            kwargs: Dict[str, Any] = {"api_key": OPENAI_API_KEY}
            if OPENAI_BASE_URL:
                kwargs["base_url"] = OPENAI_BASE_URL
            client = OpenAI(**kwargs)

            prompt = f"""You are a GRC documentation specialist. Take the following NCA Saudi cybersecurity template and adapt it per the customization instructions, preserving structure (headings, sections, numbered controls). Return only the revised document body in markdown.

TEMPLATE ({item['category']} — {item['title']}):
{base_content[:18000]}

CUSTOMIZATION INSTRUCTIONS:
{body.customizations}
"""

            completion = client.chat.completions.create(
                model=get_openai_model(),
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4,
            )
            final_content = (completion.choices[0].message.content or base_content).strip()
            ai_used = True
        except Exception:
            logger.exception("AI customization of NCA template failed; using raw template")

    doc_type = body.doc_type or CATEGORY_TO_DOC_TYPE.get(item["category"], "policy")
    title = body.title or item["title"]

    doc = GovernanceDocument(
        tenant_id=tenant_id,
        document_code=_next_document_code(tenant_id, doc_type, db),
        title=title,
        description=body.description or f"Created from NCA template: {item['title']}",
        content=final_content,
        doc_type=doc_type,
        classification=body.classification or "internal",
        current_version="1.0",
        status="draft",
        author_id=getattr(user, "id", None),
        owner_id=getattr(user, "id", None),
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "document_id": doc.id,
        "title": doc.title,
        "doc_type": doc.doc_type,
        "ai_customization_applied": ai_used,
        "source_template": {"id": template_id, "title": item["title"]},
    }


@router.post("/{template_id}/ai-draft")
def ai_draft_from_template(
    template_id: str,
    body: AIDraftFromTemplateBody,
    db: Session = Depends(get_db),
    user: GRCUser = Depends(require_auth),
):
    """Kick off async drafting using an NCA template as the structural anchor.

    Mirrors `/governance/documents/ai-draft` — returns a job id the
    frontend polls. The NCA template text is squirreled into the
    request payload as the parent-document context so the pipeline
    inherits NCA's clause numbering while still personalising for the
    tenant.
    """
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")

    tenant_id = get_user_primary_tenant(user, db)
    if not tenant_id:
        raise HTTPException(status_code=403, detail="No tenant context for drafting")

    # ── Default fast path: return the exact NCA template text as the draft.
    # No AI call, no job queue, no async polling. Senior's requirement:
    # NCA templates are official control-text that the regulator publishes
    # — re-drafting them with an LLM corrupts the wording. The frontend
    # accepts a "synchronous, no job_id" response shape; we use it here
    # so the existing UI just renders the result immediately.
    if not body.ai_redraft:
        template_text = _extract_template_text(template_id) or ""
        doc_type = body.doc_type or CATEGORY_TO_DOC_TYPE.get(item["category"], "policy")
        sections = _split_template_into_sections(template_text)
        word_count = len(template_text.split())
        if body.save_as_document:
            doc = _persist_document_from_text(
                db=db,
                tenant_id=tenant_id,
                user=user,
                title=body.title,
                doc_type=doc_type,
                classification=body.classification or "internal",
                content=template_text,
                source_template_id=template_id,
                source_template_title=item["title"],
            )
            db.commit()
            return {
                "document_id": doc.id,
                "synchronous": True,
                "source_template": {"id": template_id, "title": item["title"], "category": item["category"]},
                "generated_content": template_text,
                "suggested_title": body.title,
                "suggested_sections": sections,
                "framework_alignment": [],
                "word_count": word_count,
                "estimated_review_time": _estimate_review_minutes(word_count),
            }
        # Preview-only return — frontend opens its preview modal with this
        # payload (same shape AI flow returns on completion).
        return {
            "synchronous": True,
            "source_template": {"id": template_id, "title": item["title"], "category": item["category"]},
            "generated_content": template_text,
            "suggested_title": body.title,
            "suggested_sections": sections,
            "framework_alignment": [],
            "word_count": word_count,
            "estimated_review_time": _estimate_review_minutes(word_count),
        }

    # ── Legacy AI re-draft path (opt-in via ai_redraft=true) ─────────────
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=503, detail="AI not configured. Set AI_INTEGRATIONS_OPENAI_API_KEY.")

    reference = _extract_template_text(template_id) or ""
    parent_excerpt = reference.strip()
    if len(parent_excerpt) > 6000:
        parent_excerpt = parent_excerpt[:6000] + "\n...[truncated]"
    parent_blob = (
        f"Reference template title: {item['title']}\n"
        f"Reference template category: {item['category']}\n"
        f"Reference excerpt (use the structure/numbering as a guide but "
        f"substitute the tenant's actual values throughout):\n{parent_excerpt}"
    )
    if body.organization_context:
        parent_blob += f"\n\nUser-supplied organisation context: {body.organization_context}"
    if body.additional_requirements:
        parent_blob += f"\n\nAdditional requirements: {body.additional_requirements}"

    doc_type = body.doc_type or CATEGORY_TO_DOC_TYPE.get(item["category"], "policy")

    # Resolve tenant slug for TenantTask.
    from ..db import MasterSession
    from ..models import Tenant as MasterTenant
    master = MasterSession()
    try:
        row = master.query(MasterTenant.slug).filter(MasterTenant.id == tenant_id).first()
        tenant_slug = row[0] if row else None
    finally:
        master.close()
    if not tenant_slug:
        raise HTTPException(status_code=500, detail="Could not resolve tenant slug")

    from ..tasks.ai_drafting import create_job, dispatch_in_thread

    job_id = create_job(
        tenant_id=tenant_id,
        request_summary={
            "doc_type": doc_type,
            "title": body.title,
            "source_template": item["title"],
        },
    )
    request_payload: Dict[str, Any] = {
        "tenant_id": tenant_id,
        "doc_type": doc_type,
        "title": body.title,
        "description": body.additional_requirements,
        "framework_ids": [],
        # Stuff the NCA reference into the parent context channel — the
        # pipeline doesn't distinguish source.
        "parent_document_text": parent_blob,
        "source_template_id": template_id,
        "source_template_title": item["title"],
        "save_as_document": body.save_as_document,
        "classification": body.classification or "internal",
        "user_id": getattr(user, "id", None),
        "user_name": getattr(user, "display_name", None) or getattr(user, "username", None),
    }
    dispatch_in_thread(tenant_slug, job_id, request_payload)

    return {
        "job_id": job_id,
        "status": "queued",
        "source_template": {"id": template_id, "title": item["title"], "category": item["category"]},
        "poll_url": f"/governance/documents/ai-draft-jobs/{job_id}",
    }


@router.post("/{template_id}/compare")
def compare_with_document(
    template_id: str,
    body: CompareBody,
    db: Session = Depends(get_db),
    user: GRCUser = Depends(require_auth),
):
    """Return the template content + the user's document content for
    side-by-side rendering on the frontend, plus an optional AI gap
    analysis if the OpenAI key is configured.
    """
    item = _find_template(template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")

    tenant_id = get_user_primary_tenant(user, db)
    doc = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == body.document_id,
        GovernanceDocument.tenant_id == tenant_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    template_content = _extract_template_text(template_id)
    document_content = (doc.content or "").strip() or "(empty document)"

    gap_analysis: Optional[Dict[str, Any]] = None

    if OPENAI_API_KEY:
        try:
            from openai import OpenAI
            kwargs: Dict[str, Any] = {"api_key": OPENAI_API_KEY}
            if OPENAI_BASE_URL:
                kwargs["base_url"] = OPENAI_BASE_URL
            client = OpenAI(**kwargs)

            prompt = f"""Compare the user's document against the NCA reference template. Identify gaps, missing sections, and areas where the user document is weaker or stronger than the template. Be concise and specific.

NCA REFERENCE ({item['title']}):
{template_content[:8000]}

USER DOCUMENT ({doc.title}):
{document_content[:8000]}

Return strict JSON with keys:
- summary (2-3 sentences)
- missing_from_user_document (array of specific items present in template but absent in user doc)
- present_in_user_only (array of items in user doc not in template)
- alignment_score (0-100 integer, how well the user doc aligns with the template)
- recommended_additions (array of 3-7 specific clauses/sections to add)
"""

            completion = client.chat.completions.create(
                model=get_openai_model(),
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0.3,
            )
            gap_analysis = json.loads(completion.choices[0].message.content or "{}")
        except Exception:
            logger.exception("AI gap analysis failed; returning side-by-side without it")

    return {
        "template": {
            "id": template_id,
            "title": item["title"],
            "category": item["category"],
            "content": template_content,
        },
        "document": {
            "id": doc.id,
            "title": doc.title,
            "doc_type": doc.doc_type,
            "content": document_content,
        },
        "gap_analysis": gap_analysis,
    }
