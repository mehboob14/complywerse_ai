import os
import re
import json
import threading
from typing import List, Optional
from datetime import datetime
from difflib import SequenceMatcher
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from pydantic import BaseModel
from openai import OpenAI

from sqlalchemy import func
from ....models import (
    GovernanceDocument, GovernanceDocumentVersion, PolicyStatement,
    PolicyStatementCompliance, PolicyStatementVersion, InternalControl, GRCUser, AuditLog, get_db, SessionLocal
)
from ....routers.auth_router import require_auth, get_user_tenants
from ....schemas import ConvertStatementsRequest, InternalControlFromStatementResponse
from ..action_logger import log_governance_action

router = APIRouter(prefix="/documents", tags=["Governance - Policy Parser"])

MAX_PARSE_CHUNKS = 12
MAX_PARSE_CHARACTERS = 180000

AI_INTEGRATIONS_OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
AI_INTEGRATIONS_OPENAI_BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")


def _sanitize_policy_text(raw_text: str) -> str:
    if not raw_text:
        return ""

    text = raw_text
    text = re.sub(r"<script\b[^<]*(?:(?!</script>)<[^<]*)*</script>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<style\b[^<]*(?:(?!</style>)<[^<]*)*</style>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class ParsedStatementResponse(BaseModel):
    id: int
    document_id: int
    statement_code: Optional[str]
    statement_text: str
    statement_summary: Optional[str]
    category: Optional[str]
    priority: str
    is_mandatory: bool
    source_section: Optional[str]
    ai_confidence: Optional[float]
    ai_extracted_keywords: List[str]
    compliance_id: Optional[int]

    class Config:
        from_attributes = True


class CreatePolicyStatementRequest(BaseModel):
    statement_text: str
    statement_summary: Optional[str] = None
    category: Optional[str] = None
    priority: str = "medium"
    is_mandatory: bool = True
    source_section: Optional[str] = None
    ai_extracted_keywords: Optional[List[str]] = None


class UpdatePolicyStatementRequest(BaseModel):
    statement_text: Optional[str] = None
    statement_summary: Optional[str] = None
    category: Optional[str] = None
    priority: Optional[str] = None
    is_mandatory: Optional[bool] = None
    source_section: Optional[str] = None
    change_reason: Optional[str] = None


class RollbackStatementRequest(BaseModel):
    version_id: int


class ApplyReparseProposalsRequest(BaseModel):
    decisions: List[dict] = []


def validate_document_access(user: GRCUser, document: GovernanceDocument, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if document.tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this document"
        )


def extract_text_from_file(file_path: str, file_type: str) -> str:
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document file not found on disk"
        )
    
    extracted_text = ""
    
    if file_type == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted_text = "\n\n".join(text_parts)
    
    elif file_type in ("docx", "doc"):
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        extracted_text = "\n".join(text_parts)
    
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {file_type}. Only PDF and Word documents are supported."
        )
    
    return extracted_text


def _split_text_into_chunks(text: str, chunk_size: int = 20000, overlap: int = 500) -> List[str]:
    heading_pattern = re.compile(
        r'^(?:'
        r'(?:Section|SECTION|Chapter|CHAPTER|Article|ARTICLE|Part|PART)\s+[\dIVXivx]+'
        r'|[\d]+(?:\.[\d]+)*\s+[A-Z]'
        r'|[A-Z][A-Z\s]{5,}$'
        r')',
        re.MULTILINE
    )

    headings = list(heading_pattern.finditer(text))

    if len(headings) >= 3:
        sections = []
        for i, match in enumerate(headings):
            start = match.start()
            end = headings[i + 1].start() if i + 1 < len(headings) else len(text)
            section_text = text[start:end].strip()
            if section_text:
                sections.append(section_text)

        chunks = []
        current_chunk = ""
        for section in sections:
            if len(current_chunk) + len(section) > chunk_size and current_chunk:
                chunks.append(current_chunk)
                current_chunk = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk += "\n\n" + section
            else:
                current_chunk = (current_chunk + "\n\n" + section).strip() if current_chunk else section

        if current_chunk.strip():
            chunks.append(current_chunk)

        if len(chunks) >= 2:
            return chunks
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            break_point = text.rfind('\n\n', start, end)
            if break_point == -1 or break_point <= start:
                break_point = text.rfind('\n', start, end)
            if break_point == -1 or break_point <= start:
                break_point = text.rfind('. ', start, end)
            if break_point > start:
                end = break_point + 1
        else:
            end = len(text)

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = max(start + 1, end - overlap)

    return chunks if chunks else [text]


def _is_duplicate(new_text: str, existing_statements: List[dict], threshold: float = 0.80) -> bool:
    new_lower = new_text.lower().strip()
    for existing in existing_statements:
        existing_lower = existing.get("statement_text", "").lower().strip()
        if new_lower == existing_lower:
            return True
        if len(new_lower) > 20 and len(existing_lower) > 20:
            ratio = SequenceMatcher(None, new_lower, existing_lower).ratio()
            if ratio > threshold:
                return True
    return False


def parse_policy_statements_with_openai(text: str, document_title: str, progress_callback=None) -> List[dict]:
    if not AI_INTEGRATIONS_OPENAI_API_KEY or not AI_INTEGRATIONS_OPENAI_BASE_URL:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI integration not configured"
        )
    
    client = OpenAI(
        api_key=AI_INTEGRATIONS_OPENAI_API_KEY,
        base_url=AI_INTEGRATIONS_OPENAI_BASE_URL
    )

    cleaned_text = _sanitize_policy_text(text)
    if not cleaned_text:
        return []

    if len(cleaned_text) > MAX_PARSE_CHARACTERS:
        cleaned_text = cleaned_text[:MAX_PARSE_CHARACTERS]

    chunks = _split_text_into_chunks(cleaned_text)
    if len(chunks) > MAX_PARSE_CHUNKS:
        chunks = chunks[:MAX_PARSE_CHUNKS]

    all_statements: List[dict] = []

    system_message = (
        "You are a governance expert that performs EXHAUSTIVE extraction of every policy statement. "
        "You must capture every requirement, obligation, prohibition, and directive in the document. "
        "CRITICAL: You must copy the EXACT text from the document verbatim - never paraphrase, summarize, or rewrite any statement. "
        "Each extracted statement_text must be a direct quote from the source document. "
        "Never skip or combine statements. Always respond with valid JSON."
    )

    total_chunks = len(chunks)

    for chunk_idx, chunk_text in enumerate(chunks):
        current_chunk = chunk_idx + 1
        if progress_callback:
            progress_callback(current_chunk=current_chunk, total_chunks=total_chunks, phase="processing")

        prompt = f"""You are a governance and compliance expert analyzing policy documents to extract individual policy statements.

Analyze the following section from the document titled "{document_title}" and extract ALL distinct policy statements.

Extract EVERY individual policy statement, requirement, obligation, and mandate. Copy the EXACT wording from the document - do NOT paraphrase, summarize, rewrite, or generalize. Each extracted statement must be a direct quote from the document text. Do NOT combine statements. Each sentence that contains 'shall', 'must', 'should', 'will', 'required', 'responsible', 'prohibited', or similar directive language is a separate statement. A typical policy document contains 30-100+ statements.

You must extract AT LEAST 15 statements from this section if it contains policy language.

IMPORTANT: Extract EVERY distinct statement. Look for ALL sentences containing directive language (shall, must, should, will, required, prohibited, responsible for, etc.). Include statements from tables, appendices, and all subsections. Each distinct requirement is a separate statement.

For each policy statement you find, extract:
1. statement_text: Copy the EXACT text from the document - do NOT paraphrase, summarize, or rewrite. Use the precise wording as it appears in the document. If a statement spans multiple sentences, include all of them exactly as written.
2. statement_summary: A brief 1-2 sentence summary of what the statement requires
3. category: Categorize into one of: security, privacy, operational, compliance, governance, risk_management, hr, it, financial, legal, environmental, quality
4. priority: "critical" for mandatory security/compliance items, "high" for important requirements, "medium" for standard policies, "low" for guidelines
5. is_mandatory: true if the statement uses "shall", "must", "required", "mandatory"; false if "should", "may", "recommended", "encouraged"
6. source_section: The section or chapter title/number where this statement appears (e.g., "Section 4.2 - Access Control", "Chapter 3")
7. ai_confidence: Your confidence score from 0.0 to 1.0 in the accuracy of this extraction
8. ai_extracted_keywords: Array of 3-5 key terms/concepts from this statement (e.g., ["access control", "authentication", "password"])

Document section (part {chunk_idx + 1} of {len(chunks)}) to analyze:
---
{chunk_text}
---

Return a JSON object with a "statements" array containing all extracted policy statements. Example format:
{{
  "statements": [
    {{
      "statement_text": "All employees must complete information security awareness training within 30 days of hire.",
      "statement_summary": "Mandatory security training for new employees within 30 days",
      "category": "security",
      "priority": "high",
      "is_mandatory": true,
      "source_section": "Section 5.1 - Security Awareness",
      "ai_confidence": 0.95,
      "ai_extracted_keywords": ["security training", "awareness", "new employees", "onboarding"]
    }}
  ]
}}"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=16384
            )
            
            result_text = response.choices[0].message.content or "{}"
            result = json.loads(result_text)
            chunk_statements = result.get("statements", [])

            for stmt in chunk_statements:
                stmt_text = stmt.get("statement_text", "").strip()
                if stmt_text and not _is_duplicate(stmt_text, all_statements):
                    all_statements.append(stmt)

            if progress_callback:
                progress_callback(current_chunk=current_chunk, total_chunks=total_chunks, phase="processed")
        
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to parse OpenAI response: {str(e)}"
            )
        except Exception as e:
            error_msg = str(e)
            if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
                raise HTTPException(
                    status_code=status.HTTP_402_PAYMENT_REQUIRED,
                    detail="Cloud budget exceeded. Please upgrade your plan."
                )
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"OpenAI API error: {error_msg}"
            )

    return all_statements


_parsing_status = {}
_reparse_proposals = {}


def serialize_statement(statement: PolicyStatement, compliance_id: Optional[int] = None) -> dict:
    return {
        "id": statement.id,
        "document_id": statement.document_id,
        "document_version_id": statement.document_version_id,
        "statement_code": statement.statement_code,
        "statement_text": statement.statement_text,
        "statement_summary": statement.statement_summary,
        "category": statement.category,
        "priority": statement.priority,
        "is_mandatory": statement.is_mandatory,
        "source_section": statement.source_section,
        "source_page": statement.source_page,
        "ai_confidence": statement.ai_confidence,
        "ai_extracted_keywords": statement.ai_extracted_keywords or [],
        "ai_suggested_controls": statement.ai_suggested_controls or [],
        "status": statement.status,
        "effective_date": statement.effective_date.isoformat() if statement.effective_date else None,
        "created_at": statement.created_at.isoformat() if statement.created_at else None,
        "updated_at": statement.updated_at.isoformat() if statement.updated_at else None,
        "current_version": len(statement.versions) if hasattr(statement, 'versions') and statement.versions else 1,
        "created_by": statement.created_by,
        "compliance_id": compliance_id
    }


def _create_version_snapshot(db: Session, statement: PolicyStatement, change_type: str, changed_by: int, change_reason: str = None):
    max_version = db.query(func.max(PolicyStatementVersion.version_number)).filter(
        PolicyStatementVersion.statement_id == statement.id
    ).scalar() or 0
    
    version = PolicyStatementVersion(
        tenant_id=statement.tenant_id,
        statement_id=statement.id,
        version_number=max_version + 1,
        statement_text=statement.statement_text,
        statement_summary=statement.statement_summary,
        category=statement.category,
        sub_category=statement.sub_category,
        priority=statement.priority,
        is_mandatory=statement.is_mandatory,
        source_section=statement.source_section,
        source_page=statement.source_page,
        ai_confidence=statement.ai_confidence,
        ai_extracted_keywords=statement.ai_extracted_keywords or [],
        status=statement.status,
        change_type=change_type,
        change_reason=change_reason,
        changed_by=changed_by,
        changed_at=datetime.utcnow()
    )
    db.add(version)
    return version


@router.get("/{document_id}/parse-status")
def get_parse_status(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    status_info = _parsing_status.get(document_id, {"status": "idle"})
    return status_info


def _parse_policy_background(document_id: int, user_id: int):
    db = SessionLocal()
    try:
        document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
        if not document:
            _parsing_status[document_id] = {"status": "failed", "error": "Document not found"}
            return

        file_path = document.file_path
        file_type = document.file_type

        extracted_text = None

        # Prefer source file extraction for uploaded documents to avoid parsing rendered HTML blobs.
        if file_path and os.path.exists(file_path):
            try:
                extracted_text = extract_text_from_file(file_path, file_type)
            except Exception:
                extracted_text = None

        if (not extracted_text or not extracted_text.strip()) and document.content and document.content.strip():
            extracted_text = document.content

        if not extracted_text or not extracted_text.strip():
            _parsing_status[document_id] = {"status": "failed", "error": "Document has no content or file to parse"}
            return

        try:
            _parsing_status[document_id] = {
                "status": "parsing",
                "message": "Extracting policy statements with AI...",
                "processed_chunks": 0,
                "total_chunks": 0,
                "current_chunk": 0,
                "progress_percent": 0,
            }

            def _update_chunk_progress(current_chunk: int, total_chunks: int, phase: str = "processing"):
                processed_chunks = max(current_chunk - 1, 0) if phase == "processing" else current_chunk
                percent = int((processed_chunks / total_chunks) * 100) if total_chunks else 0
                _parsing_status[document_id] = {
                    "status": "parsing",
                    "message": f"Parsing chunk {current_chunk}/{total_chunks}...",
                    "processed_chunks": processed_chunks,
                    "total_chunks": total_chunks,
                    "current_chunk": current_chunk,
                    "progress_percent": percent,
                }

            parsed_statements_data = parse_policy_statements_with_openai(
                extracted_text,
                document.title,
                progress_callback=_update_chunk_progress,
            )

            if not parsed_statements_data:
                total_chunks = int(_parsing_status.get(document_id, {}).get("total_chunks") or 0)
                _parsing_status[document_id] = {
                    "status": "completed",
                    "total_statements": 0,
                    "message": "No statements found",
                    "processed_chunks": total_chunks,
                    "total_chunks": total_chunks,
                    "current_chunk": total_chunks,
                    "progress_percent": 100,
                }
                return

            version_id = None
            from ....models import DocumentVersion
            latest_version = db.query(DocumentVersion).filter(
                DocumentVersion.document_id == document_id
            ).order_by(DocumentVersion.version_number.desc()).first()
            if latest_version:
                version_id = latest_version.id

            existing_statements = db.query(PolicyStatement).filter(
                PolicyStatement.document_id == document_id
            ).all()

            if existing_statements:
                proposals = []
                existing_map = {}
                for s in existing_statements:
                    key = (s.source_section or "").lower().strip()
                    existing_map[key] = s

                for idx, new_data in enumerate(parsed_statements_data):
                    new_section = (new_data.get("source_section") or "").lower().strip()
                    matched_existing = existing_map.get(new_section)

                    if not matched_existing:
                        best_match = None
                        best_ratio = 0
                        for es in existing_statements:
                            ratio = SequenceMatcher(None, es.statement_text[:200].lower(), new_data.get("statement_text", "")[:200].lower()).ratio()
                            if ratio > best_ratio and ratio > 0.6:
                                best_ratio = ratio
                                best_match = es
                        matched_existing = best_match

                    proposal = {
                        "index": idx,
                        "new_statement": new_data,
                        "status": "pending",
                    }

                    if matched_existing:
                        proposal["type"] = "update"
                        proposal["existing_statement_id"] = matched_existing.id
                        proposal["existing_text"] = matched_existing.statement_text
                        proposal["similarity_score"] = SequenceMatcher(None, matched_existing.statement_text[:200].lower(), new_data.get("statement_text", "")[:200].lower()).ratio()
                    else:
                        proposal["type"] = "new"
                        proposal["existing_statement_id"] = None
                        proposal["existing_text"] = None
                        proposal["similarity_score"] = 0

                    proposals.append(proposal)

                _reparse_proposals[document_id] = {
                    "proposals": proposals,
                    "created_at": datetime.utcnow().isoformat(),
                    "total": len(proposals),
                    "new_count": sum(1 for p in proposals if p["type"] == "new"),
                    "update_count": sum(1 for p in proposals if p["type"] == "update"),
                }

                _parsing_status[document_id] = {
                    "status": "review_required",
                    "total_statements": len(parsed_statements_data),
                    "message": f"Found {len(proposals)} proposed changes. Review required before applying.",
                    "has_proposals": True,
                    "processed_chunks": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "total_chunks": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "current_chunk": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "progress_percent": 100,
                }
            else:
                for idx, statement_data in enumerate(parsed_statements_data, start=1):
                    statement_code = f"PS-{document_id:04d}-{idx:03d}"
                    policy_statement = PolicyStatement(
                        tenant_id=document.tenant_id,
                        document_id=document_id,
                        document_version_id=version_id,
                        statement_code=statement_code,
                        statement_text=statement_data.get("statement_text", "")[:10000],
                        statement_summary=statement_data.get("statement_summary", "")[:500] if statement_data.get("statement_summary") else None,
                        category=statement_data.get("category"),
                        priority=statement_data.get("priority", "medium"),
                        is_mandatory=statement_data.get("is_mandatory", True),
                        source_section=statement_data.get("source_section"),
                        ai_confidence=statement_data.get("ai_confidence"),
                        ai_extracted_keywords=statement_data.get("ai_extracted_keywords", []),
                        status="active",
                        created_by=user_id
                    )
                    db.add(policy_statement)
                    db.flush()

                    _create_version_snapshot(db, policy_statement, "initial_parse", user_id)

                    compliance_record = PolicyStatementCompliance(
                        tenant_id=document.tenant_id,
                        statement_id=policy_statement.id,
                        compliance_status="not_assessed",
                        owner_id=document.owner_id
                    )
                    db.add(compliance_record)

                db.commit()
                _parsing_status[document_id] = {
                    "status": "completed",
                    "total_statements": len(parsed_statements_data),
                    "message": f"Successfully parsed {len(parsed_statements_data)} policy statements",
                    "processed_chunks": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "total_chunks": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "current_chunk": _parsing_status.get(document_id, {}).get("total_chunks", 0),
                    "progress_percent": 100,
                }
        except Exception as e:
            db.rollback()
            _parsing_status[document_id] = {"status": "failed", "error": str(e)[:500]}
    finally:
        db.close()


@router.post("/{document_id}/parse-policy")
def parse_policy_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")

    current_status = _parsing_status.get(document_id, {})
    if current_status.get("status") == "parsing":
        return {"status": "parsing", "message": "Document is already being parsed"}

    _parsing_status[document_id] = {"status": "parsing", "message": "Starting policy statement extraction..."}
    _parsing_status[document_id].update({
        "processed_chunks": 0,
        "total_chunks": 0,
        "current_chunk": 0,
        "progress_percent": 0,
    })

    thread = threading.Thread(
        target=_parse_policy_background,
        args=(document_id, current_user.id),
        daemon=True
    )
    thread.start()

    return {
        "status": "parsing",
        "message": "Policy parsing started in background. Check status for results.",
        "document_id": document_id
    }


@router.get("/{document_id}/policy-statements")
def get_document_policy_statements(
    document_id: int,
    category: Optional[str] = None,
    priority: Optional[str] = None,
    is_mandatory: Optional[bool] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Governance document not found"
        )
    
    validate_document_access(current_user, document, db)
    
    query = db.query(PolicyStatement).filter(
        PolicyStatement.document_id == document_id
    )
    
    if category:
        query = query.filter(PolicyStatement.category == category)
    if priority:
        query = query.filter(PolicyStatement.priority == priority)
    if is_mandatory is not None:
        query = query.filter(PolicyStatement.is_mandatory == is_mandatory)
    
    statements = query.order_by(PolicyStatement.statement_code).all()
    
    result = []
    for statement in statements:
        compliance = db.query(PolicyStatementCompliance).filter(
            PolicyStatementCompliance.statement_id == statement.id
        ).first()
        result.append(serialize_statement(statement, compliance.id if compliance else None))
    
    return {
        "document_id": document_id,
        "document_title": document.title,
        "statements": result,
        "total_statements": len(result)
    }


@router.post("/{document_id}/statements")
def add_policy_statement(
    document_id: int,
    request: CreatePolicyStatementRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"[ADD STATEMENT] Received request: document_id={document_id}, user_id={current_user.id}")
    logger.info(f"[ADD STATEMENT] Request data: {request.dict()}")
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id
    ).first()

    if not document:
        logger.error(f"[ADD STATEMENT] Document {document_id} not found")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Governance document not found"
        )

    validate_document_access(current_user, document, db)

    if not request.statement_text or not request.statement_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="statement_text is required"
        )

    existing_count = db.query(PolicyStatement).filter(
        PolicyStatement.document_id == document_id
    ).count()
    next_idx = existing_count + 1
    statement_code = f"PS-{document_id:04d}-{next_idx:03d}"

    statement = PolicyStatement(
        tenant_id=document.tenant_id,
        document_id=document_id,
        statement_code=statement_code,
        statement_text=request.statement_text[:10000],
        statement_summary=request.statement_summary[:500] if request.statement_summary else None,
        category=request.category,
        priority=request.priority or "medium",
        is_mandatory=request.is_mandatory,
        source_section=request.source_section,
        ai_extracted_keywords=request.ai_extracted_keywords or [],
        status="active",
        created_by=current_user.id,
        created_at=datetime.utcnow()
    )
    db.add(statement)
    db.flush()

    _create_version_snapshot(db, statement, "manual_add", current_user.id)

    compliance_record = PolicyStatementCompliance(
        tenant_id=document.tenant_id,
        statement_id=statement.id,
        compliance_status="not_assessed",
        owner_id=document.owner_id
    )
    db.add(compliance_record)
    db.commit()
    db.refresh(statement)

    # Log the action for review
    log_governance_action(
        db=db,
        tenant_id=document.tenant_id,
        action_type="policy_statement_created",
        action_description=f"Policy statement '{statement.statement_code}' created: {statement.statement_text[:100]}...",
        entity_type="policy_statement",
        action_user_id=current_user.id,
        entity_id=statement.id,
        action_metadata={
            "document_id": document_id,
            "statement_code": statement.statement_code,
            "category": statement.category,
            "priority": statement.priority
        }
    )
    db.commit()

    logger.info(f"[ADD STATEMENT] Statement created successfully: id={statement.id}, code={statement.statement_code}")
    return serialize_statement(statement, compliance_record.id)


@router.put("/{document_id}/statements/{statement_id}")
def update_policy_statement(
    document_id: int,
    statement_id: int,
    request: UpdatePolicyStatementRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Governance document not found"
        )

    validate_document_access(current_user, document, db)

    statement = db.query(PolicyStatement).filter(
        PolicyStatement.id == statement_id,
        PolicyStatement.document_id == document_id
    ).first()

    if not statement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Policy statement not found for this document"
        )

    _create_version_snapshot(db, statement, "manual_edit", current_user.id, request.change_reason)

    if request.statement_text is not None:
        statement.statement_text = request.statement_text[:10000]
    if request.statement_summary is not None:
        statement.statement_summary = request.statement_summary[:500] if request.statement_summary else None
    if request.category is not None:
        statement.category = request.category
    if request.priority is not None:
        statement.priority = request.priority
    if request.is_mandatory is not None:
        statement.is_mandatory = request.is_mandatory
    if request.source_section is not None:
        statement.source_section = request.source_section

    statement.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(statement)

    compliance = db.query(PolicyStatementCompliance).filter(
        PolicyStatementCompliance.statement_id == statement.id
    ).first()

    return serialize_statement(statement, compliance.id if compliance else None)


@router.get("/{document_id}/statements/{statement_id}/versions")
def get_statement_versions(
    document_id: int,
    statement_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    validate_document_access(current_user, document, db)

    statement = db.query(PolicyStatement).filter(
        PolicyStatement.id == statement_id,
        PolicyStatement.document_id == document_id
    ).first()
    if not statement:
        raise HTTPException(status_code=404, detail="Statement not found")

    versions = db.query(PolicyStatementVersion).filter(
        PolicyStatementVersion.statement_id == statement_id
    ).order_by(PolicyStatementVersion.version_number.desc()).all()

    result = []
    for v in versions:
        changer_name = None
        if v.changed_by:
            changer = db.query(GRCUser).filter(GRCUser.id == v.changed_by).first()
            if changer:
                changer_name = changer.display_name
        result.append({
            "id": v.id,
            "version_number": v.version_number,
            "statement_text": v.statement_text,
            "statement_summary": v.statement_summary,
            "category": v.category,
            "priority": v.priority,
            "is_mandatory": v.is_mandatory,
            "source_section": v.source_section,
            "ai_confidence": v.ai_confidence,
            "ai_extracted_keywords": v.ai_extracted_keywords or [],
            "status": v.status,
            "change_type": v.change_type,
            "change_reason": v.change_reason,
            "changed_by": v.changed_by,
            "changed_by_name": changer_name,
            "changed_at": v.changed_at.isoformat() if v.changed_at else None,
        })

    return {
        "statement_id": statement_id,
        "current_text": statement.statement_text,
        "versions": result,
        "total_versions": len(result)
    }


@router.post("/{document_id}/statements/{statement_id}/rollback")
def rollback_statement(
    document_id: int,
    statement_id: int,
    request: RollbackStatementRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    validate_document_access(current_user, document, db)

    statement = db.query(PolicyStatement).filter(
        PolicyStatement.id == statement_id,
        PolicyStatement.document_id == document_id
    ).first()
    if not statement:
        raise HTTPException(status_code=404, detail="Statement not found")

    target_version_id = request.version_id
    if not target_version_id:
        raise HTTPException(status_code=400, detail="version_id is required")

    target_version = db.query(PolicyStatementVersion).filter(
        PolicyStatementVersion.id == target_version_id,
        PolicyStatementVersion.statement_id == statement_id
    ).first()
    if not target_version:
        raise HTTPException(status_code=404, detail="Version not found")

    _create_version_snapshot(db, statement, "rollback", current_user.id,
                           f"Rolled back to version {target_version.version_number}")

    statement.statement_text = target_version.statement_text
    statement.statement_summary = target_version.statement_summary
    statement.category = target_version.category
    statement.sub_category = target_version.sub_category
    statement.priority = target_version.priority
    statement.is_mandatory = target_version.is_mandatory
    statement.source_section = target_version.source_section
    statement.source_page = target_version.source_page
    statement.ai_confidence = target_version.ai_confidence
    statement.ai_extracted_keywords = target_version.ai_extracted_keywords
    statement.status = target_version.status
    statement.updated_at = datetime.utcnow()

    audit = AuditLog(
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="statement_rollback",
        resource_type="PolicyStatement",
        resource_id=str(statement.id),
        details=json.dumps({
            "document_id": document_id,
            "statement_code": statement.statement_code,
            "rolled_back_to_version": target_version.version_number,
            "target_version_id": target_version_id,
        })
    )
    db.add(audit)
    db.commit()
    db.refresh(statement)

    compliance = db.query(PolicyStatementCompliance).filter(
        PolicyStatementCompliance.statement_id == statement.id
    ).first()

    return {
        "message": f"Statement rolled back to version {target_version.version_number}",
        "statement": serialize_statement(statement, compliance.id if compliance else None)
    }


@router.get("/{document_id}/statements/{statement_id}/diff")
def compare_versions(
    document_id: int,
    statement_id: int,
    version_a: int,
    version_b: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    validate_document_access(current_user, document, db)

    va = db.query(PolicyStatementVersion).filter(
        PolicyStatementVersion.id == version_a,
        PolicyStatementVersion.statement_id == statement_id
    ).first()
    vb = db.query(PolicyStatementVersion).filter(
        PolicyStatementVersion.id == version_b,
        PolicyStatementVersion.statement_id == statement_id
    ).first()

    if not va or not vb:
        raise HTTPException(status_code=404, detail="One or both versions not found")

    fields_to_compare = ["statement_text", "statement_summary", "category", "priority", "is_mandatory", "source_section", "status"]
    changes = []
    for field in fields_to_compare:
        val_a = getattr(va, field)
        val_b = getattr(vb, field)
        if val_a != val_b:
            changes.append({
                "field": field,
                "version_a_value": str(val_a) if val_a is not None else None,
                "version_b_value": str(val_b) if val_b is not None else None,
            })

    import difflib
    text_diff = list(difflib.unified_diff(
        (va.statement_text or "").splitlines(keepends=True),
        (vb.statement_text or "").splitlines(keepends=True),
        fromfile=f"Version {va.version_number}",
        tofile=f"Version {vb.version_number}",
        lineterm=""
    ))

    return {
        "version_a": {"id": va.id, "version_number": va.version_number, "changed_at": va.changed_at.isoformat() if va.changed_at else None, "change_type": va.change_type},
        "version_b": {"id": vb.id, "version_number": vb.version_number, "changed_at": vb.changed_at.isoformat() if vb.changed_at else None, "change_type": vb.change_type},
        "field_changes": changes,
        "text_diff": text_diff,
        "total_changes": len(changes)
    }


@router.get("/{document_id}/reparse-proposals")
def get_reparse_proposals(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    validate_document_access(current_user, document, db)

    proposals = _reparse_proposals.get(document_id)
    if not proposals:
        return {"document_id": document_id, "proposals": [], "total": 0}

    return {"document_id": document_id, **proposals}


@router.post("/{document_id}/reparse-proposals/apply")
def apply_reparse_proposals(
    document_id: int,
    request: ApplyReparseProposalsRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(GovernanceDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    validate_document_access(current_user, document, db)

    proposals_data = _reparse_proposals.get(document_id)
    if not proposals_data:
        raise HTTPException(status_code=404, detail="No pending proposals found")

    decisions = request.decisions

    accepted = 0
    rejected = 0

    version_id = None
    from ....models import DocumentVersion
    latest_version = db.query(DocumentVersion).filter(
        DocumentVersion.document_id == document_id
    ).order_by(DocumentVersion.version_number.desc()).first()
    if latest_version:
        version_id = latest_version.id

    for decision in decisions:
        idx = decision.get("index")
        action = decision.get("action")

        if idx is None or idx >= len(proposals_data["proposals"]):
            continue

        proposal = proposals_data["proposals"][idx]

        if action == "reject":
            proposal["status"] = "rejected"
            rejected += 1
            continue

        if action == "accept":
            proposal["status"] = "accepted"
            new_data = proposal["new_statement"]

            if proposal["type"] == "update" and proposal.get("existing_statement_id"):
                statement = db.query(PolicyStatement).filter(
                    PolicyStatement.id == proposal["existing_statement_id"]
                ).first()
                if statement:
                    _create_version_snapshot(db, statement, "ai_reparse", current_user.id, "Re-parsed from document")

                    statement.statement_text = new_data.get("statement_text", statement.statement_text)[:10000]
                    statement.statement_summary = new_data.get("statement_summary", statement.statement_summary)
                    if new_data.get("statement_summary"):
                        statement.statement_summary = new_data["statement_summary"][:500]
                    statement.category = new_data.get("category", statement.category)
                    statement.priority = new_data.get("priority", statement.priority)
                    statement.is_mandatory = new_data.get("is_mandatory", statement.is_mandatory)
                    statement.source_section = new_data.get("source_section", statement.source_section)
                    statement.ai_confidence = new_data.get("ai_confidence", statement.ai_confidence)
                    statement.ai_extracted_keywords = new_data.get("ai_extracted_keywords", statement.ai_extracted_keywords)
                    statement.updated_at = datetime.utcnow()
                    accepted += 1

            elif proposal["type"] == "new":
                existing_count = db.query(PolicyStatement).filter(
                    PolicyStatement.document_id == document_id
                ).count()
                next_idx = existing_count + 1
                statement_code = f"PS-{document_id:04d}-{next_idx:03d}"

                new_statement = PolicyStatement(
                    tenant_id=document.tenant_id,
                    document_id=document_id,
                    document_version_id=version_id,
                    statement_code=statement_code,
                    statement_text=new_data.get("statement_text", "")[:10000],
                    statement_summary=new_data.get("statement_summary", "")[:500] if new_data.get("statement_summary") else None,
                    category=new_data.get("category"),
                    priority=new_data.get("priority", "medium"),
                    is_mandatory=new_data.get("is_mandatory", True),
                    source_section=new_data.get("source_section"),
                    ai_confidence=new_data.get("ai_confidence"),
                    ai_extracted_keywords=new_data.get("ai_extracted_keywords", []),
                    status="active",
                    created_by=current_user.id
                )
                db.add(new_statement)
                db.flush()

                _create_version_snapshot(db, new_statement, "ai_reparse", current_user.id, "Added from re-parse")

                compliance_record = PolicyStatementCompliance(
                    tenant_id=document.tenant_id,
                    statement_id=new_statement.id,
                    compliance_status="not_assessed",
                    owner_id=document.owner_id
                )
                db.add(compliance_record)
                accepted += 1

    audit = AuditLog(
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="reparse_proposals_applied",
        resource_type="GovernanceDocument",
        resource_id=str(document_id),
        details=json.dumps({
            "accepted": accepted,
            "rejected": rejected,
            "total_proposals": len(decisions),
        })
    )
    db.add(audit)
    db.commit()

    all_resolved = all(p["status"] != "pending" for p in proposals_data["proposals"])
    if all_resolved:
        del _reparse_proposals[document_id]
        _parsing_status[document_id] = {
            "status": "completed",
            "total_statements": accepted,
            "message": f"Applied {accepted} changes, rejected {rejected}"
        }

    return {
        "message": f"Processed {accepted + rejected} proposals: {accepted} accepted, {rejected} rejected",
        "accepted": accepted,
        "rejected": rejected,
        "all_resolved": all_resolved
    }


@router.delete("/{document_id}/statements/{statement_id}")
def delete_policy_statement(
    document_id: int,
    statement_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Governance document not found"
        )

    validate_document_access(current_user, document, db)

    statement = db.query(PolicyStatement).filter(
        PolicyStatement.id == statement_id,
        PolicyStatement.document_id == document_id
    ).first()

    if not statement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Policy statement not found for this document"
        )

    db.query(PolicyStatementCompliance).filter(
        PolicyStatementCompliance.statement_id == statement_id
    ).delete(synchronize_session='fetch')

    db.delete(statement)
    db.commit()

    return {"message": "Statement deleted"}


@router.post("/{document_id}/statements/convert-to-controls", response_model=List[InternalControlFromStatementResponse])
def convert_statements_to_controls(
    document_id: int,
    request: ConvertStatementsRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Governance document not found"
        )
    
    validate_document_access(current_user, document, db)
    
    statements = db.query(PolicyStatement).filter(
        PolicyStatement.id.in_(request.statement_ids),
        PolicyStatement.document_id == document_id
    ).all()
    
    if not statements:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No valid statements found for conversion"
        )
    
    existing_count = db.query(InternalControl).filter(
        InternalControl.tenant_id == document.tenant_id
    ).count()
    
    created_controls = []
    for idx, statement in enumerate(statements):
        control_num = existing_count + idx + 1
        control_id = f"IC-{document.tenant_id}-{control_num:04d}"
        
        name = (statement.statement_summary or statement.statement_text)[:200]
        
        category = request.category or statement.category
        priority = request.priority or statement.priority or "medium"
        
        control = InternalControl(
            tenant_id=document.tenant_id,
            control_id=control_id,
            name=name,
            description=statement.statement_text,
            category=category,
            priority=priority,
            source_document_id=document_id,
            source_statement_id=statement.id,
            status="draft",
            created_by=current_user.id
        )
        db.add(control)
        created_controls.append(control)
    
    db.commit()
    
    for control in created_controls:
        db.refresh(control)
    
    return created_controls
