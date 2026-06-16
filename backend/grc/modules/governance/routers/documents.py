from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime, timedelta
import logging
import os
import uuid
import json
import html

logger = logging.getLogger(__name__)
from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_, and_
from pydantic import BaseModel, Field
from openai import OpenAI

from ....models import (
    GovernanceDocument, GovernanceDocumentVersion, DocumentReviewer,
    DocumentApprovalStep, DocumentAuditLog, DocumentAnnotation, GRCUser,
    Tenant, PolicyStatement, InternalControl, ParsedFrameworkControl,
    UploadedFramework, CertificationJourney, PolicyReviewHistory, get_db
)
from ....routers.auth_router import require_auth, get_user_tenants, get_user_primary_tenant
from ..action_logger import log_governance_action
from ..ai_drafting import (
    build_tenant_context,
    build_framework_index,
    run_drafting_pipeline,
)

AI_INTEGRATIONS_OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
AI_INTEGRATIONS_OPENAI_BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")

GOVERNANCE_UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "uploads", "governance")
os.makedirs(GOVERNANCE_UPLOAD_DIR, exist_ok=True)

ALLOWED_FILE_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'doc': 'application/msword',
    'xls': 'application/vnd.ms-excel'
}

router = APIRouter(prefix="/documents", tags=["Governance - Documents"])


GOVERNANCE_DOMAIN_KEYWORDS: Dict[str, List[str]] = {
    "Governance and Oversight": [
        "governance", "oversight", "board", "committee", "strategy", "leadership", "management review",
        "policy management", "authority", "accountability"
    ],
    "Risk Management": [
        "risk", "assessment", "treatment", "mitigation", "register", "appetite", "tolerance", "residual"
    ],
    "Asset Management": [
        "asset", "inventory", "classification", "ownership", "lifecycle", "media", "disposal"
    ],
    "Identity and Access Management": [
        "access", "identity", "authentication", "authorization", "privilege", "user account", "segregation of duties",
        "joiner", "mover", "leaver", "password", "mfa", "remote access"
    ],
    "Security Operations and Monitoring": [
        "monitoring", "logging", "alert", "soc", "security event", "anomaly", "detection", "use case", "telemetry"
    ],
    "Incident Response and Resilience": [
        "incident", "response", "crisis", "breach", "resilience", "recovery", "continuity", "disaster", "forensic"
    ],
    "Vulnerability and Patch Management": [
        "vulnerability", "patch", "remediation", "hardening", "baseline", "configuration", "exposure", "scan"
    ],
    "Change and Release Management": [
        "change", "release", "deployment", "version", "promotion", "rollback", "cab", "testing"
    ],
    "Third-Party and Supplier Management": [
        "third party", "vendor", "supplier", "outsourcing", "service provider", "due diligence", "contract"
    ],
    "Data Protection and Privacy": [
        "data protection", "privacy", "personal data", "retention", "disposal", "encryption", "backup", "confidentiality"
    ],
    "Secure Development": [
        "development", "sdlc", "secure coding", "code review", "testing", "application security", "devsecops"
    ],
    "Physical and Environmental Security": [
        "physical", "environmental", "facility", "visitor", "perimeter", "surveillance"
    ],
    "Awareness and Competence": [
        "awareness", "training", "competence", "education", "disciplinary", "culture"
    ],
    "Compliance and Assurance": [
        "compliance", "audit", "assurance", "regulatory", "attestation", "evidence", "inspection"
    ],
}


class GovernanceDocumentCreate(BaseModel):
    document_code: Optional[str] = None
    title: str
    description: Optional[str] = None
    content: Optional[str] = None
    doc_type: str
    doc_sub_type: Optional[str] = None
    classification: str = "internal"
    parent_document_id: Optional[int] = None
    owner_id: Optional[int] = None
    author_id: Optional[int] = None
    department_id: Optional[int] = None
    effective_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None
    review_cycle_months: int = 12
    next_review_date: Optional[datetime] = None
    regulatory_scope: Optional[List[str]] = []
    framework_ids: Optional[List[int]] = []
    tags: Optional[List[str]] = []


class GovernanceDocumentUpdate(BaseModel):
    document_code: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    content: Optional[str] = None
    doc_type: Optional[str] = None
    doc_sub_type: Optional[str] = None
    classification: Optional[str] = None
    parent_document_id: Optional[int] = None
    status: Optional[str] = None
    owner_id: Optional[int] = None
    author_id: Optional[int] = None
    department_id: Optional[int] = None
    effective_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None
    review_cycle_months: Optional[int] = None
    next_review_date: Optional[datetime] = None
    regulatory_scope: Optional[List[str]] = None
    framework_ids: Optional[List[int]] = None
    tags: Optional[List[str]] = None


class BulkStatusUpdate(BaseModel):
    document_ids: List[int]
    status: str


class BulkArchive(BaseModel):
    document_ids: List[int]


def validate_tenant_access(user: GRCUser, tenant_id: int, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this tenant's data"
        )


def create_audit_log(
    db: Session,
    document_id: int,
    tenant_id: int,
    user_id: int,
    action: str,
    action_details: Optional[str] = None,
    field_changed: Optional[str] = None,
    old_value: Optional[str] = None,
    new_value: Optional[str] = None
) -> DocumentAuditLog:
    audit_log = DocumentAuditLog(
        document_id=document_id,
        tenant_id=tenant_id,
        action=action,
        action_details=action_details,
        field_changed=field_changed,
        old_value=old_value,
        new_value=new_value,
        performed_by=user_id,
        performed_at=datetime.utcnow()
    )
    db.add(audit_log)
    return audit_log


def serialize_document(doc: GovernanceDocument, db: Session = None) -> dict:
    result = {
        "id": doc.id,
        "tenant_id": doc.tenant_id,
        "document_code": doc.document_code,
        "title": doc.title,
        "description": doc.description,
        "content": doc.content,
        "file_name": doc.file_name,
        "file_size": doc.file_size,
        "file_type": doc.file_type,
        "has_file": doc.file_path is not None,
        "doc_type": doc.doc_type,
        "doc_sub_type": doc.doc_sub_type,
        "classification": doc.classification,
        "parent_document_id": doc.parent_document_id,
        "current_version": doc.current_version,
        "status": doc.status,
        "owner_id": doc.owner_id,
        "owner_name": doc.owner.display_name if doc.owner else None,
        "author_id": doc.author_id,
        "author_name": doc.author.display_name if doc.author else None,
        "department_id": doc.department_id,
        "effective_date": doc.effective_date.isoformat() if doc.effective_date else None,
        "expiry_date": doc.expiry_date.isoformat() if doc.expiry_date else None,
        "review_cycle_months": doc.review_cycle_months,
        "next_review_date": doc.next_review_date.isoformat() if doc.next_review_date else None,
        "last_reviewed_at": doc.last_reviewed_at.isoformat() if doc.last_reviewed_at else None,
        "last_reviewed_by": doc.last_reviewed_by,
        "regulatory_scope": doc.regulatory_scope or [],
        "framework_ids": doc.framework_ids or [],
        "tags": doc.tags or [],
        "approved_by": doc.approved_by,
        "approved_at": doc.approved_at.isoformat() if doc.approved_at else None,
        "published_by": doc.published_by,
        "published_at": doc.published_at.isoformat() if doc.published_at else None,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
    }
    
    if db:
        try:
            from ...compliance.schema_migrations import ensure_assigned_column
            ensure_assigned_column(db)
        except Exception:
            pass
        policy_count = db.query(PolicyStatement).filter(
            PolicyStatement.document_id == doc.id
        ).count()
        result["policy_statement_count"] = policy_count
    else:
        result["policy_statement_count"] = len(doc.policy_statements) if hasattr(doc, 'policy_statements') and doc.policy_statements else 0
    
    if hasattr(doc, 'workflow_instance') and doc.workflow_instance:
        wf = doc.workflow_instance
        result["workflow_instance"] = {
            "id": wf.id,
            "template_id": wf.template_id,
            "template_name": wf.template.name if wf.template else None,
            "current_step_id": wf.current_step_id,
            "current_step_sequence": wf.current_step_sequence,
            "current_step_name": wf.current_step.name if wf.current_step else None,
            "status": wf.status,
            "started_at": wf.started_at.isoformat() if wf.started_at else None,
            "completed_at": wf.completed_at.isoformat() if wf.completed_at else None,
        }
    else:
        result["workflow_instance"] = None
    
    return result


@router.get("")
def list_documents(
    tenant_id: Optional[int] = None,
    doc_type: Optional[str] = None,
    status_filter: Optional[str] = Query(None, alias="status"),
    owner_id: Optional[int] = None,
    classification: Optional[str] = None,
    framework_id: Optional[int] = None,
    search: Optional[str] = None,
    date_from: Optional[datetime] = None,
    date_to: Optional[datetime] = None,
    sort_by: str = "created_at",
    sort_order: str = "desc",
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"items": [], "total": 0, "skip": skip, "limit": limit}
    
    query = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner)
    ).filter(GovernanceDocument.tenant_id.in_(user_tenants))
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
        query = query.filter(GovernanceDocument.tenant_id == tenant_id)
    if doc_type:
        query = query.filter(GovernanceDocument.doc_type == doc_type)
    if status_filter:
        query = query.filter(GovernanceDocument.status == status_filter)
    if owner_id:
        query = query.filter(GovernanceDocument.owner_id == owner_id)
    if classification:
        query = query.filter(GovernanceDocument.classification == classification)
    if framework_id:
        query = query.filter(GovernanceDocument.framework_ids.contains([framework_id]))
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                GovernanceDocument.title.ilike(search_term),
                GovernanceDocument.description.ilike(search_term),
                GovernanceDocument.document_code.ilike(search_term)
            )
        )
    if date_from:
        query = query.filter(GovernanceDocument.created_at >= date_from)
    if date_to:
        query = query.filter(GovernanceDocument.created_at <= date_to)
    
    total = query.count()
    
    sort_column = getattr(GovernanceDocument, sort_by, GovernanceDocument.created_at)
    if sort_order == "asc":
        query = query.order_by(sort_column.asc())
    else:
        query = query.order_by(sort_column.desc())
    
    documents = query.offset(skip).limit(limit).all()
    
    return {
        "items": [serialize_document(doc, db) for doc in documents],
        "total": total,
        "skip": skip,
        "limit": limit
    }


@router.post("", status_code=status.HTTP_201_CREATED)
def create_document(
    document: GovernanceDocumentCreate,
    tenant_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
    else:
        tenant_id = get_user_primary_tenant(current_user, db)
        if not tenant_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User is not assigned to any tenant"
            )
    
    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()
    if not tenant:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Tenant not found"
        )
    
    if document.parent_document_id:
        parent = db.query(GovernanceDocument).filter(
            GovernanceDocument.id == document.parent_document_id,
            GovernanceDocument.tenant_id == tenant_id
        ).first()
        if not parent:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Parent document not found"
            )
    
    db_document = GovernanceDocument(
        tenant_id=tenant_id,
        document_code=document.document_code,
        title=document.title,
        description=document.description,
        content=document.content,
        doc_type=document.doc_type,
        doc_sub_type=document.doc_sub_type,
        classification=document.classification,
        parent_document_id=document.parent_document_id,
        owner_id=document.owner_id or current_user.id,
        author_id=document.author_id or current_user.id,
        department_id=document.department_id,
        effective_date=document.effective_date,
        expiry_date=document.expiry_date,
        review_cycle_months=document.review_cycle_months,
        next_review_date=document.next_review_date,
        regulatory_scope=document.regulatory_scope or [],
        framework_ids=document.framework_ids or [],
        tags=document.tags or [],
        status="draft",
        current_version="1.0"
    )
    db.add(db_document)
    db.flush()
    
    create_audit_log(
        db=db,
        document_id=db_document.id,
        tenant_id=tenant_id,
        user_id=current_user.id,
        action="created",
        action_details=f"Document '{document.title}' created"
    )
    
    db.commit()
    
    # Log the action for review
    log_governance_action(
        db=db,
        tenant_id=tenant_id,
        action_type="document_draft_created",
        action_description=f"Governance document draft created: '{document.title}' ({document.doc_type})",
        entity_type="governance_document",
        action_user_id=current_user.id,
        entity_id=db_document.id,
        action_metadata={
            "document_code": document.document_code,
            "doc_type": document.doc_type,
            "doc_sub_type": document.doc_sub_type,
            "classification": document.classification
        }
    )
    db.commit()
    
    db.refresh(db_document)
    
    return serialize_document(db_document, db)


@router.get("/policies")
def get_policies(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="policy", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/standards")
def get_standards(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="standard", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/procedures")
def get_procedures(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="procedure", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/guidelines")
def get_guidelines(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="guideline", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/charters")
def get_charters(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="charter", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/frameworks")
def get_framework_documents(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="framework", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/hierarchy")
def get_document_hierarchy(
    tenant_id: Optional[int] = None,
    parent_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return []
    
    query = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner)
    ).filter(GovernanceDocument.tenant_id.in_(user_tenants))
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
        query = query.filter(GovernanceDocument.tenant_id == tenant_id)
    
    if parent_id is None:
        query = query.filter(GovernanceDocument.parent_document_id.is_(None))
    else:
        query = query.filter(GovernanceDocument.parent_document_id == parent_id)
    
    documents = query.order_by(GovernanceDocument.doc_type, GovernanceDocument.title).all()
    
    def build_hierarchy(doc):
        result = serialize_document(doc, db)
        children = db.query(GovernanceDocument).filter(
            GovernanceDocument.parent_document_id == doc.id
        ).all()
        result["children"] = [build_hierarchy(child) for child in children]
        return result
    
    return [build_hierarchy(doc) for doc in documents]


@router.get("/reviews/upcoming")
def get_upcoming_reviews(
    days: int = 90,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    now = datetime.utcnow()
    cutoff = now + timedelta(days=days)
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.tenant_id.in_(user_tenants),
        GovernanceDocument.next_review_date != None,
        GovernanceDocument.next_review_date <= cutoff,
        GovernanceDocument.status != "retired"
    ).order_by(GovernanceDocument.next_review_date.asc()).all()
    
    results = []
    for doc in documents:
        is_overdue = doc.next_review_date and doc.next_review_date < now
        days_until = (doc.next_review_date - now).days if doc.next_review_date else None
        
        owner_name = None
        if doc.owner_id:
            owner = db.query(GRCUser).filter(GRCUser.id == doc.owner_id).first()
            if owner:
                owner_name = owner.display_name
        
        results.append({
            "document_id": doc.id,
            "title": doc.title,
            "doc_type": doc.doc_type,
            "status": doc.status,
            "review_cycle_months": doc.review_cycle_months,
            "next_review_date": doc.next_review_date.isoformat() if doc.next_review_date else None,
            "is_overdue": is_overdue,
            "days_until_review": days_until,
            "owner_name": owner_name
        })
    
    overdue_count = sum(1 for r in results if r["is_overdue"])
    upcoming_count = len(results) - overdue_count
    
    return {
        "total": len(results),
        "overdue_count": overdue_count,
        "upcoming_count": upcoming_count,
        "reviews": results
    }


@router.get("/{document_id}")
def get_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner),
        joinedload(GovernanceDocument.author),
        joinedload(GovernanceDocument.approver),
        joinedload(GovernanceDocument.versions),
        joinedload(GovernanceDocument.reviewers).joinedload(DocumentReviewer.user),
        joinedload(GovernanceDocument.approval_steps).joinedload(DocumentApprovalStep.approver),
        joinedload(GovernanceDocument.parent_document)
    ).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    result = serialize_document(document, db)
    
    result["versions"] = [
        {
            "id": v.id,
            "version_number": v.version_number,
            "change_type": v.change_type,
            "title": v.title,
            "change_summary": v.change_summary,
            "change_reason": v.change_reason,
            "status": v.status,
            "created_at": v.created_at.isoformat() if v.created_at else None,
            "created_by": v.created_by,
            "creator_name": v.creator.display_name if v.creator else None,
            "approved_by": v.approved_by,
            "approved_at": v.approved_at.isoformat() if v.approved_at else None,
        }
        for v in document.versions
    ]
    
    result["reviewers"] = [
        {
            "id": r.id,
            "user_id": r.user_id,
            "user_name": r.user.display_name if r.user else None,
            "role_type": r.role_type,
            "sequence": r.sequence,
            "is_required": r.is_required,
            "notify_on_update": r.notify_on_update,
            "notify_on_expiry": r.notify_on_expiry,
            "assigned_at": r.assigned_at.isoformat() if r.assigned_at else None,
        }
        for r in document.reviewers
    ]
    
    result["approval_steps"] = [
        {
            "id": s.id,
            "step_sequence": s.step_sequence,
            "step_name": s.step_name,
            "approval_type": s.approval_type,
            "approver_id": s.approver_id,
            "approver_name": s.approver.display_name if s.approver else None,
            "approver_role": s.approver_role,
            "status": s.status,
            "requested_at": s.requested_at.isoformat() if s.requested_at else None,
            "due_date": s.due_date.isoformat() if s.due_date else None,
            "completed_at": s.completed_at.isoformat() if s.completed_at else None,
            "comments": s.comments,
        }
        for s in document.approval_steps
    ]
    
    result["parent_document"] = None
    if document.parent_document:
        result["parent_document"] = {
            "id": document.parent_document.id,
            "title": document.parent_document.title,
            "doc_type": document.parent_document.doc_type,
        }
    
    child_documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.parent_document_id == document_id
    ).all()
    result["child_documents"] = [
        {
            "id": c.id,
            "title": c.title,
            "doc_type": c.doc_type,
            "status": c.status,
        }
        for c in child_documents
    ]
    
    return result


@router.put("/{document_id}")
def update_document(
    document_id: int,
    document_update: GovernanceDocumentUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    update_data = document_update.model_dump(exclude_unset=True)
    changed_fields = []
    
    for field, new_value in update_data.items():
        old_value = getattr(document, field)
        if old_value != new_value:
            changed_fields.append({
                "field": field,
                "old_value": str(old_value) if old_value else None,
                "new_value": str(new_value) if new_value else None
            })
            setattr(document, field, new_value)
    
    document.updated_at = datetime.utcnow()
    
    if changed_fields:
        create_audit_log(
            db=db,
            document_id=document_id,
            tenant_id=document.tenant_id,
            user_id=current_user.id,
            action="updated",
            action_details=f"Updated fields: {', '.join([c['field'] for c in changed_fields])}"
        )
    
    db.commit()
    db.refresh(document)
    
    return serialize_document(document, db)


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    child_count = db.query(GovernanceDocument).filter(
        GovernanceDocument.parent_document_id == document_id
    ).count()
    
    if child_count > 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot delete document with {child_count} child documents. Delete or reassign child documents first."
        )
    
    db.query(InternalControl).filter(InternalControl.source_document_id == document_id).delete()
    
    db.delete(document)
    db.commit()
    
    return None


@router.get("/{document_id}/review-history")
def get_review_history(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    reviews = db.query(PolicyReviewHistory).filter(
        PolicyReviewHistory.document_id == document_id,
        PolicyReviewHistory.tenant_id.in_(user_tenants)
    ).order_by(PolicyReviewHistory.created_at.desc()).all()
    
    return [{
        "id": r.id,
        "review_type": r.review_type,
        "review_status": r.review_status,
        "scheduled_date": r.scheduled_date.isoformat() if r.scheduled_date else None,
        "started_at": r.started_at.isoformat() if r.started_at else None,
        "completed_at": r.completed_at.isoformat() if r.completed_at else None,
        "reviewer_id": r.reviewer_id,
        "review_notes": r.review_notes,
        "changes_made": r.changes_made,
        "outcome": r.outcome,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    } for r in reviews]


class StartReviewRequest(BaseModel):
    review_type: str = "periodic"


class CompleteReviewRequest(BaseModel):
    review_notes: str = ""
    changes_made: str = ""
    outcome: str = "no_changes"


@router.post("/{document_id}/start-review")
def start_review(
    document_id: int,
    request_body: Optional[StartReviewRequest] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    review_type = request_body.review_type if request_body else "periodic"
    
    review = PolicyReviewHistory(
        tenant_id=document.tenant_id,
        document_id=document_id,
        review_type=review_type,
        review_status="in_progress",
        scheduled_date=document.next_review_date,
        started_at=datetime.utcnow(),
        reviewer_id=current_user.id,
    )
    db.add(review)
    db.commit()
    db.refresh(review)
    
    return {
        "id": review.id,
        "review_type": review.review_type,
        "review_status": review.review_status,
        "scheduled_date": review.scheduled_date.isoformat() if review.scheduled_date else None,
        "started_at": review.started_at.isoformat() if review.started_at else None,
        "reviewer_id": review.reviewer_id,
        "message": "Review started successfully"
    }


@router.post("/{document_id}/complete-review")
def complete_review(
    document_id: int,
    request_body: CompleteReviewRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    review = db.query(PolicyReviewHistory).filter(
        PolicyReviewHistory.document_id == document_id,
        PolicyReviewHistory.review_status == "in_progress",
        PolicyReviewHistory.tenant_id.in_(user_tenants)
    ).order_by(PolicyReviewHistory.created_at.desc()).first()
    
    if not review:
        review = PolicyReviewHistory(
            tenant_id=document.tenant_id,
            document_id=document_id,
            review_type="periodic",
            review_status="in_progress",
            started_at=datetime.utcnow(),
            reviewer_id=current_user.id,
        )
        db.add(review)
        db.flush()
    
    now = datetime.utcnow()
    review.review_status = "completed"
    review.completed_at = now
    review.review_notes = request_body.review_notes
    review.changes_made = request_body.changes_made
    review.outcome = request_body.outcome
    review.reviewer_id = current_user.id
    
    document.last_reviewed_at = now
    document.last_reviewed_by = current_user.id
    if document.review_cycle_months:
        from dateutil.relativedelta import relativedelta
        document.next_review_date = now + relativedelta(months=document.review_cycle_months)
    
    document.updated_at = now
    
    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="review_completed",
        action_details=f"Review completed with outcome: {request_body.outcome}"
    )
    
    db.commit()
    db.refresh(review)
    db.refresh(document)
    
    return {
        "id": review.id,
        "review_status": review.review_status,
        "completed_at": review.completed_at.isoformat() if review.completed_at else None,
        "outcome": review.outcome,
        "next_review_date": document.next_review_date.isoformat() if document.next_review_date else None,
        "message": "Review completed successfully"
    }


@router.post("/bulk-update-status")
def bulk_update_status(
    bulk_update: BulkStatusUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    valid_statuses = ["draft", "pending_review", "pending_approval", "approved", "published", "expired", "archived", "exception_applied"]
    if bulk_update.status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}"
        )
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.id.in_(bulk_update.document_ids),
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found"
        )
    
    updated_count = 0
    for doc in documents:
        old_status = doc.status
        doc.status = bulk_update.status
        doc.updated_at = datetime.utcnow()
        
        create_audit_log(
            db=db,
            document_id=doc.id,
            tenant_id=doc.tenant_id,
            user_id=current_user.id,
            action="status_changed",
            action_details=f"Status changed from '{old_status}' to '{bulk_update.status}'",
            field_changed="status",
            old_value=old_status,
            new_value=bulk_update.status
        )
        updated_count += 1
    
    db.commit()
    
    return {
        "message": f"Successfully updated {updated_count} documents",
        "updated_count": updated_count
    }


@router.post("/bulk-archive")
def bulk_archive(
    bulk_archive_request: BulkArchive,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.id.in_(bulk_archive_request.document_ids),
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found"
        )
    
    archived_count = 0
    for doc in documents:
        old_status = doc.status
        doc.status = "archived"
        doc.updated_at = datetime.utcnow()
        
        create_audit_log(
            db=db,
            document_id=doc.id,
            tenant_id=doc.tenant_id,
            user_id=current_user.id,
            action="archived",
            action_details=f"Document archived (previous status: '{old_status}')",
            field_changed="status",
            old_value=old_status,
            new_value="archived"
        )
        archived_count += 1
    
    db.commit()
    
    return {
        "message": f"Successfully archived {archived_count} documents",
        "archived_count": archived_count
    }


@router.put("/{document_id}/status")
def update_document_status(
    document_id: int,
    body: dict,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    new_status = body.get("status")
    valid_statuses = ["draft", "pending_review", "pending_approval", "approved", "published", "expired", "archived"]
    if new_status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}"
        )

    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    old_status = document.status
    document.status = new_status
    document.updated_at = datetime.utcnow()

    if new_status == "published":
        document.published_by = current_user.id
        document.published_at = datetime.utcnow()

    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="status_changed",
        action_details=f"Status changed from '{old_status}' to '{new_status}' by {current_user.display_name}",
        field_changed="status",
        old_value=old_status,
        new_value=new_status
    )

    db.commit()
    db.refresh(document)

    return serialize_document(document, db)


@router.post("/{document_id}/publish")
def publish_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Publish an approved document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if document.status != "approved":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only approved documents can be published. Current status: {document.status}"
        )
    
    old_status = document.status
    document.status = "published"
    document.published_by = current_user.id
    document.published_at = datetime.utcnow()
    document.updated_at = datetime.utcnow()
    
    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="published",
        action_details=f"Document published by {current_user.display_name}",
        field_changed="status",
        old_value=old_status,
        new_value="published"
    )
    
    db.commit()
    db.refresh(document)
    
    return serialize_document(document, db)


@router.get("/{document_id}/audit-logs")
def get_document_audit_logs(
    document_id: int,
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    logs = db.query(DocumentAuditLog).options(
        joinedload(DocumentAuditLog.user)
    ).filter(
        DocumentAuditLog.document_id == document_id
    ).order_by(DocumentAuditLog.performed_at.desc()).offset(skip).limit(limit).all()
    
    return [
        {
            "id": log.id,
            "action": log.action,
            "action_details": log.action_details,
            "field_changed": log.field_changed,
            "old_value": log.old_value,
            "new_value": log.new_value,
            "performed_by": log.performed_by,
            "performer_name": log.user.display_name if log.user else None,
            "performed_at": log.performed_at.isoformat() if log.performed_at else None,
            "ip_address": log.ip_address,
        }
        for log in logs
    ]


@router.post("/{document_id}/upload-file")
async def upload_document_file(
    document_id: int,
    file: UploadFile = File(...),
    change_summary: Optional[str] = Form(None),
    change_reason: Optional[str] = Form(None),
    create_new_version: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Upload a file (PDF, Word, Excel) for an existing governance document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    file_ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only PDF, Word (.doc, .docx), and Excel (.xls, .xlsx) files are supported"
        )
    
    unique_id = str(uuid.uuid4())
    safe_filename = f"{unique_id}_{file.filename}"
    file_path = os.path.join(GOVERNANCE_UPLOAD_DIR, safe_filename)
    
    content = await file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    old_file_name = document.file_name
    
    current_ver = document.current_version or "1.0"
    if create_new_version and document.file_name:
        try:
            version_parts = current_ver.split('.')
            major = int(version_parts[0]) if version_parts else 1
            minor = int(version_parts[1]) if len(version_parts) > 1 else 0
            new_version = f"{major}.{minor + 1}"
        except (ValueError, IndexError):
            new_version = "1.1"
        
        existing_versions = db.query(GovernanceDocumentVersion).filter(
            GovernanceDocumentVersion.document_id == document_id,
            GovernanceDocumentVersion.status == "current"
        ).all()
        for v in existing_versions:
            v.status = "superseded"
        
        version = GovernanceDocumentVersion(
            document_id=document_id,
            version_number=new_version,
            change_type="minor",
            title=document.title,
            content=document.content,
            file_name=file.filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_ext,
            change_summary=change_summary or f"File updated: {file.filename}",
            change_reason=change_reason,
            status="current",
            created_by=current_user.id
        )
        db.add(version)
        document.current_version = new_version
    
    document.file_name = file.filename
    document.file_path = file_path
    document.file_size = file_size
    document.file_type = file_ext
    document.updated_at = datetime.utcnow()
    
    create_audit_log(
        db=db,
        document_id=document.id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="file_uploaded",
        action_details=f"File uploaded: {file.filename} ({file_size} bytes)",
        field_changed="file_name",
        old_value=old_file_name,
        new_value=file.filename
    )
    
    # Log the action for review
    log_governance_action(
        db=db,
        tenant_id=document.tenant_id,
        action_type="document_file_uploaded",
        action_description=f"File uploaded to document: '{document.title}' - {file.filename}",
        entity_type="governance_document",
        action_user_id=current_user.id,
        entity_id=document.id,
        action_metadata={
            "document_code": document.document_code,
            "doc_type": document.doc_type,
            "file_name": file.filename,
            "file_size": file_size,
            "version": document.current_version
        }
    )
    
    db.commit()
    db.refresh(document)
    
    return {
        "message": "File uploaded successfully",
        "document": serialize_document(document, db)
    }


@router.get("/{document_id}/download-file")
def download_document_file(
    document_id: int,
    version_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Download the attached file for a governance document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if version_id:
        version = db.query(GovernanceDocumentVersion).filter(
            GovernanceDocumentVersion.id == version_id,
            GovernanceDocumentVersion.document_id == document_id
        ).first()
        
        if not version:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Version not found"
            )
        
        if not version.file_path or not os.path.exists(version.file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="File not found for this version"
            )
        
        return FileResponse(
            path=version.file_path,
            filename=version.file_name or f"document_{document_id}_v{version.version_number}",
            media_type=ALLOWED_FILE_TYPES.get(version.file_type, "application/octet-stream")
        )
    
    if not document.file_path or not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No file attached to this document"
        )
    
    return FileResponse(
        path=document.file_path,
        filename=document.file_name or f"document_{document_id}",
        media_type=ALLOWED_FILE_TYPES.get(document.file_type, "application/octet-stream")
    )


@router.post("/upload-with-file", status_code=status.HTTP_201_CREATED)
async def create_document_with_file(
    file: UploadFile = File(...),
    title: str = Form(...),
    doc_type: str = Form(...),
    description: Optional[str] = Form(None),
    document_code: Optional[str] = Form(None),
    classification: str = Form("internal"),
    owner_id: Optional[int] = Form(None),
    tenant_id: Optional[int] = Form(None),
    # JSON-encoded list of UploadedFramework ids the doc should link to.
    # Sent as a JSON string because List[int] in multipart Form isn't
    # consistently supported across FastAPI versions and front-end
    # FormData encoders. Parsed below; safe if absent or malformed.
    framework_ids: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Create a new governance document with an attached file"""
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
    else:
        tenant_id = get_user_primary_tenant(current_user, db)
        if not tenant_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User is not assigned to any tenant"
            )
    
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    file_ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only PDF, Word (.doc, .docx), and Excel (.xls, .xlsx) files are supported"
        )
    
    unique_id = str(uuid.uuid4())
    safe_filename = f"{unique_id}_{file.filename}"
    file_path = os.path.join(GOVERNANCE_UPLOAD_DIR, safe_filename)
    
    content = await file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Decode framework_ids JSON if the client supplied any. Tolerant of
    # missing / malformed input — the upload still succeeds, the doc
    # just isn't framework-tagged in that case.
    parsed_framework_ids: List[int] = []
    if framework_ids:
        try:
            raw = json.loads(framework_ids)
            if isinstance(raw, list):
                parsed_framework_ids = [int(v) for v in raw if v is not None]
        except (json.JSONDecodeError, TypeError, ValueError):
            parsed_framework_ids = []

    document = GovernanceDocument(
        tenant_id=tenant_id,
        title=title,
        description=description,
        document_code=document_code,
        doc_type=doc_type,
        classification=classification,
        owner_id=owner_id or current_user.id,
        author_id=current_user.id,
        file_name=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_ext,
        status="draft",
        current_version="1.0",
        framework_ids=parsed_framework_ids,
    )
    db.add(document)
    db.flush()
    
    version = GovernanceDocumentVersion(
        document_id=document.id,
        version_number="1.0",
        change_type="major",
        title=title,
        file_name=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_ext,
        change_summary="Initial version",
        status="current",
        created_by=current_user.id
    )
    db.add(version)
    
    create_audit_log(
        db=db,
        document_id=document.id,
        tenant_id=tenant_id,
        user_id=current_user.id,
        action="created",
        action_details=f"Document created with file: {file.filename}"
    )
    
    db.commit()
    
    # Log the action for review
    log_governance_action(
        db=db,
        tenant_id=tenant_id,
        action_type="document_uploaded",
        action_description=f"Governance document uploaded: '{document.title}' ({document.doc_type}) - File: {file.filename}",
        entity_type="governance_document",
        action_user_id=current_user.id,
        entity_id=document.id,
        action_metadata={
            "document_code": document.document_code,
            "doc_type": document.doc_type,
            "file_name": file.filename,
            "file_size": file_size,
            "classification": document.classification
        }
    )
    db.commit()
    
    db.refresh(document)
    
    return {
        "message": "Document created successfully",
        "document": serialize_document(document, db)
    }


class PolicyAIDraftRequest(BaseModel):
    doc_type: str
    title: str
    framework_ids: Optional[List[int]] = None
    regulatory_scope: Optional[List[str]] = None
    description: Optional[str] = None
    include_sections: Optional[List[str]] = None
    parent_document_id: Optional[int] = None


class PolicySuggestRequest(BaseModel):
    framework_ids: List[int]
    doc_type: Optional[str] = None


def infer_governance_domain(control: ParsedFrameworkControl) -> str:
    explicit_domain = (control.domain or control.category or "").strip()
    if explicit_domain:
        return explicit_domain

    haystack = " ".join(
        filter(
            None,
            [
                control.title,
                control.description,
                control.full_text,
                control.parent_section,
            ],
        )
    ).lower()

    best_domain = "General Governance"
    best_score = 0
    for domain, keywords in GOVERNANCE_DOMAIN_KEYWORDS.items():
        score = sum(1 for keyword in keywords if keyword in haystack)
        if score > best_score:
            best_domain = domain
            best_score = score

    return best_domain


def build_framework_control_context(
    frameworks: List[UploadedFramework],
    db: Session,
    per_framework_limit: int = 80,
    representative_controls_per_domain: int = 4,
) -> Tuple[str, List[dict], List[str]]:
    context_parts: List[str] = []
    framework_alignment: List[dict] = []
    all_domains: List[str] = []

    for framework in frameworks:
        controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework.id
        ).order_by(
            ParsedFrameworkControl.priority.desc(),
            ParsedFrameworkControl.original_reference.asc(),
            ParsedFrameworkControl.control_id.asc(),
        ).limit(per_framework_limit).all()

        if not controls:
            continue

        grouped_controls: Dict[str, dict] = {}
        framework_refs: List[str] = []

        for control in controls:
            reference = control.original_reference or control.control_id
            domain = infer_governance_domain(control)
            bucket = grouped_controls.setdefault(domain, {"count": 0, "samples": []})
            bucket["count"] += 1
            if len(bucket["samples"]) < representative_controls_per_domain:
                bucket["samples"].append({
                    "reference": reference,
                    "title": control.title,
                    "description": (control.description or control.full_text or "")[:240],
                })
            if len(framework_refs) < 15:
                framework_refs.append(reference)

        sorted_domains = sorted(
            grouped_controls.items(),
            key=lambda item: (-item[1]["count"], item[0]),
        )
        top_domains = [domain for domain, _ in sorted_domains[:8]]
        all_domains.extend(top_domains)

        context_parts.append(f"Framework: {framework.name}")
        context_parts.append(f"Total Controls Reviewed: {len(controls)}")
        context_parts.append("Primary Governance Domains and Representative Controls:")
        for domain, bucket in sorted_domains[:8]:
            context_parts.append(f"- {domain} ({bucket['count']} controls)")
            for sample in bucket["samples"]:
                sample_line = f"  - {sample['reference']}: {sample['title']}"
                if sample["description"]:
                    sample_line += f" | {sample['description']}"
                context_parts.append(sample_line)
        context_parts.append("")

        framework_alignment.append({
            "framework": framework.name,
            "controls": framework_refs,
            "domains": top_domains,
        })

    ordered_domains = list(dict.fromkeys(all_domains))
    return "\n".join(context_parts).strip(), framework_alignment, ordered_domains


def get_document_depth_instruction(doc_type: str) -> dict:
    if doc_type == "procedure":
        return {
            "minimum_words": 2200,
            "template_name": "ISO-aligned procedure",
            "section_outline": [
                "1. Document Control",
                "2. Purpose",
                "3. Scope and Applicability",
                "4. Normative References and Related Documents",
                "5. Terms and Definitions",
                "6. Roles and Responsibilities",
                "7. Preconditions and Trigger Events",
                "8. Procedure Steps",
                "9. Records and Evidence Retention",
                "10. Monitoring, Exceptions, and Escalation",
                "11. Review and Revision History"
            ],
            "statement_style": "Break activities into ordered steps, sub-steps, decision points, required records, and escalation rules."
        }

    if doc_type == "standard":
        return {
            "minimum_words": 1800,
            "template_name": "ISO-aligned standard",
            "section_outline": [
                "1. Document Control",
                "2. Purpose",
                "3. Scope and Applicability",
                "4. Normative References",
                "5. Control Objectives",
                "6. Mandatory Standard Requirements",
                "7. Roles and Responsibilities",
                "8. Measurement and Exceptions",
                "9. Enforcement",
                "10. Review and Revision History"
            ],
            "statement_style": "Write mandatory requirements as precise atomic statements with numbered sub-clauses."
        }

    if doc_type == "guideline":
        return {
            "minimum_words": 1600,
            "template_name": "ISO-aligned guideline",
            "section_outline": [
                "1. Document Control",
                "2. Purpose",
                "3. Scope and Intended Audience",
                "4. Guidance Principles",
                "5. Recommended Practices",
                "6. Roles and Responsibilities",
                "7. Examples and Implementation Notes",
                "8. Related Standards and Procedures",
                "9. Review and Revision History"
            ],
            "statement_style": "Provide practical guidance with clear recommendations, examples, and implementation notes."
        }

    return {
        "minimum_words": 1800,
        "template_name": "ISO-aligned policy",
        "section_outline": [
            "1. Document Control",
            "2. Purpose",
            "3. Scope and Applicability",
            "4. Normative References",
            "5. Terms and Definitions",
            "6. Policy Objectives",
            "7. Policy Statements",
            "8. Roles and Responsibilities",
            "9. Governance, Monitoring, and Exceptions",
            "10. Compliance and Enforcement",
            "11. Review and Revision History"
        ],
        "statement_style": "Write policy requirements as atomic statements, each addressing a single obligation, and decompose them into numbered clauses where needed."
    }


def generate_policy_with_openai(
    doc_type: str,
    title: str,
    controls_context: str,
    domain_context: List[str],
    regulatory_scope: List[str],
    description: Optional[str],
    include_sections: Optional[List[str]],
    parent_document_context: Optional[str],
) -> dict:
    if not AI_INTEGRATIONS_OPENAI_API_KEY or not AI_INTEGRATIONS_OPENAI_BASE_URL:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI integration not configured"
        )
    
    client = OpenAI(
        api_key=AI_INTEGRATIONS_OPENAI_API_KEY,
        base_url=AI_INTEGRATIONS_OPENAI_BASE_URL
    )
    
    doc_type_labels = {
        "policy": "Policy",
        "standard": "Standard", 
        "procedure": "Procedure",
        "guideline": "Guideline"
    }
    doc_label = doc_type_labels.get(doc_type, "Policy")
    depth_instruction = get_document_depth_instruction(doc_type)
    
    sections_instruction = ""
    if include_sections:
        sections_instruction = f"\n\nThe document MUST include these specific sections: {', '.join(include_sections)}"
    
    regulatory_context = ""
    if regulatory_scope:
        regulatory_context = f"\n\nThis document should align with the following regulatory frameworks: {', '.join(regulatory_scope)}"

    domain_summary_context = ""
    if domain_context:
        domain_summary_context = (
            "\n\nPrimary governance domains that must be covered in a structured way: "
            + ", ".join(domain_context)
        )
    
    description_context = ""
    if description:
        description_context = f"\n\nAdditional requirements: {description}"

    parent_context = ""
    if parent_document_context:
        parent_context = (
            "\n\nParent document context (mandatory reference baseline):\n"
            f"{parent_document_context}\n"
            "- Reuse the parent document's intent, terminology, and control objectives.\n"
            "- Keep this new document subordinate and implementation-focused relative to the parent.\n"
        )

    framework_option_context = ""
    if controls_context:
        framework_option_context = (
            "\n\nFramework-based drafting mode:\n"
            "- Maintain explicit traceability to cited control references where feasible.\n"
            "- Reflect domain obligations from framework context in structure and clause wording."
        )
    else:
        framework_option_context = (
            "\n\nBest-practice drafting mode (no framework references selected):\n"
            "- Produce a comprehensive document using banking/financial-institution security governance best practices.\n"
            "- Include practical and auditable controls suitable for regulated banks.\n"
            "- Use risk-based language, segregation of duties, approval governance, monitoring, evidence retention, and exception handling."
        )
    
    prompt = f"""You are an expert governance and compliance consultant for banking and financial institutions. Generate a professional, mature, enterprise-grade {doc_label} document titled "{title}".

{controls_context}
{parent_context}
{framework_option_context}
{regulatory_context}
{domain_summary_context}
{description_context}
{sections_instruction}

Document authoring requirements:
- Use an {depth_instruction['template_name']} structure.
- Minimum length: {depth_instruction['minimum_words']} words. Do not produce a short summary.
- The document must read like a real enterprise artifact, not generic guidance.
- The document must be general and reusable. Refer only to "the organization", "management", "personnel", "functions", or specific roles. Do not invent company names or environment-specific facts.
- Use numbered headings and numbered sub-clauses.
- {depth_instruction['statement_style']}
- Tie each major section back to the relevant framework obligations and control themes.
- Include implementation depth, governance detail, review expectations, evidence expectations, and exception handling.
- Avoid placeholders, vague statements, and one-line sections.
- Every substantive section must contain enough depth to be operationally useful. Do not produce shallow sections.
- Policy and standard sections should contain multiple atomic obligations, each expressed as one requirement per clause.
- Procedure sections must be written as detailed operational steps with role ownership, triggers, prerequisites, inputs, outputs, records, evidence, escalation points, and exception handling.
- Use clause hierarchy such as 7, 7.1, 7.1.1 when the topic requires decomposition.
- Include document control elements such as owner, approver, effective date, review frequency, and version control guidance in a general ISO-style format.
- Include monitoring, metrics, non-compliance handling, issue escalation, and periodic review expectations.
- If multiple governance domains are implicated, create domain-specific sub-sections instead of collapsing them into a generic section.
- Ensure the output is materially longer than two pages of text when rendered.

Minimum content expectations by document type:
- Policy: at least 12 atomic policy clauses across the relevant domains, with clear governance and enforcement language.
- Standard: at least 12 mandatory standard requirements with prescriptive wording and measurable expectations.
- Procedure: at least 10 ordered procedural steps plus supporting roles, records, approvals, and escalation subsections.
- Guideline: detailed practice guidance, examples, recommended methods, and implementation notes across the relevant domains.

Required section outline:
{chr(10).join(f'- {section}' for section in depth_instruction['section_outline'])}

Document quality requirements:
- Policy and standard clauses must be prescriptive and auditable.
- Procedure content must include roles, step sequence, inputs, outputs, records, and escalation points.
- If this is a Procedure and a parent Policy exists, map procedure steps clearly to policy intent and control obligations.
- If this is a Guideline, include practical implementation examples and explicit placeholders for screenshots where helpful, e.g., "[Screenshot: Admin console setting path]".
- When a topic requires hierarchy, break it into clauses, sub-clauses, and bullet lists.
- If frameworks imply multiple domains, reflect that in the structure rather than collapsing everything into generic text.
- Suggested section content must be substantial; avoid one-line summaries.
- The generated content and suggested sections must be consistent with each other.

Return a JSON object with:
{{
  "generated_content": "The full document content in markdown format",
  "suggested_title": "The recommended document title",
  "suggested_sections": [
    {{"heading": "1. Purpose", "content": "Section content..."}},
    {{"heading": "2. Scope", "content": "Section content..."}},
    {{"heading": "3. {doc_label} Statements", "content": "Section content..."}},
    {{"heading": "4. Roles and Responsibilities", "content": "Section content..."}},
    {{"heading": "5. Compliance", "content": "Section content..."}},
    {{"heading": "6. Review and Updates", "content": "Section content..."}}
  ],
  "framework_alignment": [
    {{"framework": "Framework Name", "controls": ["Control ID 1", "Control ID 2"]}}
  ]
}}"""

    try:
        # Use the same env-driven model as the async drafting pipeline so a
        # single setting (`OPENAI_DRAFT_MODEL` in .env) drives both paths.
        # Default to GPT-5; GPT-5 currently rejects custom temperature, so
        # we omit it on that family and rely on the default.
        _legacy_model = os.environ.get("OPENAI_DRAFT_MODEL", "gpt-5")
        # Layer enterprise SME craft onto the prompt and system message — the
        # legacy generator was previously a single-shot call with a vague
        # "expert consultant" system message that produced generic output.
        # We inject the doc-type-specific drafting craft block + banking
        # reality block so the legacy path produces work matching the
        # multi-stage pipeline's quality bar.
        from ..ai_drafting.enterprise_craft import (
            SME_SYSTEM_ADDENDUM,
            enterprise_drafting_block,
        )
        _craft_block = enterprise_drafting_block(doc_type)
        _system_msg = (
            "You are an expert governance and compliance consultant that "
            "creates mature ISO-aligned governance documents for enterprise "
            "organizations. Always respond with valid JSON. Never return "
            "brief or generic output."
            + SME_SYSTEM_ADDENDUM
        )
        _user_msg = prompt + "\n\n" + _craft_block + (
            "\n\nFinal reminder: replace every generic governance verb "
            "(`ensure`, `make sure`, `consider`, `where possible`, `as "
            "appropriate`) with a prescriptive verb and a named owner. "
            "Output must read like a real bank's policy library — not a "
            "SaaS startup blog."
        )
        # GPT-5 / o-series are reasoning models — `max_completion_tokens`
        # has to cover BOTH hidden reasoning AND visible output. 12k was OK
        # on gpt-4o (no reasoning trace); for GPT-5 we double it so a long
        # policy body doesn't get cut to empty after the model finishes
        # thinking. Reasoning models also reject custom `temperature`.
        _is_reasoning = _legacy_model.lower().startswith(("gpt-5", "o1", "o3", "o4"))
        _legacy_kwargs: Dict[str, Any] = {
            "model": _legacy_model,
            "messages": [
                {"role": "system", "content": _system_msg},
                {"role": "user", "content": _user_msg},
            ],
            "response_format": {"type": "json_object"},
            "max_completion_tokens": 24000 if _is_reasoning else 12000,
        }

        def _call(kwargs: Dict[str, Any]):
            resp = client.chat.completions.create(**kwargs)
            txt = resp.choices[0].message.content
            return resp, txt

        try:
            response, result_text = _call(_legacy_kwargs)
        except Exception as primary_exc:  # noqa: BLE001
            logger.warning("Legacy generator primary model %r failed: %s", _legacy_model, primary_exc)
            response, result_text = None, None

        # Auto-fallback when the primary returns empty content (GPT-5 ate the
        # whole budget on reasoning) or raises. Falls back to gpt-4o so the
        # user still gets a real document instead of an empty-body stub.
        if not result_text:
            _fallback_model = os.environ.get("OPENAI_DRAFT_FALLBACK_MODEL", "gpt-4o")
            if _fallback_model and _fallback_model != _legacy_model:
                logger.info(
                    "Legacy generator falling back from %r to %r after empty primary output",
                    _legacy_model, _fallback_model,
                )
                _fallback_kwargs = dict(_legacy_kwargs)
                _fallback_kwargs["model"] = _fallback_model
                _fallback_kwargs["temperature"] = 0.4  # gpt-4o accepts this
                _fallback_kwargs["max_completion_tokens"] = 12000
                try:
                    response, result_text = _call(_fallback_kwargs)
                except Exception as fb_exc:  # noqa: BLE001
                    logger.exception("Legacy generator fallback %r also failed: %s", _fallback_model, fb_exc)

        if not result_text:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=(
                    "OpenAI returned empty content from both the primary model "
                    f"({_legacy_model}) and the fallback. Check the API key "
                    "billing/quota, OPENAI_DRAFT_MODEL value, and server logs."
                ),
            )
        result = json.loads(result_text)
        return result
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse OpenAI response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"OpenAI API error: {error_msg}"
        )


def _normalize_doc_title(s: str) -> str:
    """Normalize a document title for fuzzy duplicate detection.

    Lowercase, drop everything that isn't a letter or digit, collapse runs.
    "ISO 27001 Information Security Policy v2.0" → "iso27001informationsecuritypolicy".
    Used as the first-pass exact-match guard before the semantic dedup AI call.
    """
    return "".join(ch for ch in (s or "").lower() if ch.isalnum())


# Words to strip when computing a token-based similarity score. They appear
# in nearly every governance-document title and would inflate match scores
# for unrelated titles otherwise.
_TITLE_STOPWORDS = {
    "the", "and", "of", "for", "on", "in", "to", "a", "an", "is", "or", "by",
    "with", "from", "as", "at", "policy", "policies", "procedure", "procedures",
    "standard", "standards", "guideline", "guidelines", "charter", "framework",
    "document", "documents", "v", "version", "iso", "nist", "soc", "pci", "dss",
    "hipaa", "gdpr", "sama", "nca", "sbp",
}

# Common security/governance abbreviations expanded so a title using the
# acronym is still tokenized in a way that overlaps with the spelled-out
# form. e.g. "ISMS Policy" → tokens {information, security, management,
# system}; "Information Security Management System" → same set.
_TITLE_SYNONYMS = {
    "isms": "information security management system",
    "infosec": "information security",
    "infosecurity": "information security",
    "bcp": "business continuity plan",
    "bcm": "business continuity management",
    "drp": "disaster recovery plan",
    "dr": "disaster recovery",
    "iam": "identity access management",
    "mfa": "multi factor authentication",
    "2fa": "two factor authentication",
    "ir": "incident response",
    "irm": "incident response management",
    "ac": "access control",
    "acm": "access control management",
    "cm": "change management",
    "vm": "vulnerability management",
    "tprm": "third party risk management",
    "vrm": "vendor risk management",
    "dlp": "data loss prevention",
    "byod": "bring your own device",
    "soc": "security operations centre",
    "soa": "statement of applicability",
    "ai": "artificial intelligence",
    "rmp": "risk management plan",
    "ssp": "system security plan",
    "appsec": "application security",
    "siem": "security information event management",
}


def _title_token_set(title: str) -> set:
    """Tokenize a title for similarity comparison: lowercase, expand synonyms,
    drop stopwords and short tokens. Returns a set of meaningful words.

    "ISMS Policy v2" → {information, security, management, system}
    "Information Security Management System Policy" → {information, security, management, system}
    """
    if not title:
        return set()
    raw = title.lower()
    # Replace non-alphanumerics with spaces, then split.
    cleaned = "".join(ch if ch.isalnum() else " " for ch in raw)
    tokens = []
    for tok in cleaned.split():
        # Expand abbreviations into their constituent words.
        expansion = _TITLE_SYNONYMS.get(tok)
        if expansion:
            tokens.extend(expansion.split())
        else:
            tokens.append(tok)
    # Drop stopwords + 1-char numerics like "v2" leftover digits.
    return {
        t for t in tokens
        if t not in _TITLE_STOPWORDS and len(t) > 1 and not t.isdigit()
    }


def _title_similarity(a: str, b: str) -> float:
    """Jaccard similarity over the tokenized representations. 1.0 = identical
    meaningful tokens, 0.0 = nothing in common."""
    ta, tb = _title_token_set(a), _title_token_set(b)
    if not ta or not tb:
        return 0.0
    inter = ta & tb
    union = ta | tb
    return len(inter) / len(union) if union else 0.0


# When a suggestion's tokenized similarity to ANY existing title is at or
# above this threshold, treat it as a duplicate without calling the AI.
# Conservative threshold: this stage only catches *very* close matches
# (>= 85% Jaccard overlap on meaningful tokens). The semantic GPT stage
# downstream catches the looser cases with full context.
#
# Why so conservative: at 0.5 we hit false positives like
# "Information Security Management Policy" vs "Supplier Security
# Management Policy" — 2 of 4 shared tokens (security, management) =
# 0.5 = exact threshold hit, but they're semantically different policies
# (internal IT security vs third-party/vendor security). Raising to 0.85
# keeps obvious cases (e.g. identical token sets after synonym expansion)
# while routing genuine gray-area calls to the semantic stage that can
# reason about scope.
_TITLE_SIMILARITY_HARD_DUP = 0.85


# Document statuses that count as "the org already has this on the books".
# Drafts and pending-review docs count too, because the user is actively
# building them — re-suggesting them would be noise. Archived/expired/exception
# rows are intentionally excluded so a retired policy gets re-recommended.
_EXISTING_DOC_STATUSES = {
    "draft", "pending_review", "pending_approval", "approved", "published"
}


def _semantic_dedup_against_existing(
    client: OpenAI,
    suggestions: List[Dict],
    existing_documents: List[Dict],
) -> Tuple[set, List[Dict]]:
    """Ask GPT-4o which of the new suggestions are substantially equivalent
    to a document the tenant already has.

    Returns (set_of_indices_to_drop, list_of_match_dicts). Each match dict
    has shape {suggested_title, matched_existing_id, matched_existing_title,
    reason}. Failure-mode is "drop nothing" — we never want a flaky AI call
    to silently remove legitimate suggestions, so any error returns empty
    results.
    """
    if not suggestions or not existing_documents:
        return set(), []

    suggestion_payload = [
        {
            "i": idx,
            "title": str(s.get("title") or "").strip(),
            "doc_type": str(s.get("doc_type") or "").strip(),
            "description": str(s.get("description") or "")[:300],
        }
        for idx, s in enumerate(suggestions)
    ]
    existing_payload = [
        {
            "title": d.get("title") or "",
            "doc_type": d.get("doc_type") or "",
        }
        for d in existing_documents
    ]

    system_msg = (
        "You are a governance librarian. You output ONLY strict JSON. Your "
        "job is to AGGRESSIVELY identify duplicate document suggestions. "
        "When in doubt, mark as duplicate — the user is frustrated when the "
        "system suggests something they already have. False-positive (over-"
        "deduplicating) is preferable to false-negative (re-suggesting an "
        "existing doc)."
    )
    user_msg = f"""Two lists below: NEW (suggested documents to potentially add) and EXISTING (documents the organisation already maintains).

For each NEW item, decide whether it COVERS THE SAME PRIMARY TOPIC as ANY EXISTING item. If yes → duplicate, drop it.

DUPLICATE (drop) when ANY of these is true:
- Same subject, different abbreviation: "Information Security Policy" vs "InfoSec Policy" vs "ISMS Policy" vs "ISMS Manual"
- Same subject, reordered words: "Access Control Policy" vs "Policy on Access Control" vs "User Access Management Policy"
- Same subject, different doc_type label: existing "Access Control Procedure" vs new "Access Control Policy" — same scope, drop
- Sub-topic the parent obviously covers: existing "Information Security Policy" vs new "Password Policy" / "Encryption Policy" / "Acceptable Use Policy" → drop unless the new one is a deeply specialised standard/procedure with operational detail beyond the parent
- Renamed equivalent: "Disaster Recovery Plan" vs "Business Continuity Plan" when only one of the two would normally be maintained
- Singular/plural, hyphen, casing, version variants: "Risk Management Policy" vs "Risk Management Policies"
- Localized variants: "Data Protection Policy" vs "Privacy Policy" vs "GDPR Policy" — same primary topic
- Standards covering same domain: existing "Cryptographic Controls Standard" vs new "Encryption Standard" — same topic

NOT duplicate (keep):
- Genuinely distinct primary topics: "Acceptable Use Policy" vs "Vendor Risk Management Policy"
- Materially narrower OPERATIONAL document with concrete procedure not in the parent: existing "Information Security Policy" vs new "Vulnerability Scanning Procedure" → keep, the procedure has operational steps the policy doesn't
- Different lifecycle phase: existing "Incident Response Policy" vs new "Incident Response Tabletop Exercise Procedure" → keep

Return STRICT JSON ONLY:
{{"duplicates": [{{"i": <int index from NEW>, "matches_existing_title": "<existing title>", "reason": "<one short sentence>"}}]}}

Only include entries that ARE duplicates. Use the integer "i" from NEW exactly as given.

EXISTING:
{json.dumps(existing_payload, ensure_ascii=False)}

NEW:
{json.dumps(suggestion_payload, ensure_ascii=False)}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},
            temperature=0.0,
            max_tokens=2000,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        # Build a title → existing-doc lookup so we can attach the matched
        # existing record to each drop. Used by callers to render the skip
        # alongside the doc that justified it.
        existing_by_title = {(d.get("title") or "").strip().lower(): d for d in existing_documents}
        dropped_indices: set = set()
        match_records: List[Dict] = []
        for entry in (parsed.get("duplicates") or []):
            try:
                idx = int(entry.get("i"))
            except (TypeError, ValueError):
                continue
            if not (0 <= idx < len(suggestions)):
                continue
            suggested_title = (suggestions[idx].get("title") or "").strip()
            matched_title = (entry.get("matches_existing_title") or "").strip()
            reason_text = (entry.get("reason") or "").strip()[:200]
            matched_doc = existing_by_title.get(matched_title.lower()) if matched_title else None
            # Only emit a skip when we can pin it to a real existing doc.
            # If the AI returned a duplicate flag but the matched title
            # doesn't resolve to one of EXISTING, we drop the entry — the
            # user explicitly asked us not to show inaccurate skips.
            if not matched_doc or not suggested_title:
                continue
            dropped_indices.add(idx)
            match_records.append({
                "suggested_title": suggested_title,
                "matched_existing_id": matched_doc.get("id"),
                "matched_existing_title": matched_doc.get("title"),
                "matched_existing_doc_type": matched_doc.get("doc_type"),
                "matched_existing_status": matched_doc.get("status"),
                "reason": reason_text or "Semantic match against an existing document.",
                "match_type": "semantic",
            })
        return dropped_indices, match_records
    except Exception:
        # Defensive: never let a flaky dedup call hide real suggestions.
        return set(), []


def _fetch_existing_documents_for_frameworks(
    framework_ids: List[int],
    tenant_ids: List[int],
    requested_doc_type: Optional[str],
    db: Session,
) -> List[Dict]:
    """Return active documents in the tenant's library, used as the
    deduplication pool for AI policy suggestions.

    Important: we intentionally do NOT filter by framework_ids here. Many
    tenants upload policies/procedures without tagging them to a specific
    framework, so a framework-scoped query would miss "Information Security
    Policy" sitting untagged in the library and the AI would happily
    re-suggest it. Dedup against the WHOLE library; the cost is small (a
    tenant rarely has more than a few hundred docs) and the false-negative
    cost (re-suggesting a doc that already exists) is what the user is
    complaining about.

    Documents tagged to one of the requested frameworks are surfaced first
    in the response (so the UI's "Already covered" panel still shows the
    framework-relevant ones at the top).
    """
    if not tenant_ids:
        return []

    query = db.query(GovernanceDocument).filter(
        GovernanceDocument.tenant_id.in_(tenant_ids),
    )
    if requested_doc_type:
        query = query.filter(GovernanceDocument.doc_type == requested_doc_type)

    rows = query.all()
    fid_set = set(framework_ids or [])
    out: List[Dict] = []
    for d in rows:
        if (d.status or "").lower() not in _EXISTING_DOC_STATUSES:
            continue
        # Best-effort normalize JSON framework_ids. Tolerant of stringy ids
        # stored in older rows.
        normalized: set = set()
        for v in (d.framework_ids or []):
            try:
                normalized.add(int(v))
            except (TypeError, ValueError):
                continue
        # Whether this doc is tagged to one of the requested frameworks
        # (used purely for ranking — it's still in the dedup pool either way).
        is_framework_tagged = bool(fid_set & normalized) if fid_set else False
        out.append({
            "id": d.id,
            "title": d.title,
            "doc_type": d.doc_type,
            "status": d.status,
            "framework_ids": list(normalized),
            "is_framework_tagged": is_framework_tagged,
        })
    # Surface framework-tagged docs first, then everything else. Stable order
    # within each group.
    out.sort(key=lambda d: (0 if d["is_framework_tagged"] else 1, d["title"] or ""))
    return out


@router.post("/ai-suggest-policies")
def suggest_policies_for_framework(
    request: PolicySuggestRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Given framework IDs, use AI to suggest what policies, procedures, and standards can be developed.

    The suggestion is gap-aware: existing tenant documents linked to the same
    framework(s) are listed in the prompt as "already covered", and any AI
    suggestion whose normalized title collides with an existing title is
    filtered out before the response is returned.
    """
    if not request.framework_ids:
        raise HTTPException(status_code=400, detail="At least one framework must be selected")

    frameworks = db.query(UploadedFramework).filter(
        UploadedFramework.id.in_(request.framework_ids)
    ).all()

    if not frameworks:
        raise HTTPException(status_code=404, detail="No frameworks found")

    controls_summary, framework_alignment, domain_context = build_framework_control_context(frameworks, db)
    
    if not AI_INTEGRATIONS_OPENAI_API_KEY or not AI_INTEGRATIONS_OPENAI_BASE_URL:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI integration not configured"
        )
    
    client = OpenAI(
        api_key=AI_INTEGRATIONS_OPENAI_API_KEY,
        base_url=AI_INTEGRATIONS_OPENAI_BASE_URL
    )
    
    framework_names = ", ".join([f.name for f in frameworks])
    requested_doc_type = (request.doc_type or "").strip().lower() or None
    allowed_doc_types = {"policy", "standard", "procedure", "guideline", "charter", "framework"}
    if requested_doc_type and requested_doc_type not in allowed_doc_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported doc_type '{request.doc_type}'. Allowed values: {', '.join(sorted(allowed_doc_types))}"
        )

    # Pull the documents the tenant already has covering these framework(s)
    # so we can both (a) tell the model not to re-suggest them and (b) hard
    # filter any near-duplicates the model emits anyway.
    user_tenant_ids = get_user_tenants(current_user, db)
    existing_documents = _fetch_existing_documents_for_frameworks(
        request.framework_ids, user_tenant_ids, requested_doc_type, db
    )
    existing_normalized_titles = {
        _normalize_doc_title(d["title"]) for d in existing_documents if d.get("title")
    }
    existing_documents_section = ""
    if existing_documents:
        # Compact representation — title + type is enough to dedupe against.
        # Group by doc_type so the prompt reads cleanly even with 200+ rows.
        grouped: Dict[str, List[str]] = {}
        for d in existing_documents:
            grouped.setdefault(d.get("doc_type") or "document", []).append(d["title"])
        lines = []
        for dt, titles in sorted(grouped.items()):
            lines.append(f"- {dt}s ({len(titles)}): " + "; ".join(sorted(set(titles))))
        existing_documents_section = (
            "\n\nDOCUMENTS THE ORG ALREADY HAS (do NOT re-suggest these or close paraphrases):\n"
            + "\n".join(lines)
        )

    doc_type_instruction = (
        f"Suggest only {requested_doc_type} documents. Do not include any other document types."
        if requested_doc_type
        else "Cover policy, standard, procedure, and guideline needs systematically."
    )
    suggestion_count_instruction = (
        "Provide 12-20 suggestions, ordered by priority (high first)."
        if requested_doc_type
        else "Provide 18-30 suggestions covering policies, procedures, standards, and guidelines. Order by priority (high first) and ensure the list covers the major domains implied by the controls. The output must be domain-balanced rather than over-indexed on generic information security titles."
    )

    prompt = f"""You are an expert governance and compliance consultant. Based on the following regulatory framework controls, suggest the governance documents an organization should create to achieve domain-specific compliance coverage.

{controls_summary}{existing_documents_section}

Analysis requirements:
- Cluster the controls into governance domains before suggesting documents.
- Suggest documents only when they are justified by the control themes.
- DO NOT suggest a document that is already on the "DOCUMENTS THE ORG ALREADY HAS" list above. Skip it AND skip close paraphrases (e.g. if "Information Security Policy" exists, do not suggest "ISMS Policy" or "InfoSec Policy"; if "Access Control Policy" exists, do not suggest "User Access Management Policy"). Treat the existing list as the canonical coverage for those topics — only propose documents that fill genuine gaps.
- Prefer specific domain-aware titles over generic titles.
- {doc_type_instruction}
- Avoid duplicate or overlapping suggestions unless one is clearly a parent policy and another is a detailed procedure.
- Keep the suggested document titles general and reusable. Do not make them company-specific.
- Make procedures and standards domain-specific and operationally meaningful, not generic supporting documents.
- Where applicable, distinguish parent policy documents from subordinate standards, procedures, or guidelines.
- Cover these major inferred domains: {', '.join(domain_context) if domain_context else 'Use the control themes provided.'}

For each suggestion, provide:
- A clear, professional title
- A detailed description of what the document should cover
- The primary governance domain it belongs to
- The priority level (high, medium, low) based on regulatory importance
- Which specific control references from the framework it addresses
- Why this document is needed for the cited control themes
- The key sections the document should contain, reflecting ISO-style structure and the relevant domain obligations

Return a JSON object with this structure:
{{
  "framework_names": "{framework_names}",
  "suggestions": [
    {{
      "doc_type": "policy",
      "title": "Information Security Policy",
      "description": "Establishes the organization's approach to information security, defining objectives, principles, and responsibilities for protecting information assets.",
            "domain": "Information Security Governance",
      "priority": "high",
      "relevant_controls": ["Control-1.1", "Control-1.2"],
                        "coverage_rationale": "Addresses governance, ownership, and protection requirements in the cited controls.",
      "key_sections": ["Purpose", "Scope", "Policy Statements", "Roles and Responsibilities", "Compliance", "Review"]
    }}
  ],
  "total_suggestions": 0
}}

{suggestion_count_instruction}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert governance and compliance consultant. Always respond with valid JSON and prioritize domain accuracy, clause coverage, and non-generic titles."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=12000
        )
        
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        if isinstance(result, dict) and isinstance(result.get("suggestions"), list):
            if requested_doc_type:
                result["suggestions"] = [
                    suggestion
                    for suggestion in result.get("suggestions", [])
                    if str(suggestion.get("doc_type", "")).strip().lower() == requested_doc_type
                ]

            raw_suggestions = result.get("suggestions", [])

            # Three-stage dedup pipeline against the tenant's existing
            # documents:
            #  1. Exact normalized-title match — covers casing/punctuation/
            #     version-suffix variants. Free, cheap.
            #  2. Token-set Jaccard with synonym expansion — covers cases
            #     like "ISMS Policy" ≈ "Information Security Policy" and
            #     "InfoSec Policy" ≈ "Information Security Policy". Free.
            #  3. Semantic AI dedup — for the remaining suggestions, ask
            #     GPT-4o whether each is substantially equivalent to any
            #     existing document. Catches the long-tail cases the
            #     heuristics miss without AI hallucinations.
            existing_titles_for_compare = [d["title"] for d in existing_documents if d.get("title")]
            existing_norm_lookup = {
                _normalize_doc_title(d["title"]): d for d in existing_documents if d.get("title")
            }

            # Every entry in skipped_matches carries the *matched existing
            # doc* so the UI never has to guess which platform document
            # justified the skip. If no real existing doc can be pinned to
            # a suggestion, that suggestion is NOT recorded as a skip
            # (we'd rather pass through to the AI list than show an
            # inaccurate "we skipped X" claim — which was the prior bug).
            after_stage1: List[Dict] = []
            skipped_matches: List[Dict] = []

            stage1_count = 0
            for suggestion in raw_suggestions:
                title = str(suggestion.get("title") or "").strip()
                norm = _normalize_doc_title(title)
                matched = existing_norm_lookup.get(norm) if norm else None
                if matched and title:
                    skipped_matches.append({
                        "suggested_title": title,
                        "matched_existing_id": matched.get("id"),
                        "matched_existing_title": matched.get("title"),
                        "matched_existing_doc_type": matched.get("doc_type"),
                        "matched_existing_status": matched.get("status"),
                        "reason": "Identical normalised title already exists.",
                        "match_type": "exact_normalized",
                    })
                    stage1_count += 1
                    continue
                after_stage1.append(suggestion)

            stage2_count = 0
            after_stage2: List[Dict] = []
            for suggestion in after_stage1:
                title = str(suggestion.get("title") or "").strip()
                matched_doc = None
                best_score = 0.0
                for existing in existing_documents:
                    existing_title = existing.get("title") or ""
                    if not existing_title:
                        continue
                    score = _title_similarity(title, existing_title)
                    if score >= _TITLE_SIMILARITY_HARD_DUP and score > best_score:
                        best_score = score
                        matched_doc = existing
                if matched_doc and title:
                    skipped_matches.append({
                        "suggested_title": title,
                        "matched_existing_id": matched_doc.get("id"),
                        "matched_existing_title": matched_doc.get("title"),
                        "matched_existing_doc_type": matched_doc.get("doc_type"),
                        "matched_existing_status": matched_doc.get("status"),
                        "reason": f"Title shares {int(best_score * 100)}% of meaningful tokens with an existing document.",
                        "match_type": "token_overlap",
                    })
                    stage2_count += 1
                    continue
                after_stage2.append(suggestion)

            # Stage 3: semantic dedup pass. Skip the call entirely when
            # there's nothing to compare against or no remaining
            # suggestions — saves a round-trip and tokens.
            after_stage3 = after_stage2
            stage3_count = 0
            if existing_titles_for_compare and after_stage2:
                semantic_skip_ids, semantic_matches = _semantic_dedup_against_existing(
                    client=client,
                    suggestions=after_stage2,
                    existing_documents=existing_documents,
                )
                if semantic_skip_ids:
                    after_stage3 = [s for i, s in enumerate(after_stage2) if i not in semantic_skip_ids]
                    skipped_matches.extend(semantic_matches)
                    stage3_count = len(semantic_matches)

            result["suggestions"] = after_stage3
            result["total_suggestions"] = len(after_stage3)

            result["framework_alignment"] = framework_alignment
            result["domains_covered"] = domain_context
            # Surface the gap-awareness summary so the UI can show "12 already
            # covered" alongside "18 missing documents to create".
            result["already_covered"] = [
                {"id": d["id"], "title": d["title"], "doc_type": d["doc_type"], "status": d["status"]}
                for d in existing_documents
            ]
            result["already_covered_count"] = len(existing_documents)
            # New authoritative field: only contains skips that are pinned
            # to a real existing doc. The UI renders this directly.
            result["skipped_matches"] = skipped_matches
            # Legacy: titles only. Kept for any old client that still
            # reads it. Derived from skipped_matches so it can't drift.
            result["skipped_duplicate_titles"] = [m["suggested_title"] for m in skipped_matches]
            result["skipped_breakdown"] = {
                "exact_normalized": stage1_count,
                "token_overlap": stage2_count,
                "semantic": stage3_count,
            }
        return result
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse AI response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI API error: {error_msg}"
        )


@router.post("/ai-draft")
def generate_policy_ai_draft(
    request: PolicyAIDraftRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Kick off an async drafting job.

    The synchronous pipeline took 30–90s for bank-grade output and was
    blowing past HTTP timeouts at the proxy. This endpoint now:

      1. Validates the request (parent doc, tenant context).
      2. Allocates a Redis-tracked job id.
      3. Dispatches `ai_drafting.generate_draft` to Celery.
      4. Returns `{job_id, status: 'queued'}` immediately.

    The frontend polls `GET /governance/documents/ai-draft-jobs/{job_id}`
    every second or two until `status == 'completed'|'failed'`.

    Falls back to in-process execution if Celery isn't reachable — the
    user-visible behaviour is identical, just slower without an opt-out.
    """
    tenant_id = get_user_primary_tenant(current_user, db)
    if not tenant_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="No tenant context for drafting",
        )

    # Validate parent document up-front so the user gets immediate
    # feedback rather than a failed job 60s later.
    if request.parent_document_id:
        user_tenants = get_user_tenants(current_user, db)
        parent_exists = db.query(GovernanceDocument.id).filter(
            GovernanceDocument.id == request.parent_document_id,
            GovernanceDocument.tenant_id.in_(user_tenants),
        ).first()
        if not parent_exists:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Parent document not found",
            )

    from ....tasks.ai_drafting import create_job

    request_payload = {
        "tenant_id": tenant_id,
        "doc_type": request.doc_type,
        "title": request.title,
        "description": request.description,
        "framework_ids": request.framework_ids or [],
        "parent_document_id": request.parent_document_id,
        "include_sections": request.include_sections or [],
    }
    job_id = create_job(
        tenant_id=tenant_id,
        request_summary={
            "doc_type": request.doc_type,
            "title": request.title,
            "framework_ids": request.framework_ids or [],
        },
    )

    # Resolve tenant slug for the TenantTask first-arg contract.
    from ....db import MasterSession
    from ....models import Tenant as MasterTenant
    master = MasterSession()
    try:
        row = master.query(MasterTenant.slug).filter(MasterTenant.id == tenant_id).first()
        tenant_slug = row[0] if row else None
    finally:
        master.close()

    if not tenant_slug:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Could not resolve tenant slug for async drafting",
        )

    # Run in a daemon thread inside this FastAPI process. This avoids the
    # "broker accepts the task but no worker is consuming" failure mode
    # we hit when relying on Celery. The thread writes progress to Redis
    # the same way the Celery task would; the polling endpoint is
    # identical for both code paths.
    from ....tasks.ai_drafting import dispatch_in_thread
    dispatch_in_thread(tenant_slug, job_id, request_payload)
    return {
        "job_id": job_id,
        "status": "queued",
        "poll_url": f"/governance/documents/ai-draft-jobs/{job_id}",
    }


@router.get("/ai-draft-jobs/{job_id}")
def get_ai_draft_job(
    job_id: str,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    """Poll a running / completed drafting job.

    Returns the same payload `generate_draft` writes to Redis — frontend
    keeps polling until `status in ('completed', 'failed')` and then
    reads the `result` field for the generated document.
    """
    from ....tasks.ai_drafting import get_job

    tenant_id = get_user_primary_tenant(current_user, db)
    payload = get_job(job_id)
    if not payload:
        raise HTTPException(status_code=404, detail="Drafting job not found or expired")
    # Tenant scoping — never let a tenant peek at another tenant's job.
    if payload.get("tenant_id") and payload["tenant_id"] != tenant_id:
        raise HTTPException(status_code=404, detail="Drafting job not found or expired")
    return payload


class CompareWithDocumentBody(BaseModel):
    target_document_id: int


@router.post("/{document_id}/compare-with-document")
def compare_with_document(
    document_id: int,
    body: CompareWithDocumentBody,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    """Side-by-side comparison of two governance documents owned by the tenant,
    with an optional AI gap analysis when OpenAI is configured.

    Mirrors the response shape of `/governance/nca-templates/{id}/compare` so
    the frontend compare modal can render either source identically.
    """
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        raise HTTPException(status_code=403, detail="No tenant access")

    source = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants),
    ).first()
    if not source:
        raise HTTPException(status_code=404, detail="Source document not found")

    target = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == body.target_document_id,
        GovernanceDocument.tenant_id.in_(user_tenants),
    ).first()
    if not target:
        raise HTTPException(status_code=404, detail="Target document not found")

    if source.id == target.id:
        raise HTTPException(status_code=400, detail="Cannot compare a document with itself")

    source_content = (source.content or "").strip() or "(empty document)"
    target_content = (target.content or "").strip() or "(empty document)"

    gap_analysis = None
    if AI_INTEGRATIONS_OPENAI_API_KEY:
        try:
            client_kwargs = {"api_key": AI_INTEGRATIONS_OPENAI_API_KEY}
            if AI_INTEGRATIONS_OPENAI_BASE_URL:
                client_kwargs["base_url"] = AI_INTEGRATIONS_OPENAI_BASE_URL
            client = OpenAI(**client_kwargs)

            prompt = f"""Compare two governance documents. Identify gaps, missing sections, redundancies, and areas where one is weaker or stronger than the other. Be concise and specific.

SOURCE DOCUMENT ({source.title} — {source.doc_type}):
{source_content[:8000]}

TARGET DOCUMENT ({target.title} — {target.doc_type}):
{target_content[:8000]}

Return strict JSON with keys:
- summary (2-3 sentences)
- missing_from_source (array of items present in target but absent in source)
- missing_from_target (array of items present in source but absent in target)
- alignment_score (0-100 integer, how well the two documents align in scope and content)
- recommended_additions (array of 3-7 specific clauses/sections the source should add to fully align with the target)
"""

            completion = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0.3,
            )
            gap_raw = completion.choices[0].message.content or "{}"
            gap_analysis = json.loads(gap_raw)
            # Map to the same shape the NCA-template compare endpoint returns
            if "missing_from_source" in gap_analysis:
                gap_analysis["missing_from_user_document"] = gap_analysis.pop("missing_from_source")
            if "missing_from_target" in gap_analysis:
                gap_analysis["present_in_user_only"] = gap_analysis.pop("missing_from_target")
        except Exception:
            import logging
            logging.getLogger(__name__).exception("Document-vs-document AI gap analysis failed")

    return {
        # 'template' key kept for frontend shape parity with the NCA compare endpoint
        "template": {
            "id": f"doc-{target.id}",
            "title": target.title,
            "category": target.doc_type or "Document",
            "content": target_content,
        },
        "document": {
            "id": source.id,
            "title": source.title,
            "doc_type": source.doc_type,
            "content": source_content,
        },
        "gap_analysis": gap_analysis,
    }


@router.get("/{document_id}/view-html")
def get_document_html(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Convert uploaded document to HTML for in-platform viewing"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # If document has a file, always prefer file-based conversion (better HTML)
    if document.file_path and os.path.exists(document.file_path):
        html_content = ""
        file_type = document.file_type or ""
        
        try:
            if file_type == "pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(document.file_path)
                sections = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        lines = page_text.split('\n')
                        page_html_parts = []
                        for line in lines:
                            stripped = line.strip()
                            if not stripped:
                                continue
                            if stripped.isupper() and len(stripped) < 100:
                                page_html_parts.append(f'<h3 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-weight: 600;">{html.escape(stripped)}</h3>')
                            elif len(stripped) < 80 and not stripped.endswith('.') and not stripped.endswith(','):
                                page_html_parts.append(f'<h4 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(stripped)}</h4>')
                            else:
                                page_html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
                        
                        sections.append(f'<div class="page" style="margin-bottom: 2rem; padding-bottom: 1rem; border-bottom: 1px solid #334155;"><div style="color: #475569; font-size: 0.75rem; margin-bottom: 0.5rem;">Page {i+1}</div>{"".join(page_html_parts)}</div>')
                html_content = "".join(sections)
            
            elif file_type in ("docx", "doc"):
                from docx import Document as DocxDocument
                doc = DocxDocument(document.file_path)
                parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = (para.style.name or "").lower() if para.style else ""
                    if "heading 1" in style_name:
                        parts.append(f'<h1 style="color: #f1f5f9; margin-top: 2rem; margin-bottom: 0.75rem; font-size: 1.5rem; font-weight: 700;">{html.escape(text)}</h1>')
                    elif "heading 2" in style_name:
                        parts.append(f'<h2 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-size: 1.25rem; font-weight: 600;">{html.escape(text)}</h2>')
                    elif "heading 3" in style_name:
                        parts.append(f'<h3 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.5rem; font-size: 1.1rem; font-weight: 500;">{html.escape(text)}</h3>')
                    elif "heading" in style_name:
                        parts.append(f'<h4 style="color: #cbd5e1; margin-top: 0.75rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(text)}</h4>')
                    elif "title" in style_name:
                        parts.append(f'<h1 style="color: #f1f5f9; margin-top: 1rem; margin-bottom: 1rem; font-size: 1.75rem; font-weight: 700; text-align: center;">{html.escape(text)}</h1>')
                    elif "list" in style_name:
                        parts.append(f'<li style="color: #94a3b8; margin-left: 1.5rem; margin-bottom: 0.25rem; line-height: 1.6;">{html.escape(text)}</li>')
                    else:
                        is_bold = all(run.bold for run in para.runs if run.text.strip()) if para.runs else False
                        if is_bold and len(text) < 100:
                            parts.append(f'<h4 style="color: #cbd5e1; margin-top: 0.75rem; margin-bottom: 0.25rem; font-weight: 600;">{html.escape(text)}</h4>')
                        else:
                            parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(text)}</p>')
                
                for table in doc.tables:
                    table_html = '<table style="width: 100%; border-collapse: collapse; margin: 1rem 0; border: 1px solid #334155;">'
                    for i, row in enumerate(table.rows):
                        table_html += '<tr>'
                        for cell in row.cells:
                            tag = 'th' if i == 0 else 'td'
                            style = 'padding: 0.5rem; border: 1px solid #334155; color: #94a3b8;'
                            if i == 0:
                                style += ' background-color: #1e293b; color: #e2e8f0; font-weight: 600;'
                            table_html += f'<{tag} style="{style}">{html.escape(cell.text.strip())}</{tag}>'
                        table_html += '</tr>'
                    table_html += '</table>'
                    parts.append(table_html)
                
                html_content = "".join(parts)
            
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        except HTTPException:
            raise
        except Exception as e:
            if document.content:
                lines = document.content.split('\n')
                html_parts = []
                for line in lines:
                    stripped = line.strip()
                    if not stripped:
                        continue
                    html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
                return {
                    "document_id": document.id,
                    "title": document.title,
                    "content_type": "html",
                    "html": "".join(html_parts),
                    "file_type": document.file_type
                }
            raise HTTPException(status_code=500, detail=f"Error converting document: {str(e)}")
        
        return {
            "document_id": document.id,
            "title": document.title,
            "content_type": "html",
            "html": html_content,
            "file_type": file_type
        }
    
    # Fall back to content field if no file
    if document.content:
        lines = document.content.split('\n')
        html_parts = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.isupper() and len(stripped) < 100:
                html_parts.append(f'<h3 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-weight: 600;">{html.escape(stripped)}</h3>')
            elif len(stripped) < 80 and not stripped.endswith('.') and not stripped.endswith(','):
                html_parts.append(f'<h4 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(stripped)}</h4>')
            else:
                html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
        return {
            "document_id": document.id,
            "title": document.title,
            "content_type": "html",
            "html": "".join(html_parts),
            "file_type": document.file_type
        }

    raise HTTPException(status_code=404, detail="Document file not found")


# ============================================================================
# Document annotations / remarks
# ----------------------------------------------------------------------------
# Auditors and reviewers can attach remarks to a document. Two anchor kinds
# are supported today:
#   - text_range : anchored to a character offset range in the document's
#                   plain-text representation (works for `content`-based
#                   docs, markdown, and the in-browser viewer's text mode)
#   - general    : free-form comment with no specific anchor
#
# Threading is intentionally NOT supported in v1 — flat annotations are
# easier to anchor and render; reply chains can be added later by
# introducing `parent_annotation_id` without a breaking change.
# ============================================================================

class AnnotationCreateRequest(BaseModel):
    anchor_kind: str = Field("general", pattern=r"^(text_range|general)$")
    anchor_data: Optional[Dict[str, Any]] = None
    comment: str = Field(..., min_length=1, max_length=4000)


class AnnotationUpdateRequest(BaseModel):
    comment: Optional[str] = Field(None, min_length=1, max_length=4000)
    status: Optional[str] = Field(None, pattern=r"^(open|resolved)$")


def _serialize_annotation(a: DocumentAnnotation, user_lookup: Dict[int, GRCUser]) -> dict:
    user = user_lookup.get(a.user_id)
    return {
        "id": a.id,
        "document_id": a.document_id,
        "anchor_kind": a.anchor_kind,
        "anchor_data": a.anchor_data or {},
        "comment": a.comment,
        "status": a.status,
        "user_id": a.user_id,
        "user_name": (user.display_name or user.email) if user else None,
        "created_at": a.created_at,
        "updated_at": a.updated_at,
    }


def _assert_doc_access(document_id: int, user_tenants: List[int], db: Session) -> GovernanceDocument:
    doc = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants),
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found in your tenant.")
    return doc


@router.get("/{document_id}/annotations")
def list_document_annotations(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    doc = _assert_doc_access(document_id, user_tenants, db)
    rows = db.query(DocumentAnnotation).filter(
        DocumentAnnotation.document_id == doc.id,
        DocumentAnnotation.tenant_id == doc.tenant_id,
    ).order_by(DocumentAnnotation.created_at.asc()).all()
    user_ids = list({r.user_id for r in rows})
    users = db.query(GRCUser).filter(GRCUser.id.in_(user_ids)).all() if user_ids else []
    user_lookup = {u.id: u for u in users}
    return {
        "annotations": [_serialize_annotation(a, user_lookup) for a in rows],
        "total": len(rows),
    }


@router.post("/{document_id}/annotations", status_code=status.HTTP_201_CREATED)
def create_document_annotation(
    document_id: int,
    payload: AnnotationCreateRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    doc = _assert_doc_access(document_id, user_tenants, db)
    annotation = DocumentAnnotation(
        tenant_id=doc.tenant_id,
        document_id=doc.id,
        user_id=current_user.id,
        anchor_kind=payload.anchor_kind,
        anchor_data=payload.anchor_data or {},
        comment=payload.comment.strip(),
        status="open",
    )
    db.add(annotation)
    db.commit()
    db.refresh(annotation)
    return _serialize_annotation(annotation, {current_user.id: current_user})


@router.put("/{document_id}/annotations/{annotation_id}")
def update_document_annotation(
    document_id: int,
    annotation_id: int,
    payload: AnnotationUpdateRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    doc = _assert_doc_access(document_id, user_tenants, db)
    annotation = db.query(DocumentAnnotation).filter(
        DocumentAnnotation.id == annotation_id,
        DocumentAnnotation.document_id == doc.id,
        DocumentAnnotation.tenant_id == doc.tenant_id,
    ).first()
    if not annotation:
        raise HTTPException(status_code=404, detail="Annotation not found.")
    # Comment edits restricted to the author; status can be flipped by
    # any tenant member (so a different reviewer can mark resolved).
    if payload.comment is not None:
        if annotation.user_id != current_user.id:
            raise HTTPException(status_code=403, detail="Only the author can edit this remark.")
        annotation.comment = payload.comment.strip()
    if payload.status is not None:
        annotation.status = payload.status
    db.commit()
    db.refresh(annotation)
    user = db.query(GRCUser).filter(GRCUser.id == annotation.user_id).first()
    return _serialize_annotation(annotation, {annotation.user_id: user} if user else {})


@router.delete("/{document_id}/annotations/{annotation_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document_annotation(
    document_id: int,
    annotation_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth),
):
    user_tenants = get_user_tenants(current_user, db)
    doc = _assert_doc_access(document_id, user_tenants, db)
    annotation = db.query(DocumentAnnotation).filter(
        DocumentAnnotation.id == annotation_id,
        DocumentAnnotation.document_id == doc.id,
        DocumentAnnotation.tenant_id == doc.tenant_id,
    ).first()
    if not annotation:
        raise HTTPException(status_code=404, detail="Annotation not found.")
    if annotation.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Only the author can delete this remark.")
    db.delete(annotation)
    db.commit()
    return None
