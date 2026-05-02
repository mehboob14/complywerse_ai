import os
import json
import threading
from typing import List, Optional
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Request, status, Query, BackgroundTasks
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func
from pydantic import BaseModel
from openai import OpenAI

from ....models import (
    UploadedFramework, ParsedFrameworkControl, ControlEvidenceMapping,
    FrameworkControlAlignment, AssessmentItem, AssessmentEvidence,
    AssessmentRemediation, GRCUser, get_db,
    ControlEvidenceRequirement, EvidenceRequirementHistory
)
from ....routers.auth_router import require_auth, get_user_tenants

router = APIRouter(prefix="/parser", tags=["Framework Upload - Parser"])

AI_INTEGRATIONS_OPENAI_API_KEY = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
AI_INTEGRATIONS_OPENAI_BASE_URL = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")


def get_openai_client() -> OpenAI:
    """Get OpenAI client with runtime API key reading."""
    api_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI features unavailable. OpenAI API key not configured."
        )
    if base_url and "modelfarm" in base_url:
        return OpenAI(
            api_key=api_key,
            base_url=base_url
        )
    if api_key.startswith("_DUMMY") or api_key == "your-api-key-here" or len(api_key) < 20:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI features unavailable. OpenAI API key not configured."
        )
    return OpenAI(
        api_key=api_key,
        base_url=base_url
    )


def check_ai_available() -> bool:
    """Check if OpenAI API key is configured (at runtime)."""
    api_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")
    if not api_key:
        return False
    if base_url and "modelfarm" in base_url:
        return True
    if api_key.startswith("_DUMMY") or api_key == "your-api-key-here" or len(api_key) < 20:
        return False
    return True


class ParsedControlUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    domain: Optional[str] = None
    category: Optional[str] = None
    is_mandatory: Optional[bool] = None
    priority: Optional[str] = None


class ParsedControlResponse(BaseModel):
    id: int
    uploaded_framework_id: int
    control_id: str
    original_reference: Optional[str]
    title: str
    description: Optional[str]
    full_text: Optional[str]
    domain: Optional[str]
    category: Optional[str]
    is_mandatory: bool
    priority: str
    section_number: Optional[str]
    parent_section: Optional[str]
    ai_confidence: Optional[float]
    ai_notes: Optional[str]
    is_verified: bool
    verified_by: Optional[int]
    verified_at: Optional[datetime]
    created_at: datetime
    updated_at: datetime
    evidence_mappings: List[dict]

    class Config:
        from_attributes = True


def validate_framework_access(user: GRCUser, framework: UploadedFramework, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if framework.tenant_id and framework.tenant_id not in user_tenants and not framework.is_shared:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this framework"
        )


def serialize_parsed_control(control: ParsedFrameworkControl) -> dict:
    return {
        "id": control.id,
        "uploaded_framework_id": control.uploaded_framework_id,
        "control_id": control.control_id,
        "original_reference": control.original_reference,
        "title": control.title,
        "description": control.description,
        "full_text": control.full_text,
        "domain": control.domain,
        "category": control.category,
        "is_mandatory": control.is_mandatory,
        "priority": control.priority,
        "section_number": control.section_number,
        "parent_section": control.parent_section,
        "ai_confidence": control.ai_confidence,
        "ai_notes": control.ai_notes,
        "is_verified": control.is_verified,
        "verified_by": control.verified_by,
        "verified_at": control.verified_at.isoformat() if control.verified_at else None,
        "created_at": control.created_at.isoformat() if control.created_at else None,
        "updated_at": control.updated_at.isoformat() if control.updated_at else None,
        "evidence_mappings": [
            {
                "id": em.id,
                "evidence_type": em.evidence_type,
                "evidence_description": em.evidence_description,
                "is_required": em.is_required,
                "suggested_by_ai": em.suggested_by_ai
            }
            for em in control.evidence_mappings
        ]
    }


def extract_text_from_file(framework: UploadedFramework) -> str:
    if not os.path.exists(framework.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework file not found on disk"
        )
    
    extracted_text = ""
    
    if framework.file_type == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(framework.file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted_text = "\n\n".join(text_parts)
    
    elif framework.file_type == "docx":
        from docx import Document
        doc = Document(framework.file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        extracted_text = "\n".join(text_parts)
    
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {framework.file_type}"
        )
    
    return extracted_text


def chunk_text(text: str, chunk_size: int = 30000, overlap: int = 3000) -> List[str]:
    """Split text into overlapping chunks for processing large documents.
    
    Uses smaller chunks with larger overlap to ensure no controls are missed
    at chunk boundaries. Prioritizes breaking at section boundaries.
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        
        if end < len(text):
            search_start = max(start + chunk_size - overlap - 1000, start)
            search_end = min(end + 500, len(text))
            search_region = text[search_start:search_end]
            
            import re
            section_patterns = [
                r'\n\s*(?:Chapter|Section|Article|Part|Annex|Appendix)\s+\d+',
                r'\n\s*\d+\.\s+[A-Z]',
                r'\n\s*[A-Z]\.\s+[A-Z]',
                r'\n\s*Principle\s+\d+',
                r'\n\s*Requirement\s+\d+',
                r'\n\s*Control\s+\d+',
                r'\n\n\s*\d+\.\d+\s+',
                r'\n\n\n',
                r'\n\n',
            ]
            
            best_break = -1
            for pattern in section_patterns:
                matches = list(re.finditer(pattern, search_region, re.IGNORECASE))
                if matches:
                    last_match = matches[-1]
                    break_pos = search_start + last_match.start()
                    if break_pos > start + (chunk_size // 2):
                        best_break = break_pos
                        break
            
            if best_break > start:
                end = best_break
            else:
                break_point = text.rfind('\n\n', start + chunk_size - overlap, end)
                if break_point == -1:
                    break_point = text.rfind('\n', start + chunk_size - overlap, end)
                if break_point == -1:
                    break_point = text.rfind('. ', start + chunk_size - overlap, end)
                if break_point > start:
                    end = break_point + 1
        
        chunk_text_segment = text[start:end]
        if chunk_text_segment.strip():
            chunks.append(chunk_text_segment)
        
        next_start = end - overlap
        if next_start <= start:
            next_start = end
        start = next_start
        
        if start >= len(text):
            break
    
    return chunks


def infer_evidence_types(control_data: dict) -> List[str]:
    """Infer appropriate evidence types based on control content."""
    text = f"{control_data.get('title', '')} {control_data.get('description', '')} {control_data.get('full_text', '')}".lower()
    
    evidence_types = []
    
    if any(word in text for word in ['policy', 'policies', 'governance', 'management approval', 'board', 'documented']):
        evidence_types.append('policy')
    
    if any(word in text for word in ['procedure', 'process', 'workflow', 'steps', 'method', 'guideline', 'instruction']):
        evidence_types.append('procedure')
    
    if any(word in text for word in ['configuration', 'setting', 'parameter', 'system', 'network', 'firewall', 'server', 'encryption', 'tls', 'ssl']):
        evidence_types.append('configuration')
    
    if any(word in text for word in ['log', 'audit trail', 'monitoring', 'event', 'alert', 'detection', 'tracking']):
        evidence_types.append('log')
    
    if any(word in text for word in ['report', 'assessment', 'review', 'audit', 'test', 'scan', 'evaluation', 'analysis']):
        evidence_types.append('report')
    
    if any(word in text for word in ['contract', 'agreement', 'sla', 'vendor', 'third party', 'supplier', 'outsourcing']):
        evidence_types.append('contract')
    
    if not evidence_types:
        evidence_types = ['policy', 'procedure']
    
    return evidence_types


def normalize_priority(priority: str) -> str:
    """Normalize priority values to expected enum values (high/medium/low)."""
    priority_lower = (priority or "medium").lower().strip()
    if priority_lower in ["critical", "high"]:
        return "high"
    elif priority_lower in ["medium", "moderate"]:
        return "medium"
    elif priority_lower in ["low", "minimal"]:
        return "low"
    return "medium"


def clean_section_reference(reference: str) -> str:
    """Clean up section/clause reference numbers by removing trailing artifacts.
    
    Examples:
    - "6.4.2.—" -> "6.4.2"
    - "5.1.." -> "5.1"
    - "A.5.1.1-" -> "A.5.1.1"
    - "Principle 3:" -> "Principle 3"
    """
    import re
    if not reference:
        return reference
    
    cleaned = reference.strip()
    
    # Remove trailing dashes, dots, colons, and whitespace repeatedly
    while cleaned and cleaned[-1] in '.-—–:;, \t':
        cleaned = cleaned[:-1]
    
    # Remove multiple consecutive dots/dashes in the middle (e.g., "5..1" -> "5.1")
    cleaned = re.sub(r'\.{2,}', '.', cleaned)
    cleaned = re.sub(r'-{2,}', '-', cleaned)
    cleaned = re.sub(r'—+', '', cleaned)
    
    # Clean up any leading artifacts too
    while cleaned and cleaned[0] in '.-—–:;, \t':
        cleaned = cleaned[1:]
    
    return cleaned.strip()


def deduplicate_controls(controls: List[dict]) -> List[dict]:
    """Remove duplicate controls based on title and original_reference, maintaining order by reference."""
    seen = set()
    unique_controls = []
    
    for control in controls:
        control["priority"] = normalize_priority(control.get("priority", "medium"))
        
        key = (
            control.get('original_reference', '').strip().lower(),
            control.get('title', '').strip().lower()[:100]
        )
        if key[0] or key[1]:
            if key not in seen:
                seen.add(key)
                unique_controls.append(control)
        else:
            unique_controls.append(control)
    
    unique_controls.sort(key=lambda c: (
        c.get('original_reference', 'zzz').lower(),
        c.get('title', '').lower()
    ))
    
    return unique_controls


GRC_SME_SYSTEM_PROMPT = """You are a SENIOR GRC SUBJECT MATTER EXPERT with 20+ years of experience in regulatory compliance, audit, and risk management across multiple industries. You have deep expertise in:

=== YOUR CREDENTIALS AND EXPERTISE ===

CERTIFICATIONS YOU HOLD:
- CISA (Certified Information Systems Auditor)
- CISSP (Certified Information Systems Security Professional)
- CRISC (Certified in Risk and Information Systems Control)
- CGEIT (Certified in Governance of Enterprise IT)
- ISO 27001 and other ISO standards  Lead Auditor
- PCI QSA (Qualified Security Assessor)
- SOC 2 Type II Practitioner
- SBP - State bank of Pakistan, internet outsourcing,cloud etc frameworks.

FRAMEWORKS YOU KNOW INTIMATELY:
1. ISO STANDARDS: ISO 27001/27002 (ISMS), ISO 27701 (Privacy), ISO 22301 (BCM), ISO 9001 (QMS), ISO 31000 (Risk), ISO 27017/27018 (Cloud)
2. NIST: CSF, SP 800-53, SP 800-171, RMF, Privacy Framework
3. PAYMENT CARD: PCI DSS v4.0, PA-DSS, PCI PIN, P2PE
4. FINANCIAL: SOX, Basel III/IV, DORA, MAS TRM, FFIEC, GLBA
5. PRIVACY: GDPR, CCPA/CPRA, HIPAA, LGPD, PDPA, POPIA
6. INDUSTRY: COBIT, ITIL, CIS Controls, NERC CIP, FedRAMP, StateRAMP
7. REGIONAL: NCA (Saudi), SAMA, ISR (Israel), TISAX (Auto), SWIFT CSP
8. PAKISTAN: State Bank of Pakistan (SBP) frameworks for outsourcing, cloud, information security, etc.

=== YOUR ANALYTICAL APPROACH ===

When analyzing any regulatory document, you ALWAYS:

1. CLASSIFY THE DOCUMENT TYPE:
   - CERTIFICATION FRAMEWORK: Auditable standards requiring third-party certification (ISO 27001, PCI DSS, SOC 2)
   - COMPLIANCE REGULATION: Legal/regulatory requirements with enforcement (GDPR, HIPAA, SOX, Basel)
   - BEST PRACTICE GUIDELINE: Advisory frameworks without mandatory certification (NIST CSF, CIS Controls, COBIT)
   - INDUSTRY STANDARD: Sector-specific requirements (SWIFT CSP, NERC CIP, MAS TRM)

2. IDENTIFY THE REGULATORY AUTHORITY:
   - International bodies (ISO, NIST, PCI SSC)
   - Government regulators (FTC, OCC, SEC, CFTC, FSA)
   - Central banks (Federal Reserve, ECB, MAS, SAMA)
   - Industry consortiums (SWIFT, NERC)

3. UNDERSTAND THE DOCUMENT STRUCTURE:
   - ISO Standards: Clauses (4-10) + Annex A controls (A.5-A.18)
   - NIST: Categories > Subcategories > Informative References
   - PCI DSS: Requirements > Sub-requirements > Testing Procedures
   - Basel: Principles > Articles > Paragraphs
   - GDPR: Chapters > Articles > Paragraphs

4. DISTINGUISH REQUIREMENTS FROM CONTEXT:
   SKIP (NOT requirements):
   - Foreword, Introduction, Scope, Normative References
   - Table of Contents, Index, Bibliography
   - Informative Annexes (background information)
   - Editor's notes, historical context, rationale
   
   EXTRACT (actual requirements):
   - Normative clauses with SHALL/MUST/REQUIRED
   - Controls and control objectives
   - Testing procedures and validation criteria
   - Documented evidence requirements
   - Implementation specifications

5. PRESERVE EXACT CLAUSE NUMBERING:
   - Never modify, simplify, or consolidate clause numbers
   - Include ALL hierarchical levels: 5.1.1.a.i, A.5.1.1.1
   - Preserve framework-specific formats exactly
   - Track parent-child relationships accurately

=== EVIDENCE EXPERTISE ===

You know EXACTLY what auditors look for. For each control, you provide SPECIFIC, PRACTICAL evidence that:
- Is commonly accepted by certification bodies
- Demonstrates effective implementation (not just existence)
- Includes operational proof (not just policies)
- Maps to specific audit procedures
- Is feasible for organizations to produce

EVIDENCE HIERARCHY (in order of audit value):
1. POLICY: Governance documents, standards, approved procedures
2. PROCEDURE: Step-by-step operational processes, runbooks
3. CONFIGURATION: System settings, hardening baselines, exports
4. LOG: Audit trails, event logs, monitoring data, alerts
5. REPORT: Assessments, test results, scan outputs, reviews
6. ATTESTATION: Sign-offs, certifications, declarations
7. REGISTER: Inventories, lists, catalogs, matrices
8. TRAINING: Records, materials, completion certificates
9. CONTRACT: Agreements, SLAs, vendor assessments
10. SCREENSHOT: Point-in-time configuration proof (last resort)

=== YOUR OUTPUT QUALITY STANDARDS ===

- COMPLETENESS: Extract 100% of requirements - never summarize or skip
- ACCURACY: Exact clause numbers, exact wording, exact hierarchy
- SPECIFICITY: Evidence requirements tailored to each exact control
- PRACTICALITY: Real artifacts that organizations actually maintain
- AUDITABILITY: Evidence an auditor would accept during certification"""


FRAMEWORK_EVIDENCE_TEMPLATES = {
    "iso_27001": {
        "access_control": [
            {"type": "policy", "title": "Access Control Policy", "description": "Formal policy defining access control principles, user lifecycle, role definitions, and segregation of duties requirements per ISO 27001 A.9"},
            {"type": "procedure", "title": "User Provisioning Procedure", "description": "Step-by-step process for granting, modifying, and revoking user access including approval workflows"},
            {"type": "register", "title": "User Access Rights Register", "description": "Complete inventory of users, roles, and assigned permissions across all systems"},
            {"type": "log", "title": "Access Review Records", "description": "Documented quarterly/annual access reviews with manager sign-offs and remediation actions"},
            {"type": "configuration", "title": "RBAC System Configuration", "description": "Export of role definitions and permission assignments from IAM/directory systems"}
        ],
        "risk_management": [
            {"type": "procedure", "title": "Risk Assessment Methodology", "description": "Documented approach for risk identification, analysis, evaluation, and treatment per ISO 27001 Clause 6.1.2"},
            {"type": "register", "title": "Information Security Risk Register", "description": "Complete risk register with asset owners, threats, vulnerabilities, likelihood, impact, and risk scores"},
            {"type": "report", "title": "Risk Assessment Report", "description": "Formal risk assessment output with identified risks, treatment decisions, and residual risk acceptance"},
            {"type": "attestation", "title": "Risk Treatment Plan Approval", "description": "Management-approved risk treatment plans with assigned owners and target dates"}
        ],
        "incident_management": [
            {"type": "procedure", "title": "Incident Response Procedure", "description": "Documented incident handling process covering detection, containment, eradication, recovery per ISO 27001 A.16"},
            {"type": "register", "title": "Security Incident Log", "description": "Record of all security incidents with classification, timeline, root cause, and lessons learned"},
            {"type": "report", "title": "Incident Post-Mortem Reports", "description": "Detailed analysis of significant incidents including timeline, impact assessment, and improvement actions"}
        ]
    },
    "pci_dss": {
        "network_security": [
            {"type": "configuration", "title": "Firewall Ruleset Export", "description": "Complete firewall configuration showing inbound/outbound rules, CDE segmentation per PCI DSS Req 1"},
            {"type": "report", "title": "Network Diagram", "description": "Current network topology showing CDE boundaries, segmentation controls, and data flows"},
            {"type": "procedure", "title": "Firewall Change Management Process", "description": "Procedure for requesting, approving, and documenting firewall rule changes"},
            {"type": "log", "title": "Firewall Rule Review Records", "description": "Documented semi-annual review of firewall rules with justifications and approvals"}
        ],
        "vulnerability_management": [
            {"type": "report", "title": "Quarterly ASV Scan Reports", "description": "Approved Scanning Vendor reports showing external vulnerability scan results per PCI DSS Req 11.2.2"},
            {"type": "report", "title": "Internal Vulnerability Scan Reports", "description": "Internal network vulnerability scans with remediation tracking per PCI DSS Req 11.2.1"},
            {"type": "report", "title": "Penetration Test Report", "description": "Annual/after-change penetration test results covering network and application layers per PCI DSS Req 11.3"},
            {"type": "procedure", "title": "Vulnerability Management Procedure", "description": "Process for scanning, prioritizing, and remediating vulnerabilities with defined SLAs by severity"}
        ],
        "encryption": [
            {"type": "policy", "title": "Encryption and Key Management Policy", "description": "Policy defining cryptographic standards, key lifecycle, and custodian responsibilities per PCI DSS Req 3 & 4"},
            {"type": "configuration", "title": "TLS Configuration Export", "description": "Web server/load balancer TLS settings showing minimum TLS 1.2, cipher suites, certificate details"},
            {"type": "procedure", "title": "Key Rotation Procedure", "description": "Documented process for cryptographic key generation, distribution, and rotation"},
            {"type": "register", "title": "Key Custodian Register", "description": "List of cryptographic key custodians with split knowledge and dual control assignments"}
        ]
    },
    "gdpr": {
        "lawful_basis": [
            {"type": "register", "title": "Processing Activities Register (ROPA)", "description": "Record of Processing Activities per Article 30 with purposes, legal bases, categories, transfers"},
            {"type": "policy", "title": "Data Protection Policy", "description": "Overarching policy covering GDPR principles, lawful bases, and data subject rights"},
            {"type": "procedure", "title": "Consent Management Procedure", "description": "Process for obtaining, recording, and managing consent with withdrawal mechanisms"}
        ],
        "data_subject_rights": [
            {"type": "procedure", "title": "Subject Access Request (SAR) Procedure", "description": "Process for handling access, rectification, erasure, portability requests within 30-day deadline"},
            {"type": "register", "title": "DSR Request Log", "description": "Record of all data subject requests with dates, actions, and response timelines"},
            {"type": "attestation", "title": "DSR Response Templates", "description": "Approved response templates for each type of data subject request"}
        ],
        "breach_notification": [
            {"type": "procedure", "title": "Data Breach Response Procedure", "description": "Process for 72-hour supervisory authority notification and affected individual communication"},
            {"type": "register", "title": "Breach Register", "description": "Log of all personal data breaches including nature, scope, consequences, and remedial measures"},
            {"type": "report", "title": "Breach Notification Records", "description": "Copies of supervisory authority notifications and data subject communications for actual breaches"}
        ]
    },
    "nist_csf": {
        "identify": [
            {"type": "register", "title": "IT Asset Inventory", "description": "Complete inventory of hardware, software, data, and personnel assets per ID.AM"},
            {"type": "report", "title": "Business Impact Analysis", "description": "BIA documenting critical business processes and recovery priorities per ID.BE"},
            {"type": "policy", "title": "Cybersecurity Policy", "description": "Overarching policy aligned to NIST CSF functions with roles and responsibilities per ID.GV"}
        ],
        "protect": [
            {"type": "procedure", "title": "Identity and Access Management Procedures", "description": "IAM processes covering provisioning, authentication, and authorization per PR.AC"},
            {"type": "training", "title": "Security Awareness Training Records", "description": "Training completion records and materials for workforce security awareness per PR.AT"},
            {"type": "configuration", "title": "Security Baseline Configurations", "description": "Hardening standards and compliance reports for systems per PR.IP"}
        ],
        "detect": [
            {"type": "procedure", "title": "Security Monitoring Procedures", "description": "Processes for continuous monitoring, event correlation, and alert handling per DE.CM"},
            {"type": "log", "title": "SIEM Alert Samples", "description": "Sample security event logs and alerts demonstrating detection capability per DE.AE"},
            {"type": "report", "title": "Detection Capability Assessment", "description": "Assessment of detection coverage across MITRE ATT&CK or similar framework"}
        ]
    },
    "sox": {
        "financial_controls": [
            {"type": "matrix", "title": "Risk Control Matrix (RCM)", "description": "Mapping of financial statement risks to controls with control owners and testing procedures"},
            {"type": "procedure", "title": "Financial Close Procedures", "description": "Month/quarter/year-end close procedures with reconciliation and review steps"},
            {"type": "attestation", "title": "Management Review Sign-offs", "description": "Evidence of management review and approval of financial statements and reconciliations"}
        ],
        "it_general_controls": [
            {"type": "procedure", "title": "Change Management Procedure", "description": "IT change management process with segregation between development and production"},
            {"type": "log", "title": "Access Provisioning/Deprovisioning Logs", "description": "Records of access grants and revocations for financially significant applications"},
            {"type": "report", "title": "Privileged Access Review", "description": "Quarterly review of privileged access to financially significant systems"}
        ]
    },
    "basel": {
        "capital_requirements": [
            {"type": "report", "title": "Capital Adequacy Report", "description": "Regulatory capital calculations showing CET1, Tier 1, and Total Capital ratios"},
            {"type": "procedure", "title": "RWA Calculation Methodology", "description": "Documented methodology for calculating Risk-Weighted Assets across credit, market, operational risk"},
            {"type": "register", "title": "Capital Instruments Register", "description": "Inventory of regulatory capital instruments with terms and eligibility analysis"}
        ],
        "risk_management": [
            {"type": "policy", "title": "Enterprise Risk Management Framework", "description": "Board-approved ERM framework covering risk appetite, governance, and three lines of defense"},
            {"type": "report", "title": "ICAAP Documentation", "description": "Internal Capital Adequacy Assessment Process documentation and results"},
            {"type": "register", "title": "Material Risk Register", "description": "Inventory of material risks with quantification, monitoring, and mitigation strategies"}
        ]
    }
}


def get_framework_category(framework_name: str, text_sample: str) -> str:
    """Determine the framework category for tailored evidence generation."""
    name_lower = framework_name.lower()
    text_lower = text_sample[:5000].lower()
    
    if "27001" in name_lower or "27002" in name_lower or "isms" in text_lower:
        return "iso_27001"
    elif "pci" in name_lower or "payment card" in text_lower or "cardholder" in text_lower:
        return "pci_dss"
    elif "gdpr" in name_lower or "general data protection" in text_lower:
        return "gdpr"
    elif "nist" in name_lower or "cybersecurity framework" in text_lower:
        return "nist_csf"
    elif "sox" in name_lower or "sarbanes" in text_lower:
        return "sox"
    elif "basel" in name_lower or "capital adequacy" in text_lower:
        return "basel"
    elif "hipaa" in name_lower or "health insurance portability" in text_lower:
        return "hipaa"
    else:
        return "general"


def extract_document_structure(text: str, framework_name: str) -> dict:
    """First pass: Extract the document's structure and classify the framework type using GRC SME expertise."""
    if not check_ai_available():
        return {"sections": [], "total_expected_controls": 0}
    
    client = get_openai_client()
    
    sample_text = text[:25000] if len(text) > 25000 else text
    
    try:
        import time
        start_time = time.time()
        print(f"[PARSE] Extracting document structure with OpenAI (gpt-4o)...", flush=True)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": GRC_SME_SYSTEM_PROMPT},
                {"role": "user", "content": f"""As a Senior GRC SME, analyze this regulatory framework document "{framework_name}" and provide comprehensive structural analysis.

Document excerpt:
---
{sample_text}
---

Perform deep analysis and return JSON with:

1. "framework_classification": {{
     "type": "certification|compliance_regulation|best_practice|industry_standard",
     "description": "Why this classification applies",
     "regulatory_authority": "Issuing body (e.g., ISO, NIST, PCI SSC, regulators)",
     "jurisdiction": "Global/Regional/National applicability",
     "certification_body_required": true/false
   }}

2. "document_structure": {{
     "has_normative_clauses": true/false,
     "has_informative_annexes": true/false,
     "main_requirement_sections": ["List of sections containing actual requirements"],
     "skip_sections": ["List of sections to SKIP - intro, foreword, scope, TOC, bibliography"],
     "control_hierarchy_depth": 1-5 (how deep the clause numbering goes)
   }}

3. "sections": Array of ALL major sections/chapters with their exact numbering

4. "control_patterns": {{
     "primary_pattern": "Main numbering format (e.g., 'N.N.N', 'A.N.N.N', 'Requirement N.N.N')",
     "secondary_patterns": ["Additional patterns used"],
     "examples": ["5.1.1", "A.5.1.1", "Requirement 1.1.1a"]
   }}

5. "requirement_indicators": {{
     "mandatory_keywords": ["shall", "must", "required", etc.],
     "advisory_keywords": ["should", "may", "recommended", etc.]
   }}

6. "total_expected_controls": Estimated total number of extractable requirements
7. "framework_type": "ISO|NIST|PCI|Banking|Privacy|Financial|Industry"
8. "evidence_focus_areas": ["Key areas where evidence will be most important"]"""}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=4096,
            temperature=0
        )
        
        elapsed = time.time() - start_time
        print(f"[PARSE] Document structure extracted in {elapsed:.1f}s", flush=True)
        
        result = json.loads(response.choices[0].message.content or "{}")
        expected = result.get("total_expected_controls", 0)
        print(f"[PARSE] Expected controls from document structure: {expected}", flush=True)
        return result
    except Exception as e:
        print(f"[PARSE] Error extracting document structure: {e}", flush=True)
        return {"sections": [], "total_expected_controls": 0}


def extract_controls_lightweight(text: str, framework_name: str, chunk_number: int = 1, total_chunks: int = 1, doc_structure: dict = None) -> List[dict]:
    """Extract controls with minimal fields for maximum quantity.
    
    This is Pass 1 of the two-pass extraction approach. It uses a lightweight
    output format to maximize the number of controls extracted per API call.
    """
    client = get_openai_client()
    
    chunk_context = ""
    if total_chunks > 1:
        chunk_context = f"\n\nThis is chunk {chunk_number} of {total_chunks}. Extract ALL controls from this chunk."
    
    structure_hints = ""
    if doc_structure:
        control_patterns = doc_structure.get("control_patterns", {})
        examples = control_patterns.get("examples", [])[:5]
        if examples:
            structure_hints = f"\nNumbering patterns in this document: {', '.join(examples)}"
    
    prompt = f"""EXTRACTION QUANTITY GOAL: You MUST extract at least 15-25 controls from this chunk.
If you find fewer than 10, you are MISSING requirements - re-read more carefully.

EVERY "shall", "must", "should", "require", "ensure", "maintain", "implement", "establish", "document", "review", "test", "assess", "monitor", "report" statement is a SEPARATE control.

SPLITTING EXAMPLES (CRITICAL - follow these):
- "shall implement, maintain, and review" = 3 SEPARATE controls (one for implement, one for maintain, one for review)
- "must (a) document, (b) test, (c) update" = 3 SEPARATE controls
- "requirements 5.1, 5.2, 5.3" = 3 SEPARATE controls (one for each number)
- Any bullet points or numbered lists = SEPARATE control for EACH item
- "shall ensure X and Y" = 2 SEPARATE controls
- "(i) first, (ii) second, (iii) third" = 3 SEPARATE controls

DOCUMENT: "{framework_name}"{chunk_context}{structure_hints}

SKIP: Foreword, Introduction, Table of Contents, Definitions, Bibliography, page headers/footers.
EXTRACT: Every SHALL, MUST, SHOULD, REQUIRE statement as a SEPARATE control.

OUTPUT FORMAT - Return JSON with "controls" array. Each control needs ONLY these 7 fields:
{{
  "original_reference": "exact clause number (e.g., '5.1.a', 'A.9.2.3', 'Principle 4.2')",
  "title": "brief descriptive title, max 100 chars",
  "full_text": "verbatim requirement text from document",
  "is_mandatory": true for shall/must/required, false for should/may,
  "domain": "one word: Governance|Security|Risk|Access|Operations|Data|Compliance|Vendor|Network|Incident|BCP|HR|Physical|Crypto|Asset",
  "category": "specific sub-category",
  "priority": "high|medium|low"
}}

DOCUMENT TEXT:
---
{text}
---

Remember: Target 15-25+ controls per chunk. Split compound requirements. Every bullet point and sub-item with an obligation is a separate control."""

    try:
        import time
        start_time = time.time()
        print(f"[PARSE] Lightweight extraction chunk {chunk_number}/{total_chunks}...", flush=True)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a compliance expert extracting regulatory requirements. Extract the MAXIMUM number of individual controls by splitting compound requirements. Each shall/must/should statement is a separate control."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_tokens=16384,
            temperature=0
        )
        
        elapsed = time.time() - start_time
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        controls = result.get("controls", [])
        print(f"[PARSE] Chunk {chunk_number}: extracted {len(controls)} controls in {elapsed:.1f}s", flush=True)
        
        return controls
    
    except json.JSONDecodeError as e:
        print(f"[PARSE] JSON decode error in chunk {chunk_number}: {e}", flush=True)
        return []
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        print(f"[PARSE] Error in chunk {chunk_number}: {error_msg}", flush=True)
        return []


def enhance_extracted_controls(controls: List[dict], framework_name: str) -> List[dict]:
    """Enhance lightweight controls with additional details.
    
    This is Pass 2 of the two-pass extraction approach. It adds:
    - Detailed description
    - Evidence requirements (2-3 per control)
    - Control type, implementation frequency
    - AI confidence and notes
    
    Processes controls in batches of 10 for efficiency.
    """
    if not check_ai_available():
        for control in controls:
            control.setdefault("description", control.get("full_text", "")[:500])
            control.setdefault("evidence_requirements", [])
            control.setdefault("control_type", "preventive")
            control.setdefault("implementation_frequency", "continuous")
            control.setdefault("ai_confidence", 0.8)
            control.setdefault("ai_notes", None)
            control.setdefault("parent_reference", None)
            control.setdefault("hierarchy_level", 1)
            control["evidence_types"] = infer_evidence_types(control)
        return controls
    
    client = get_openai_client()
    
    enhanced_controls = []
    batch_size = 10
    total_batches = (len(controls) + batch_size - 1) // batch_size
    
    for batch_num, i in enumerate(range(0, len(controls), batch_size), start=1):
        batch = controls[i:i + batch_size]
        print(f"[PARSE] Enhancing batch {batch_num}/{total_batches} ({len(batch)} controls)...", flush=True)
        
        controls_json = json.dumps(batch, indent=2)
        
        prompt = f"""For framework "{framework_name}", enhance these controls with additional audit-ready details.

INPUT CONTROLS:
{controls_json}

For EACH control, add these fields while keeping all existing fields:
- description: Plain English explanation (1-2 sentences)
- parent_reference: Parent clause number if this is a sub-item, null otherwise
- hierarchy_level: 1-5 (1=top clause, 2=sub, 3=sub-sub, etc.)
- control_type: "preventive" | "detective" | "corrective" | "directive"
- implementation_frequency: "one-time" | "daily" | "weekly" | "monthly" | "quarterly" | "annual" | "continuous" | "event-driven"
- evidence_requirements: Array of 2-3 evidence items, each with:
  {{
    "type": "policy|procedure|configuration|log|report|contract|attestation|register",
    "title": "specific evidence name",
    "description": "what auditor looks for",
    "is_required": true/false
  }}
- ai_confidence: 0.0-1.0 (1.0 for clear SHALL, lower for implicit)
- ai_notes: Any relevant notes or null

Return JSON with "controls" array containing the enhanced controls with ALL fields."""

        try:
            import time
            start_time = time.time()
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a GRC expert adding audit-ready evidence requirements to compliance controls."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_tokens=16384,
                temperature=0
            )
            
            elapsed = time.time() - start_time
            result_text = response.choices[0].message.content or "{}"
            result = json.loads(result_text)
            enhanced_batch = result.get("controls", batch)
            
            for original, enhanced in zip(batch, enhanced_batch):
                merged = {**original, **enhanced}
                merged.setdefault("description", merged.get("full_text", "")[:500])
                merged.setdefault("evidence_requirements", [])
                merged.setdefault("control_type", "preventive")
                merged.setdefault("implementation_frequency", "continuous")
                merged.setdefault("ai_confidence", 0.8)
                merged.setdefault("ai_notes", None)
                merged.setdefault("parent_reference", None)
                merged.setdefault("hierarchy_level", 1)
                if not merged.get("evidence_types"):
                    merged["evidence_types"] = infer_evidence_types(merged)
                
                seen_types = set()
                unique_evidence = []
                for ev in merged.get("evidence_requirements", []):
                    ev_type = ev.get("type", "document") if isinstance(ev, dict) else "document"
                    if ev_type not in seen_types:
                        seen_types.add(ev_type)
                        unique_evidence.append(ev)
                merged["evidence_requirements"] = unique_evidence
                
                enhanced_controls.append(merged)
            
            print(f"[PARSE] Batch {batch_num} enhanced in {elapsed:.1f}s", flush=True)
            
        except Exception as e:
            print(f"[PARSE] Enhancement error in batch {batch_num}: {e}. Using defaults.", flush=True)
            for control in batch:
                control.setdefault("description", control.get("full_text", "")[:500])
                control.setdefault("evidence_requirements", [])
                control.setdefault("control_type", "preventive")
                control.setdefault("implementation_frequency", "continuous")
                control.setdefault("ai_confidence", 0.8)
                control.setdefault("ai_notes", None)
                control.setdefault("parent_reference", None)
                control.setdefault("hierarchy_level", 1)
                control["evidence_types"] = infer_evidence_types(control)
                enhanced_controls.append(control)
    
    return enhanced_controls


def parse_with_openai(text: str, framework_name: str, chunk_number: int = 1, total_chunks: int = 1, doc_structure: dict = None) -> List[dict]:
    client = get_openai_client()
    
    chunk_context = ""
    if total_chunks > 1:
        chunk_context = f"""

=== DOCUMENT CHUNK CONTEXT ===
- This is chunk {chunk_number} of {total_chunks} from a large document
- You MUST extract EVERY control from this section - do not skip any
- If a control spans chunk boundaries, extract what you see with a note
- Previous/next chunks will capture overlapping content
- CRITICAL: Do not assume content from other chunks - focus on THIS text only"""
    
    structure_context = ""
    if doc_structure:
        fw_classification = doc_structure.get("framework_classification", {})
        doc_struct = doc_structure.get("document_structure", {})
        control_patterns = doc_structure.get("control_patterns", {})
        skip_sections = doc_struct.get("skip_sections", [])
        
        structure_context = f"""

=== FRAMEWORK ANALYSIS (from structural pre-scan) ===

FRAMEWORK CLASSIFICATION:
- Type: {fw_classification.get('type', 'Unknown')}
- Regulatory Authority: {fw_classification.get('regulatory_authority', 'Unknown')}
- Jurisdiction: {fw_classification.get('jurisdiction', 'Unknown')}
- Requires Certification: {fw_classification.get('certification_body_required', False)}

DOCUMENT STRUCTURE:
- Control Hierarchy Depth: {doc_struct.get('control_hierarchy_depth', 3)} levels deep
- Has Normative Clauses: {doc_struct.get('has_normative_clauses', True)}
- Has Informative Annexes: {doc_struct.get('has_informative_annexes', False)}

NUMBERING PATTERNS TO LOOK FOR:
- Primary Pattern: {control_patterns.get('primary_pattern', 'N.N.N')}
- Examples: {', '.join(control_patterns.get('examples', [])[:5])}

SECTIONS TO SKIP (non-requirements):
{chr(10).join(['- ' + s for s in skip_sections[:10]]) if skip_sections else '- None identified'}

SECTIONS WITH ACTUAL REQUIREMENTS:
{chr(10).join(['- ' + s for s in doc_struct.get('main_requirement_sections', [])[:10]])}"""
    
    prompt = f"""As a SENIOR GRC SUBJECT MATTER EXPERT, perform EXHAUSTIVE and INTELLIGENT extraction of ALL compliance requirements from this regulatory framework document.

=== DOCUMENT: "{framework_name}" ==={chunk_context}{structure_context}

=== CRITICAL EXTRACTION RULE ===
CRITICAL: Extract EVERY SINGLE 'shall', 'must', 'should' statement as a SEPARATE control. Do NOT consolidate multiple requirements into one control.

EXAMPLE OF SPLITTING:
If text says: "The organization shall (a) maintain logs, (b) review logs monthly, and (c) retain logs for 1 year"
→ Create 3 SEPARATE controls:
  1. Control for "(a) maintain logs"
  2. Control for "(b) review logs monthly" 
  3. Control for "(c) retain logs for 1 year"

TARGET: For a 30,000 character chunk of regulatory text, you should typically find 25-40 individual controls. If you find fewer than 15, you are likely consolidating too much - re-read the text more carefully.

=== YOUR EXPERT ANALYSIS PROCESS ===

STEP 1: SKIP NON-REQUIREMENT CONTENT
Intelligently skip these sections (they contain NO auditable requirements):
- Foreword, Introduction, Scope (unless scope contains SHALL statements)
- Normative References, Terms and Definitions (unless defining obligations)
- Table of Contents, Index, Bibliography, Acknowledgments
- Informative Annexes (background/guidance only, not normative)
- Historical context, rationale explanations, "Note:" sections (unless they contain SHALL)
- Page headers, footers, version information

STEP 2: IDENTIFY REQUIREMENT LANGUAGE
Extract ONLY when you find:
- MANDATORY: "shall", "must", "is required", "are required to", "needs to", "is mandatory"
- CONDITIONAL MANDATORY: "shall, where applicable", "must, unless documented"
- ADVISORY: "should", "may", "is recommended", "is encouraged", "it is advisable"

STEP 3: PRESERVE EXACT HIERARCHICAL STRUCTURE
Capture the COMPLETE clause hierarchy as it appears in the document:

For ISO-style documents:
- "4.1" (Clause 4.1)
- "4.1.1" (Sub-clause)
- "4.1.1.1" (Sub-sub-clause)
- "A.5.1.1" (Annex A control)
- "A.5.1.1.a" or "A.5.1.1 a)" (Lettered item)

For PCI DSS-style documents:
- "Requirement 1.1" (Main requirement)
- "1.1.1" (Sub-requirement)
- "1.1.1.a" (Testing procedure)
- "1.1.1.a.i" (Detailed procedure)

For NIST-style documents:
- "ID.AM-1" (Subcategory)
- "PR.AC-1" (Subcategory)
- Preserve the exact category codes

For Banking/Basel-style documents:
- "Principle 1" (High-level principle)
- "Principle 1.1" or "Principle 1, paragraph 2"
- "Article 3.2.1" (Detailed article)

For GDPR/Regulation-style:
- "Article 5(1)(a)" (Exact legal citation)
- "Article 32(1)" (Specific paragraph)

CRITICAL: Never consolidate, simplify, or modify clause numbers. Copy EXACTLY as written.

STEP 4: GRANULAR EXTRACTION
- Extract EACH "shall/must/should" statement as a SEPARATE control
- If one clause has 5 requirements, create 5 controls with references like "4.1.a", "4.1.b", etc.
- Sub-clauses within a bullet point still need individual extraction
- Testing procedures (in PCI DSS, for example) are separate from requirements

STEP 5: SPLIT COMPOUND REQUIREMENTS
- If a single clause contains multiple obligations (a, b, c or i, ii, iii), create SEPARATE controls for each
- Example: "shall implement and maintain" = 2 controls (implement + maintain)
- Example: "shall establish, document, and review" = 3 controls
- Each control should have ONE clear testable requirement
- Look for conjunctions like "and", "or", commas separating distinct obligations

STEP 6: EXTRACT SUB-POINTS
- Bullet points under a requirement are separate controls
- Lettered items (a), (b), (c) are separate controls
- Numbered sub-items 1), 2), 3) are separate controls
- Roman numerals (i), (ii), (iii) are separate controls
- Even if grouped under one parent, each sub-point with its own obligation is a distinct control

=== OUTPUT FORMAT FOR EACH CONTROL ===

{{
  "original_reference": "EXACT clause number as it appears (e.g., '4.1.2', 'A.5.1.1', 'Requirement 1.1.1.a', 'Article 32(1)(a)')",
  "parent_reference": "Parent clause if this is a sub-item (e.g., '4.1' is parent of '4.1.2')",
  "hierarchy_level": 1-5 (1=top clause, 2=sub-clause, 3=sub-sub-clause, etc.),
  "title": "Clear, descriptive title (max 200 chars) - not just the clause number",
  "description": "Plain English explanation of what this control requires and why",
  "full_text": "COMPLETE VERBATIM text of the requirement - copy exactly as written including any notes",
  "domain": "One of: Governance|Risk Management|Security|Access Control|Incident Management|Business Continuity|Data Protection|Compliance|Operations|Third Party|Capital & Liquidity|Credit Risk|Human Resources|Physical Security|Network Security|Application Security|Asset Management|Cryptography|Communications Security|Supplier Management",
  "category": "Specific sub-category (e.g., 'Access Control Policy', 'User Access Management', 'Privileged Access')",
  "is_mandatory": true for shall/must/required, false for should/may/recommended,
  "priority": "critical|high|medium|low based on risk and regulatory importance",
  "control_type": "preventive|detective|corrective|directive",
  "implementation_frequency": "one-time|daily|weekly|monthly|quarterly|annual|continuous|event-driven",
  "evidence_requirements": [
    {{
      "type": "policy|procedure|configuration|log|report|contract|attestation|register|matrix|plan|screenshot|training|assessment|interview|observation",
      "title": "Specific evidence artifact name (e.g., 'Information Security Policy v2.0', 'User Access Review Q4 2024')",
      "description": "Detailed description of what the auditor will look for in this evidence - reference the SPECIFIC control requirement",
      "artifact_examples": ["Practical examples of what to provide"],
      "review_frequency": "How often this evidence should be produced/updated",
      "is_required": true/false (true for primary evidence, false for supporting)
    }}
  ],
  "testing_procedure": "How an auditor would test/verify this control",
  "ai_confidence": 0.0-1.0 (1.0 for clear SHALL statements, lower for implicit requirements),
  "ai_notes": "Any extraction notes (e.g., 'Requirement spans two paragraphs', 'Implicit from context')"
}}

=== FRAMEWORK-SPECIFIC EVIDENCE GUIDANCE ===

Provide PRACTICAL, AUDITABLE evidence that auditors actually request. Examples:

FOR ISO 27001 A.9 ACCESS CONTROL:
- NOT: "Access Control Document" (too vague)
- YES: "Access Control Policy (version-controlled, management-approved) covering: logical access principles, user lifecycle management, privileged access, access review requirements per A.9.1"
- YES: "Quarterly User Access Review spreadsheet showing: review date, reviewer name, systems reviewed, exceptions found, remediation actions with target dates"
- YES: "Active Directory group membership export showing RBAC implementation"

FOR PCI DSS REQUIREMENT 1 (FIREWALLS):
- NOT: "Firewall documentation" (too vague)  
- YES: "Firewall ruleset export from [vendor] showing: deny-all default, explicit permit rules with business justification, no any-any rules per Req 1.1.1"
- YES: "Semi-annual firewall rule review minutes with: rules reviewed, justifications validated, rules removed as obsolete per Req 1.1.7"
- YES: "Network diagram (Visio/Lucidchart) showing CDE boundaries, firewall placement, and all connections per Req 1.1.2"

FOR GDPR ARTICLE 30 (RECORDS OF PROCESSING):
- NOT: "Processing register" (too vague)
- YES: "Record of Processing Activities (ROPA) spreadsheet containing: processing purposes, lawful basis for each, data categories, retention periods, recipients, transfer safeguards per Art 30(1)"
- YES: "Data Protection Impact Assessment (DPIA) for [system name] with: risk assessment, necessity evaluation, safeguards implemented per Art 35"

FOR SOX IT GENERAL CONTROLS:
- NOT: "Change management evidence" (too vague)
- YES: "Change Advisory Board (CAB) meeting minutes showing: change requests reviewed, approvals documented, segregation between requester and approver"
- YES: "ServiceNow change ticket samples (10 samples across quarter) showing: request, testing, approval, implementation, post-implementation review"

=== DOCUMENT TEXT TO ANALYZE ===
---
{text}
---

=== MANDATORY OUTPUT FIELDS ===

Every control in your output MUST include ALL of these fields (no exceptions):
- original_reference: string (REQUIRED - exact clause number)
- parent_reference: string or null (REQUIRED - parent clause number if sub-item)  
- hierarchy_level: integer 1-5 (REQUIRED - 1=top, 2=sub, 3=sub-sub, etc.)
- title: string (REQUIRED)
- description: string (REQUIRED)
- full_text: string (REQUIRED - verbatim text)
- domain: string (REQUIRED)
- category: string (REQUIRED)
- is_mandatory: boolean (REQUIRED)
- priority: string (REQUIRED - critical/high/medium/low)
- control_type: string (REQUIRED - preventive/detective/corrective/directive)
- implementation_frequency: string (REQUIRED - one-time/daily/weekly/monthly/quarterly/annual/continuous/event-driven)
- evidence_requirements: array (REQUIRED - at least 1 evidence item with type, title, description, artifact_examples array, review_frequency, is_required)
- testing_procedure: string (REQUIRED - how auditor would verify)
- ai_confidence: float (REQUIRED - 0.0-1.0)
- ai_notes: string or null (optional notes)

=== FINAL EXTRACTION CHECKLIST ===

Before completing, verify you have:
[ ] Skipped introductory/non-normative content appropriately
[ ] Extracted EVERY SHALL/MUST/SHOULD statement as a SEPARATE control (NOT consolidated)
[ ] Split compound requirements - each (a), (b), (c) item is its own control
[ ] Split verb phrases - "shall implement and maintain" = 2 controls
[ ] Preserved EXACT clause numbering including all hierarchy levels (a, b, c, i, ii, iii)
[ ] Set parent_reference and hierarchy_level for EVERY control (null parent for top-level)
[ ] Provided SPECIFIC, PRACTICAL evidence with artifact_examples array for each control
[ ] Included testing_procedure for EVERY control (how auditor verifies compliance)
[ ] Set control_type and implementation_frequency for EVERY control
[ ] Set appropriate ai_confidence levels (1.0 for explicit SHALL, lower for implicit)

MINIMUM EXTRACTION TARGET: For a 30,000 character chunk, you should typically find 25-40 individual controls. If you find fewer than 15, re-read the text more carefully - you are likely consolidating multiple requirements into single controls.

For a typical regulatory document, expect to extract 50-500+ controls. 
If you extract fewer than 20, you may have missed requirements - review the text again.

CRITICAL: Return a JSON object with a "controls" array. Each control MUST have ALL mandatory fields listed above.

Example control structure:
{{
  "original_reference": "A.5.1.1",
  "parent_reference": "A.5.1",
  "hierarchy_level": 3,
  "title": "Information Security Policies",
  "description": "A set of policies for information security shall be defined, approved by management, published and communicated to employees and relevant external parties.",
  "full_text": "A.5.1.1 Policies for information security: A set of policies for information security shall be defined, approved by management, published and communicated to employees and relevant external parties.",
  "domain": "Governance",
  "category": "Information Security Policies",
  "is_mandatory": true,
  "priority": "high",
  "control_type": "directive",
  "implementation_frequency": "annual",
  "evidence_requirements": [
    {{
      "type": "policy",
      "title": "Information Security Policy Document",
      "description": "Version-controlled policy document with management approval signature, covering security principles, roles, and responsibilities",
      "artifact_examples": ["Information Security Policy v2.0.pdf", "ISMS Policy Manual"],
      "review_frequency": "Annual",
      "is_required": true
    }},
    {{
      "type": "attestation",
      "title": "Policy Communication Records",
      "description": "Evidence of policy distribution to employees - email notifications, intranet announcements, training acknowledgments",
      "artifact_examples": ["Employee acknowledgment forms", "Email distribution logs"],
      "review_frequency": "Upon update",
      "is_required": true
    }}
  ],
  "testing_procedure": "1. Obtain current information security policy. 2. Verify management approval (signature/date). 3. Sample 10 employees and verify receipt/acknowledgment. 4. Confirm policy published on intranet.",
  "ai_confidence": 1.0,
  "ai_notes": null
}}"""

    try:
        import time
        start_time = time.time()
        print(f"[PARSE] Calling OpenAI API (model: gpt-4o)...", flush=True)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": GRC_SME_SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_tokens=16384,
            temperature=0
        )
        
        elapsed = time.time() - start_time
        print(f"[PARSE] OpenAI API response received in {elapsed:.1f}s", flush=True)
        
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        controls = result.get("controls", [])
        print(f"[PARSE] Parsed {len(controls)} controls from API response", flush=True)
        
        for control in controls:
            if not control.get("evidence_types") or len(control.get("evidence_types", [])) == 0:
                control["evidence_types"] = infer_evidence_types(control)
        
        return controls
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse OpenAI response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"OpenAI API error: {error_msg}"
        )


def parse_document_with_chunking(text: str, framework_name: str) -> tuple:
    """Parse a document using a three-pass approach for comprehensive extraction.
    
    Pass 1: Extract document structure to understand numbering patterns
    Pass 2: Lightweight extraction of ALL controls with minimal fields (maximizes quantity)
    Pass 3: Enhance unique controls with detailed evidence requirements
    
    Returns:
        tuple: (enhanced_controls, doc_structure)
    """
    print(f"[PARSE] Starting document parsing for: {framework_name}", flush=True)
    print(f"[PARSE] Document text length: {len(text):,} characters", flush=True)
    
    print(f"[PARSE] Pass 1: Extracting document structure...", flush=True)
    doc_structure = extract_document_structure(text, framework_name)
    expected_controls = doc_structure.get('total_expected_controls', 50)
    print(f"[PARSE] Document structure extracted. Expected controls: {expected_controls}", flush=True)
    
    chunks = chunk_text(text, chunk_size=20000, overlap=2500)
    print(f"[PARSE] Document split into {len(chunks)} chunks for processing", flush=True)
    
    all_controls = []
    
    print(f"[PARSE] Pass 2: Lightweight extraction (maximizing control quantity)...", flush=True)
    for idx, chunk in enumerate(chunks, start=1):
        print(f"[PARSE] Extracting chunk {idx}/{len(chunks)} ({len(chunk):,} chars)...", flush=True)
        chunk_controls = extract_controls_lightweight(
            chunk, 
            framework_name, 
            chunk_number=idx, 
            total_chunks=len(chunks),
            doc_structure=doc_structure
        )
        all_controls.extend(chunk_controls)
        
        if len(chunk_controls) < 10 and len(chunk) > 5000:
            print(f"[PARSE] WARNING: Only {len(chunk_controls)} controls from chunk {idx}. May need review.", flush=True)
    
    print(f"[PARSE] Lightweight extraction complete. Total raw controls: {len(all_controls)}", flush=True)
    
    print(f"[PARSE] Deduplicating controls...", flush=True)
    unique_controls = deduplicate_controls(all_controls)
    print(f"[PARSE] Deduplication complete. Unique controls: {len(unique_controls)}", flush=True)
    
    if expected_controls > 0 and len(unique_controls) < expected_controls * 0.3:
        print(f"[PARSE] WARNING: Extracted {len(unique_controls)} controls but expected ~{expected_controls}. Review document for missed requirements.", flush=True)
    
    print(f"[PARSE] Pass 3: Enhancing controls with evidence requirements...", flush=True)
    enhanced_controls = enhance_extracted_controls(unique_controls, framework_name)
    print(f"[PARSE] Enhancement complete. Final control count: {len(enhanced_controls)}", flush=True)
    
    return enhanced_controls, doc_structure


def update_parsing_heartbeat(db: Session, framework_id: int, stage: str):
    """Update the framework's updated_at timestamp as a heartbeat during parsing.
    
    This allows staleness detection to accurately determine if parsing is active.
    """
    try:
        fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
        if fw:
            fw.updated_at = datetime.utcnow()
            db.commit()
            print(f"[PARSE] Heartbeat: {stage}", flush=True)
    except Exception as e:
        print(f"[PARSE] Heartbeat update failed: {e}", flush=True)


def _run_background_parsing_body(db, framework_id: int, file_path: str, file_type: str, framework_name: str, tenant_slug: str):
    """Body of the framework-parse job. Takes an open tenant-scoped session."""
    print(f"[PARSE] Background parsing started for framework ID: {framework_id}", flush=True)
    print(f"[PARSE] Framework name: {framework_name} | File: {file_path} ({file_type}) | Tenant: {tenant_slug}", flush=True)

    try:
        extracted_text = ""
        if file_type == "pdf":
            print(f"[PARSE] Extracting text from PDF...", flush=True)
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            print(f"[PARSE] PDF has {len(reader.pages)} pages", flush=True)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted_text = "\n\n".join(text_parts)
            print(f"[PARSE] Text extraction complete. Total characters: {len(extracted_text):,}", flush=True)
            update_parsing_heartbeat(db, framework_id, "PDF text extracted")
        elif file_type == "docx":
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            extracted_text = "\n".join(text_parts)
            update_parsing_heartbeat(db, framework_id, "DOCX text extracted")
        else:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "failed"
                fw.parse_error = f"Unsupported file type: {file_type}"
                db.commit()
            return
        
        if not extracted_text.strip():
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "failed"
                fw.parse_error = "No text could be extracted from the document"
                db.commit()
            return
        
        parsed_controls_data, doc_structure = parse_document_with_chunking(extracted_text, framework_name)
        update_parsing_heartbeat(db, framework_id, f"AI parsing complete - {len(parsed_controls_data) if parsed_controls_data else 0} controls")
        
        if not parsed_controls_data:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "parsed"
                fw.parsed_at = datetime.utcnow()
                fw.parse_error = "No controls found in document"
                db.commit()
            return
        
        existing_controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework_id
        ).all()
        existing_control_ids = [c.id for c in existing_controls]
        
        if existing_control_ids:
            assessment_item_ids = db.query(AssessmentItem.id).filter(
                AssessmentItem.parsed_control_id.in_(existing_control_ids)
            ).all()
            ai_ids = [a.id for a in assessment_item_ids]
            
            if ai_ids:
                db.query(AssessmentRemediation).filter(
                    AssessmentRemediation.assessment_item_id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
                
                db.query(AssessmentEvidence).filter(
                    AssessmentEvidence.assessment_item_id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
                
                db.query(AssessmentItem).filter(
                    AssessmentItem.id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
            
            db.query(FrameworkControlAlignment).filter(
                FrameworkControlAlignment.parsed_control_id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
            
            db.query(ControlEvidenceMapping).filter(
                ControlEvidenceMapping.parsed_control_id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
            
            db.query(ParsedFrameworkControl).filter(
                ParsedFrameworkControl.id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
        
        db.flush()
        
        for idx, control_data in enumerate(parsed_controls_data, start=1):
            raw_reference = control_data.get("original_reference", "")
            cleaned_reference = clean_section_reference(raw_reference) if raw_reference else None
            
            control_id = cleaned_reference if cleaned_reference else f"CTRL-{idx:03d}"
            
            parent_ref = control_data.get("parent_reference", "")
            cleaned_parent_ref = clean_section_reference(parent_ref) if parent_ref else None
            
            hierarchy_level = control_data.get("hierarchy_level", 1)
            if not isinstance(hierarchy_level, int) or hierarchy_level < 1:
                hierarchy_level = 1
            elif hierarchy_level > 5:
                hierarchy_level = 5
            
            testing_procedure = control_data.get("testing_procedure", "")
            ai_notes = control_data.get("ai_notes", "")
            
            full_ai_notes = ai_notes
            if testing_procedure:
                full_ai_notes = f"Testing: {testing_procedure}" + (f"\n{ai_notes}" if ai_notes else "")
            
            parsed_control = ParsedFrameworkControl(
                uploaded_framework_id=framework_id,
                control_id=control_id,
                original_reference=cleaned_reference,
                title=control_data.get("title", "Untitled Control")[:500],
                description=control_data.get("description"),
                full_text=control_data.get("full_text"),
                domain=control_data.get("domain"),
                category=control_data.get("category"),
                is_mandatory=control_data.get("is_mandatory", True),
                priority=control_data.get("priority", "medium"),
                section_number=cleaned_reference,
                parent_section=cleaned_parent_ref,
                ai_confidence=control_data.get("ai_confidence"),
                ai_notes=full_ai_notes[:1000] if full_ai_notes else None,
                is_verified=False
            )
            db.add(parsed_control)
            db.flush()
            
            evidence_requirements = control_data.get("evidence_requirements", [])
            # Always persist evidence_requirements as JSON so certification router can read them
            if evidence_requirements:
                parsed_control.evidence_requirements = evidence_requirements
                for ev_req in evidence_requirements:
                    if isinstance(ev_req, dict):
                        ev_type = ev_req.get("type", "document")
                        ev_title = ev_req.get("title", "")
                        ev_description = ev_req.get("description", "")
                        ev_is_required = ev_req.get("is_required", True)
                        ev_examples = ev_req.get("artifact_examples", [])
                        ev_frequency = ev_req.get("review_frequency", "")
                        
                        full_description = ev_title
                        if ev_description:
                            full_description = f"{ev_title}: {ev_description}" if ev_title else ev_description
                        if ev_examples and isinstance(ev_examples, list):
                            full_description += f" Examples: {', '.join(ev_examples[:3])}"
                        if ev_frequency:
                            full_description += f" (Frequency: {ev_frequency})"
                        
                        valid_types = [
                            "policy", "procedure", "configuration", "log", "report", 
                            "contract", "attestation", "register", "matrix", "plan", 
                            "screenshot", "training", "assessment", "document",
                            "interview", "observation"
                        ]
                        if ev_type not in valid_types:
                            ev_type = "document"
                        
                        evidence_mapping = ControlEvidenceMapping(
                            parsed_control_id=parsed_control.id,
                            evidence_type=ev_type,
                            evidence_description=full_description[:1000] if full_description else None,
                            is_required=ev_is_required if isinstance(ev_is_required, bool) else True,
                            suggested_by_ai=True
                        )
                        db.add(evidence_mapping)
            else:
                evidence_types = control_data.get("evidence_types", [])
                for evidence_type in evidence_types:
                    if evidence_type in ["policy", "procedure", "configuration", "log", "report", "contract"]:
                        evidence_mapping = ControlEvidenceMapping(
                            parsed_control_id=parsed_control.id,
                            evidence_type=evidence_type,
                            is_required=True,
                            suggested_by_ai=True
                        )
                        db.add(evidence_mapping)
        
        fw_final = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
        if fw_final:
            fw_final.upload_status = "parsed"
            fw_final.parsed_at = datetime.utcnow()
            fw_final.parse_error = None
            
            # CRITICAL: Ensure document_structure always has at least minimal content for phases
            # The uploaded-framework-only architecture requires phases to come from document_structure
            if not doc_structure or not isinstance(doc_structure, dict) or not doc_structure.get("sections"):
                print(f"[PARSE] Enforcing minimal document_structure - AI extraction returned empty or malformed", flush=True)
                doc_structure = {
                    "sections": [
                        {
                            "name": f"{framework_name} - Full Document",
                            "number": "1",
                            "description": "Complete framework document"
                        }
                    ],
                    "total_expected_controls": len(parsed_controls_data) if parsed_controls_data else 0,
                    "framework_type": "imported_framework",
                    "note": "Minimal structure created due to empty AI extraction"
                }
                print(f"[PARSE] Created minimal document_structure with {len(doc_structure.get('sections', []))} section(s)", flush=True)
            
            fw_final.document_structure = doc_structure
        
        db.commit()
        
    except Exception as e:
        error_msg = str(e)
        try:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "failed"
                fw.parse_error = error_msg[:500]
                db.commit()
        except Exception:
            pass
        raise


def run_background_parsing(framework_id: int, file_path: str, file_type: str, framework_name: str, tenant_slug: str):
    """Legacy in-process entry point. Opens its own tenant session and calls
    `_run_background_parsing_body`. Kept for any threaded callers; new
    dispatches go through `grc.tasks.frameworks.parse_framework`."""
    from ....db import open_tenant_session
    db = open_tenant_session(tenant_slug)
    try:
        return _run_background_parsing_body(db, framework_id, file_path, file_type, framework_name, tenant_slug)
    finally:
        db.close()


CLASSIFICATION_SYSTEM_PROMPT = """You are a GRC (Governance, Risk, and Compliance) expert specializing in regulatory frameworks and certification standards. Your task is to analyze a framework document and classify it accurately.

CLASSIFICATION CRITERIA:

1. CERTIFICATION FRAMEWORK:
   - Requires formal third-party audit/assessment leading to a certificate
   - Has specific audit procedures, qualified assessor requirements, and validity periods
   - Organizations receive a formal certificate or attestation upon successful assessment
   - Examples: PCI DSS (QSA assessment, ROC), ISO 27001 (certification body audit), SOC 2 (CPA audit), SWIFT CSP (annual attestation)

2. COMPLIANCE FRAMEWORK:
   - Regulatory/legal requirements or best practice guidelines that require adherence
   - May have audits or assessments but NO formal certificate is issued
   - Compliance is typically demonstrated through self-assessment, regulatory reporting, or inspection
   - Examples: GDPR (law), SAMA CSF (regulatory guideline), NIST CSF (voluntary), COBIT (best practice)

REFERENCE EXAMPLES:
- PCI DSS → CERTIFICATION (requires QSA assessment, ROC, certificate of compliance)
- SWIFT CSCF → CERTIFICATION (requires SWIFT CSP certification, annual attestation)
- ISO 27001 → CERTIFICATION (requires certification body audit, formal certificate)
- SOC 2 → CERTIFICATION (requires CPA audit, attestation report)
- SAMA CSF → COMPLIANCE (Saudi regulatory guideline, no formal certification)
- GDPR → COMPLIANCE (EU law, no certification - compliance is legal requirement)
- NIST CSF → COMPLIANCE (voluntary framework, no formal certification)
- SBP Guidelines → COMPLIANCE (Pakistan central bank regulations)
- SABIC CyberTrust → COMPLIANCE (vendor requirements, no formal certification)
- ARAMCO SACS → COMPLIANCE (third-party requirements, no formal certification)
- COBIT → COMPLIANCE (ISACA best practice framework, no certification required)

Analyze the document text carefully and provide a comprehensive classification with high confidence."""


def classify_framework_with_ai(text: str, framework_name: str) -> dict:
    """Use OpenAI to classify a framework as certification or compliance."""
    client = get_openai_client()
    
    analysis_text = text[:15000]
    
    prompt = f"""Analyze the following framework document and classify it.

FRAMEWORK NAME: {framework_name}

DOCUMENT TEXT (first ~15000 characters):
---
{analysis_text}
---

Based on the document content, provide a comprehensive classification. Return a JSON object with the following structure:

{{
    "classification": "certification" or "compliance",
    "classification_confidence": 0.0-1.0 (how confident you are in the classification),
    "classification_reasoning": "Detailed explanation of why this is classified as certification or compliance",
    
    "framework_purpose": "What this framework aims to achieve",
    "framework_scope": "Who/what this framework applies to (industries, organization types, geographies)",
    "framework_objectives": ["List", "of", "key", "objectives"],
    "target_audience": "Who should implement this framework",
    
    "certification_body": "Organization that issues the certification (null if compliance framework)",
    "certification_validity_period": "How long certification is valid, e.g., '1 year', '3 years' (null if compliance)",
    "certification_levels": ["Tier/level options if applicable, otherwise null"],
    "certification_lifecycle": {{
        "preparation": "Description of preparation phase",
        "assessment": "Description of assessment/audit phase",
        "remediation": "Description of remediation phase if gaps found",
        "certification": "Description of certification issuance",
        "maintenance": "Description of ongoing maintenance requirements"
    }},
    "required_artifacts": ["policies", "procedures", "controls", "records", "evidence types needed"],
    
    "regulatory_authority": "Who enforces this (null if certification framework)",
    "compliance_deadline": "When compliance is required if mentioned (null if not mentioned or certification)",
    "penalty_for_non_compliance": "Consequences of non-compliance (null if certification framework)",
    "adoption_approach": ["Step 1: ...", "Step 2: ...", "Recommended implementation steps"]
}}

IMPORTANT:
- For CERTIFICATION frameworks: Fill in certification_* fields, set regulatory_authority/penalty_for_non_compliance/adoption_approach to null
- For COMPLIANCE frameworks: Fill in regulatory_authority/penalty_for_non_compliance/adoption_approach, set certification_* fields to null
- Be specific and detailed in your analysis
- If information is not found in the document, provide reasonable inferences based on framework type"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": CLASSIFICATION_SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_tokens=4096,
            temperature=0
        )
        
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        return result
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse classification response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Classification failed: {error_msg}"
        )


@router.post("/{framework_id}/classify")
def classify_framework(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Classify a framework as certification or compliance before parsing.
    
    This endpoint analyzes the framework document using AI to determine:
    - Whether it's a certification framework (requires formal audit/certificate)
    - Or a compliance framework (regulatory requirement without formal certification)
    
    The classification results are stored in the framework record and include:
    - Classification type and confidence
    - Framework purpose, scope, objectives, and target audience
    - Type-specific metadata (certification body, validity period, or regulatory authority, penalties)
    """
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    if framework.upload_status == "classifying":
        return {
            "message": "Classification already in progress",
            "framework_id": framework_id,
            "status": "classifying"
        }
    
    if not os.path.exists(framework.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework file not found on disk"
        )
    
    framework.upload_status = "classifying"
    db.commit()
    
    try:
        extracted_text = extract_text_from_file(framework)
        
        if not extracted_text.strip():
            framework.upload_status = "uploaded"
            framework.parse_error = "No text could be extracted from the document"
            db.commit()
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text could be extracted from the document"
            )
        
        classification_result = classify_framework_with_ai(extracted_text, framework.name)
        
        framework.classification = classification_result.get("classification")
        framework.classification_confidence = classification_result.get("classification_confidence")
        framework.classification_reasoning = classification_result.get("classification_reasoning")
        
        framework.framework_purpose = classification_result.get("framework_purpose")
        framework.framework_scope = classification_result.get("framework_scope")
        framework.framework_objectives = classification_result.get("framework_objectives")
        framework.target_audience = classification_result.get("target_audience")
        
        if framework.classification == "certification":
            framework.certification_body = classification_result.get("certification_body")
            framework.certification_validity_period = classification_result.get("certification_validity_period")
            framework.certification_levels = classification_result.get("certification_levels")
            framework.certification_lifecycle = classification_result.get("certification_lifecycle")
            framework.required_artifacts = classification_result.get("required_artifacts")
            framework.regulatory_authority = None
            framework.compliance_deadline = None
            framework.penalty_for_non_compliance = None
            framework.adoption_approach = None
        else:
            framework.regulatory_authority = classification_result.get("regulatory_authority")
            deadline_str = classification_result.get("compliance_deadline")
            if deadline_str and isinstance(deadline_str, str) and deadline_str.lower() not in ["null", "none", ""]:
                try:
                    from dateutil import parser as date_parser
                    framework.compliance_deadline = date_parser.parse(deadline_str)
                except Exception:
                    framework.compliance_deadline = None
            framework.penalty_for_non_compliance = classification_result.get("penalty_for_non_compliance")
            framework.adoption_approach = classification_result.get("adoption_approach")
            framework.certification_body = None
            framework.certification_validity_period = None
            framework.certification_levels = None
            framework.certification_lifecycle = None
            framework.required_artifacts = classification_result.get("required_artifacts")
        
        framework.upload_status = "classified"
        framework.updated_at = datetime.utcnow()
        db.commit()
        
        return {
            "message": "Framework classified successfully",
            "framework_id": framework_id,
            "status": "classified",
            "classification": framework.classification,
            "classification_confidence": framework.classification_confidence,
            "classification_reasoning": framework.classification_reasoning,
            "framework_purpose": framework.framework_purpose,
            "framework_scope": framework.framework_scope,
            "framework_objectives": framework.framework_objectives,
            "target_audience": framework.target_audience,
            "certification_body": framework.certification_body,
            "certification_validity_period": framework.certification_validity_period,
            "certification_levels": framework.certification_levels,
            "certification_lifecycle": framework.certification_lifecycle,
            "required_artifacts": framework.required_artifacts,
            "regulatory_authority": framework.regulatory_authority,
            "compliance_deadline": framework.compliance_deadline.isoformat() if framework.compliance_deadline else None,
            "penalty_for_non_compliance": framework.penalty_for_non_compliance,
            "adoption_approach": framework.adoption_approach
        }
    
    except HTTPException:
        db.rollback()
        try:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw and fw.upload_status == "classifying":
                fw.upload_status = "uploaded"
                db.commit()
        except Exception:
            pass
        raise
    except Exception as e:
        db.rollback()
        error_msg = str(e)
        try:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "uploaded"
                fw.parse_error = f"Classification failed: {error_msg[:500]}"
                db.commit()
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Classification failed: {error_msg}"
        )


@router.post("/{framework_id}/parse")
def parse_framework_document(
    framework_id: int,
    background_tasks: BackgroundTasks,
    http_request: Request,
    classify_first: bool = Query(False, description="Run classification before parsing if not already classified"),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Start parsing a framework document in the background.
    
    Returns immediately with status 'parsing'. The actual parsing runs in
    the background. Poll the framework status to check when parsing completes.
    
    Args:
        classify_first: If True and the framework is not already classified,
                       run classification before parsing. Default is False.
    """
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    if framework.upload_status == "parsing":
        return {
            "message": "Parsing already in progress",
            "framework_id": framework_id,
            "status": "parsing"
        }
    
    if framework.upload_status == "classifying":
        return {
            "message": "Classification in progress. Parsing will start after classification.",
            "framework_id": framework_id,
            "status": "classifying"
        }
    
    file_path = framework.file_path
    file_type = framework.file_type
    framework_name = framework.name
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework file not found on disk"
        )
    
    if classify_first and not framework.classification:
        try:
            framework.upload_status = "classifying"
            db.commit()
            
            extracted_text = extract_text_from_file(framework)
            
            if extracted_text.strip():
                classification_result = classify_framework_with_ai(extracted_text, framework.name)
                
                framework.classification = classification_result.get("classification")
                framework.classification_confidence = classification_result.get("classification_confidence")
                framework.classification_reasoning = classification_result.get("classification_reasoning")
                framework.framework_purpose = classification_result.get("framework_purpose")
                framework.framework_scope = classification_result.get("framework_scope")
                framework.framework_objectives = classification_result.get("framework_objectives")
                framework.target_audience = classification_result.get("target_audience")
                
                if framework.classification == "certification":
                    framework.certification_body = classification_result.get("certification_body")
                    framework.certification_validity_period = classification_result.get("certification_validity_period")
                    framework.certification_levels = classification_result.get("certification_levels")
                    framework.certification_lifecycle = classification_result.get("certification_lifecycle")
                    framework.required_artifacts = classification_result.get("required_artifacts")
                else:
                    framework.regulatory_authority = classification_result.get("regulatory_authority")
                    deadline_str = classification_result.get("compliance_deadline")
                    if deadline_str and isinstance(deadline_str, str) and deadline_str.lower() not in ["null", "none", ""]:
                        try:
                            from dateutil import parser as date_parser
                            framework.compliance_deadline = date_parser.parse(deadline_str)
                        except Exception:
                            pass
                    framework.penalty_for_non_compliance = classification_result.get("penalty_for_non_compliance")
                    framework.adoption_approach = classification_result.get("adoption_approach")
                    framework.required_artifacts = classification_result.get("required_artifacts")
                
                db.commit()
        except Exception as e:
            print(f"[PARSE] Classification failed during parse: {e}", flush=True)
    
    framework.upload_status = "parsing"
    framework.parse_error = None
    db.commit()

    tenant_slug = getattr(http_request.state, "tenant_slug", None)
    if not tenant_slug:
        raise HTTPException(status_code=400, detail="Tenant context required")

    from ....tasks.base import tenant_rate_limit, RateLimitExceeded
    try:
        tenant_rate_limit(tenant_slug, bucket="framework_parse")
    except RateLimitExceeded:
        raise HTTPException(status_code=429, detail="Too many framework parses queued; try again shortly")
    from ....tasks.frameworks import parse_framework as _parse_task
    _parse_task.delay(tenant_slug, framework_id, file_path, file_type, framework_name)

    return {
        "message": "Parsing started in background. Refresh the page periodically to check status.",
        "framework_id": framework_id,
        "status": "parsing"
    }


@router.get("/{framework_id}/parse-status")
def get_parse_status(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Check the parsing status of a framework."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    controls_count = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).count() if framework.upload_status == "parsed" else 0
    
    return {
        "framework_id": framework_id,
        "status": framework.upload_status,
        "parse_error": framework.parse_error,
        "parsed_at": framework.parsed_at.isoformat() if framework.parsed_at else None,
        "controls_count": controls_count
    }


@router.post("/{framework_id}/retry-parse")
def retry_framework_parsing(
    framework_id: int,
    http_request: Request,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Reset a stuck framework and restart parsing.
    
    Use this to recover from:
    - Frameworks stuck in 'parsing' status due to server restart (>10 minutes)
    - Failed parsing that needs to be retried
    - Any framework that didn't complete parsing properly
    
    Safety: Only allows retry if parsing is stale (>10 min) or status is failed/uploaded/text_extracted
    """
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    old_status = framework.upload_status
    
    # Check if retry is allowed based on status and staleness
    if old_status == 'parsing':
        # Check if parsing is stale (running for more than 10 minutes)
        if framework.updated_at:
            time_since_update = (datetime.utcnow() - framework.updated_at).total_seconds()
            if time_since_update < 600:  # 10 minutes = 600 seconds
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail=f"Parsing is still in progress (started {int(time_since_update)}s ago). Wait at least 10 minutes before retrying."
                )
    elif old_status in ['parsed', 'published', 'completed']:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Framework already has status '{old_status}'. No retry needed."
        )
    
    # Clear any partial parsed data to prevent duplicates
    existing_control_ids = db.query(ParsedFrameworkControl.id).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).all()
    
    if existing_control_ids:
        ctrl_ids = [c.id for c in existing_control_ids]
        # Delete associated assessment items
        assessment_item_ids = db.query(AssessmentItem.id).filter(
            AssessmentItem.parsed_control_id.in_(ctrl_ids)
        ).all()
        ai_ids = [a.id for a in assessment_item_ids]
        
        if ai_ids:
            db.query(AssessmentRemediation).filter(
                AssessmentRemediation.assessment_item_id.in_(ai_ids)
            ).delete(synchronize_session="fetch")
            db.query(AssessmentEvidence).filter(
                AssessmentEvidence.assessment_item_id.in_(ai_ids)
            ).delete(synchronize_session="fetch")
            db.query(AssessmentItem).filter(
                AssessmentItem.id.in_(ai_ids)
            ).delete(synchronize_session="fetch")
        
        db.query(FrameworkControlAlignment).filter(
            FrameworkControlAlignment.parsed_control_id.in_(ctrl_ids)
        ).delete(synchronize_session="fetch")
        db.query(ControlEvidenceMapping).filter(
            ControlEvidenceMapping.parsed_control_id.in_(ctrl_ids)
        ).delete(synchronize_session="fetch")
        db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.id.in_(ctrl_ids)
        ).delete(synchronize_session="fetch")
        db.flush()
    
    # Reset status and start fresh parsing
    framework.upload_status = "parsing"
    framework.parse_error = None
    framework.updated_at = datetime.utcnow()  # Update timestamp to track staleness
    db.commit()
    
    file_path = framework.file_path
    file_type = framework.file_type
    framework_name = framework.name
    
    if not os.path.exists(file_path):
        framework.upload_status = "failed"
        framework.parse_error = "Framework file not found on disk"
        db.commit()
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework file not found on disk"
        )

    tenant_slug = getattr(http_request.state, "tenant_slug", None)
    if not tenant_slug:
        raise HTTPException(status_code=400, detail="Tenant context required")

    from ....tasks.base import tenant_rate_limit, RateLimitExceeded
    try:
        tenant_rate_limit(tenant_slug, bucket="framework_parse")
    except RateLimitExceeded:
        raise HTTPException(status_code=429, detail="Too many framework parses queued; try again shortly")
    from ....tasks.frameworks import parse_framework as _parse_task
    _parse_task.delay(tenant_slug, framework_id, file_path, file_type, framework_name)
    
    return {
        "message": f"Parsing restarted. Previous status was '{old_status}'.",
        "framework_id": framework_id,
        "status": "parsing"
    }


@router.post("/{framework_id}/parse-sync")
def parse_framework_document_sync(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Parse a framework document synchronously (for small documents only).
    
    Warning: This may timeout for large documents. Use the async /parse endpoint instead.
    """
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    file_path = framework.file_path
    file_type = framework.file_type
    framework_name = framework.name
    
    framework.upload_status = "parsing"
    db.commit()
    
    try:
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Framework file not found on disk"
            )
        
        extracted_text = ""
        if file_type == "pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted_text = "\n\n".join(text_parts)
        elif file_type == "docx":
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            extracted_text = "\n".join(text_parts)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_type}"
            )
        
        if not extracted_text.strip():
            db.rollback()
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "parsed"
                fw.parsed_at = datetime.utcnow()
                fw.parse_error = "No text could be extracted from the document"
                db.commit()
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text could be extracted from the document"
            )
        
        parsed_controls_data, doc_structure = parse_document_with_chunking(extracted_text, framework_name)
        
        if not parsed_controls_data:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "parsed"
                fw.parsed_at = datetime.utcnow()
                fw.parse_error = "No controls found in document"
                db.commit()
            return {"message": "No controls found", "controls": []}
        
        fw_check = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
        if not fw_check:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Framework was deleted during parsing"
            )
        
        existing_controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework_id
        ).all()
        existing_control_ids = [c.id for c in existing_controls]
        
        if existing_control_ids:
            assessment_item_ids = db.query(AssessmentItem.id).filter(
                AssessmentItem.parsed_control_id.in_(existing_control_ids)
            ).all()
            ai_ids = [a.id for a in assessment_item_ids]
            
            if ai_ids:
                db.query(AssessmentRemediation).filter(
                    AssessmentRemediation.assessment_item_id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
                
                db.query(AssessmentEvidence).filter(
                    AssessmentEvidence.assessment_item_id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
                
                db.query(AssessmentItem).filter(
                    AssessmentItem.id.in_(ai_ids)
                ).delete(synchronize_session="fetch")
            
            db.query(FrameworkControlAlignment).filter(
                FrameworkControlAlignment.parsed_control_id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
            
            db.query(ControlEvidenceMapping).filter(
                ControlEvidenceMapping.parsed_control_id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
            
            db.query(ParsedFrameworkControl).filter(
                ParsedFrameworkControl.id.in_(existing_control_ids)
            ).delete(synchronize_session="fetch")
        
        db.flush()
        
        created_controls = []
        for idx, control_data in enumerate(parsed_controls_data, start=1):
            raw_reference = control_data.get("original_reference", "")
            cleaned_reference = clean_section_reference(raw_reference) if raw_reference else None
            
            control_id = cleaned_reference if cleaned_reference else f"CTRL-{idx:03d}"
            
            parent_ref = control_data.get("parent_reference", "")
            cleaned_parent_ref = clean_section_reference(parent_ref) if parent_ref else None
            
            hierarchy_level = control_data.get("hierarchy_level", 1)
            if not isinstance(hierarchy_level, int) or hierarchy_level < 1:
                hierarchy_level = 1
            elif hierarchy_level > 5:
                hierarchy_level = 5
            
            testing_procedure = control_data.get("testing_procedure", "")
            ai_notes = control_data.get("ai_notes", "")
            
            full_ai_notes = ai_notes
            if testing_procedure:
                full_ai_notes = f"Testing: {testing_procedure}" + (f"\n{ai_notes}" if ai_notes else "")
            
            parsed_control = ParsedFrameworkControl(
                uploaded_framework_id=framework_id,
                control_id=control_id,
                original_reference=cleaned_reference,
                title=control_data.get("title", "Untitled Control")[:500],
                description=control_data.get("description"),
                full_text=control_data.get("full_text"),
                domain=control_data.get("domain"),
                category=control_data.get("category"),
                is_mandatory=control_data.get("is_mandatory", True),
                priority=control_data.get("priority", "medium"),
                section_number=cleaned_reference,
                parent_section=cleaned_parent_ref,
                ai_confidence=control_data.get("ai_confidence"),
                ai_notes=full_ai_notes[:1000] if full_ai_notes else None,
                is_verified=False
            )
            db.add(parsed_control)
            db.flush()
            
            evidence_requirements = control_data.get("evidence_requirements", [])
            # Always persist evidence_requirements as JSON so certification router can read them
            if evidence_requirements:
                parsed_control.evidence_requirements = evidence_requirements
                for ev_req in evidence_requirements:
                    if isinstance(ev_req, dict):
                        ev_type = ev_req.get("type", "document")
                        ev_title = ev_req.get("title", "")
                        ev_description = ev_req.get("description", "")
                        ev_is_required = ev_req.get("is_required", True)
                        ev_examples = ev_req.get("artifact_examples", [])
                        ev_frequency = ev_req.get("review_frequency", "")
                        
                        full_description = ev_title
                        if ev_description:
                            full_description = f"{ev_title}: {ev_description}" if ev_title else ev_description
                        if ev_examples and isinstance(ev_examples, list):
                            full_description += f" Examples: {', '.join(ev_examples[:3])}"
                        if ev_frequency:
                            full_description += f" (Frequency: {ev_frequency})"
                        
                        valid_types = [
                            "policy", "procedure", "configuration", "log", "report", 
                            "contract", "attestation", "register", "matrix", "plan", 
                            "screenshot", "training", "assessment", "document",
                            "interview", "observation"
                        ]
                        if ev_type not in valid_types:
                            ev_type = "document"
                        
                        evidence_mapping = ControlEvidenceMapping(
                            parsed_control_id=parsed_control.id,
                            evidence_type=ev_type,
                            evidence_description=full_description[:1000] if full_description else None,
                            is_required=ev_is_required if isinstance(ev_is_required, bool) else True,
                            suggested_by_ai=True
                        )
                        db.add(evidence_mapping)
            else:
                evidence_types = control_data.get("evidence_types", [])
                for evidence_type in evidence_types:
                    if evidence_type in ["policy", "procedure", "configuration", "log", "report", "contract"]:
                        evidence_mapping = ControlEvidenceMapping(
                            parsed_control_id=parsed_control.id,
                            evidence_type=evidence_type,
                            is_required=True,
                            suggested_by_ai=True
                        )
                        db.add(evidence_mapping)
            
            created_controls.append(parsed_control)
        
        fw_final = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
        if fw_final:
            fw_final.upload_status = "parsed"
            fw_final.parsed_at = datetime.utcnow()
            fw_final.parse_error = None
            fw_final.document_structure = doc_structure
        
        db.commit()
        
        for control in created_controls:
            db.refresh(control)
        
        return {
            "message": f"Successfully parsed {len(created_controls)} controls",
            "framework_id": framework_id,
            "controls_count": len(created_controls),
            "controls": [serialize_parsed_control(c) for c in created_controls]
        }
    
    except HTTPException:
        db.rollback()
        try:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "failed"
                fw.parse_error = "Parsing failed"
                db.commit()
        except Exception:
            pass
        raise
    except Exception as e:
        db.rollback()
        error_msg = str(e)
        try:
            fw = db.query(UploadedFramework).filter(UploadedFramework.id == framework_id).first()
            if fw:
                fw.upload_status = "failed"
                fw.parse_error = error_msg[:500]
                db.commit()
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Parsing failed: {error_msg}"
        )


@router.get("/{framework_id}/controls")
def list_parsed_controls(
    framework_id: int,
    domain: Optional[str] = None,
    category: Optional[str] = None,
    is_verified: Optional[bool] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    query = db.query(ParsedFrameworkControl).options(
        joinedload(ParsedFrameworkControl.evidence_mappings)
    ).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    )
    
    if domain:
        query = query.filter(ParsedFrameworkControl.domain == domain)
    if category:
        query = query.filter(ParsedFrameworkControl.category == category)
    if is_verified is not None:
        query = query.filter(ParsedFrameworkControl.is_verified == is_verified)
    
    total = query.count()
    controls = query.order_by(ParsedFrameworkControl.control_id).offset(skip).limit(limit).all()
    
    return {
        "items": [serialize_parsed_control(c) for c in controls],
        "total": total,
        "skip": skip,
        "limit": limit,
        "framework_id": framework_id,
        "framework_name": framework.name
    }


@router.put("/controls/{control_id}")
def update_parsed_control(
    control_id: int,
    update_data: ParsedControlUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    control = db.query(ParsedFrameworkControl).options(
        joinedload(ParsedFrameworkControl.evidence_mappings),
        joinedload(ParsedFrameworkControl.uploaded_framework)
    ).filter(
        ParsedFrameworkControl.id == control_id
    ).first()
    
    if not control:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Parsed control not found"
        )
    
    validate_framework_access(current_user, control.uploaded_framework, db)
    
    if update_data.title is not None:
        control.title = update_data.title[:500]
    if update_data.description is not None:
        control.description = update_data.description
    if update_data.domain is not None:
        control.domain = update_data.domain
    if update_data.category is not None:
        control.category = update_data.category
    if update_data.is_mandatory is not None:
        control.is_mandatory = update_data.is_mandatory
    if update_data.priority is not None:
        if update_data.priority in ["high", "medium", "low"]:
            control.priority = update_data.priority
    
    control.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(control)
    
    return serialize_parsed_control(control)


@router.post("/controls/{control_id}/verify")
def verify_parsed_control(
    control_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    control = db.query(ParsedFrameworkControl).options(
        joinedload(ParsedFrameworkControl.evidence_mappings),
        joinedload(ParsedFrameworkControl.uploaded_framework)
    ).filter(
        ParsedFrameworkControl.id == control_id
    ).first()
    
    if not control:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Parsed control not found"
        )
    
    validate_framework_access(current_user, control.uploaded_framework, db)
    
    control.is_verified = True
    control.verified_by = current_user.id
    control.verified_at = datetime.utcnow()
    control.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(control)
    
    return {
        "message": "Control verified successfully",
        "control": serialize_parsed_control(control)
    }


@router.delete("/controls/{control_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_parsed_control(
    control_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    control = db.query(ParsedFrameworkControl).options(
        joinedload(ParsedFrameworkControl.uploaded_framework)
    ).filter(
        ParsedFrameworkControl.id == control_id
    ).first()
    
    if not control:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Parsed control not found"
        )
    
    validate_framework_access(current_user, control.uploaded_framework, db)
    
    db.delete(control)
    db.commit()
    
    return None


def _enhance_controls_body(db, framework_id: int, framework_name: str):
    """Body of the AI-enhance job. Takes an open tenant-scoped session."""
    try:
        if not check_ai_available():
            print("[ENHANCE] OpenAI API key not configured", flush=True)
            return
        
        client = get_openai_client()
        
        controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework_id
        ).all()
        
        if not controls:
            print(f"[ENHANCE] No controls found for framework {framework_id}", flush=True)
            return
        
        framework = db.query(UploadedFramework).filter(
            UploadedFramework.id == framework_id
        ).first()
        
        if framework:
            framework.upload_status = "enhancing"
            db.commit()
        
        print(f"[ENHANCE] Starting enhancement for {len(controls)} controls in {framework_name}", flush=True)
        
        batch_size = 10
        total_batches = (len(controls) + batch_size - 1) // batch_size
        
        for batch_num, i in enumerate(range(0, len(controls), batch_size), start=1):
            batch = controls[i:i + batch_size]
            print(f"[ENHANCE] Processing batch {batch_num}/{total_batches} ({len(batch)} controls)...", flush=True)
            
            controls_data = []
            for c in batch:
                controls_data.append({
                    "id": c.id,
                    "control_id": c.control_id,
                    "title": c.title,
                    "description": c.description or "",
                    "full_text": c.full_text or ""
                })
            
            controls_json = json.dumps(controls_data, indent=2)
            
            prompt = f"""For framework "{framework_name}", generate audit-ready evidence requirements for these controls.

INPUT CONTROLS:
{controls_json}

For EACH control, provide:
- id: Keep the same ID from input
- evidence_requirements: Array of 2-4 evidence items, each with:
  {{
    "type": "policy|procedure|configuration|log|report|contract|attestation|register|screenshot|interview|test_results",
    "title": "specific evidence name (e.g., 'Access Control Policy Document')",
    "description": "what auditor looks for and how to obtain",
    "is_required": true/false
  }}

Be specific and practical. Example evidence types:
- policy: Written policies (Information Security Policy, Access Control Policy)
- procedure: Step-by-step procedures (Incident Response Procedure, Change Management Procedure)
- configuration: System configs (Firewall rules, AD group memberships, encryption settings)
- log: Audit logs (Login logs, change logs, access logs)
- report: Periodic reports (Vulnerability scan reports, compliance dashboards)
- attestation: Signed acknowledgments (Training completion, policy acceptance)
- test_results: Test evidence (Penetration test reports, DR test results)

Return JSON with "controls" array containing objects with "id" and "evidence_requirements"."""

            try:
                import time
                start_time = time.time()
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are a GRC expert adding audit-ready evidence requirements to compliance controls. Be specific and practical."},
                        {"role": "user", "content": prompt}
                    ],
                    response_format={"type": "json_object"},
                    max_tokens=16384,
                    temperature=0
                )
                
                elapsed = time.time() - start_time
                result_text = response.choices[0].message.content or "{}"
                result = json.loads(result_text)
                enhanced_batch = result.get("controls", [])
                
                id_to_evidence = {int(item["id"]): item.get("evidence_requirements", []) for item in enhanced_batch}
                
                matched_count = 0
                for control in batch:
                    if control.id in id_to_evidence:
                        evidence_reqs = id_to_evidence[control.id]
                        if evidence_reqs:
                            control.evidence_requirements = evidence_reqs
                            control.updated_at = datetime.utcnow()
                            matched_count += 1
                
                print(f"[ENHANCE] Matched {matched_count}/{len(batch)} controls with evidence", flush=True)
                
                db.commit()
                print(f"[ENHANCE] Batch {batch_num} completed in {elapsed:.1f}s", flush=True)
                
            except Exception as e:
                print(f"[ENHANCE] Error in batch {batch_num}: {str(e)}", flush=True)
                continue
        
        if framework:
            framework.upload_status = "published"
            framework.updated_at = datetime.utcnow()
            db.commit()
        
        print(f"[ENHANCE] Enhancement complete for framework {framework_id}", flush=True)
        return {"status": "completed", "framework_id": framework_id}
    except Exception as e:
        print(f"[ENHANCE] Error: {str(e)}", flush=True)
        raise


def enhance_framework_controls_background(framework_id: int, framework_name: str, tenant_slug: str):
    """Legacy in-process entry; new dispatches use the Celery task."""
    from ....db import open_tenant_session
    db = open_tenant_session(tenant_slug)
    try:
        return _enhance_controls_body(db, framework_id, framework_name)
    finally:
        db.close()


@router.post("/frameworks/{framework_id}/enhance")
def enhance_framework_with_evidence(
    framework_id: int,
    background_tasks: BackgroundTasks,
    http_request: Request,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Enhance all controls in a framework with AI-generated evidence requirements."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    if framework.upload_status == "enhancing":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Framework is already being enhanced"
        )
    
    control_count = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).count()
    
    if control_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No controls found to enhance"
        )
    
    controls_with_evidence = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id,
        ParsedFrameworkControl.evidence_requirements != None,
        func.jsonb_array_length(ParsedFrameworkControl.evidence_requirements) > 0
    ).count()
    
    tenant_slug = getattr(http_request.state, "tenant_slug", None)
    if not tenant_slug:
        raise HTTPException(status_code=400, detail="Tenant context required")

    from ....tasks.base import tenant_rate_limit, RateLimitExceeded
    try:
        tenant_rate_limit(tenant_slug, bucket="framework_enhance")
    except RateLimitExceeded:
        raise HTTPException(status_code=429, detail="Too many enhancement runs queued; try again shortly")
    from ....tasks.frameworks import enhance_framework_controls as _enhance_task
    _enhance_task.delay(tenant_slug, framework_id, framework.name)
    
    return {
        "message": "Enhancement started",
        "framework_id": framework_id,
        "framework_name": framework.name,
        "total_controls": control_count,
        "controls_with_evidence": controls_with_evidence,
        "estimated_time_minutes": max(1, (control_count // 10) * 0.5)
    }


@router.get("/frameworks/{framework_id}/enhancement-status")
def get_enhancement_status(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Get the enhancement status for a framework."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    total_controls = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).count()
    
    controls_with_evidence = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id,
        ParsedFrameworkControl.evidence_requirements != None,
        func.jsonb_array_length(ParsedFrameworkControl.evidence_requirements) > 0
    ).count()
    
    return {
        "framework_id": framework_id,
        "framework_name": framework.name,
        "status": framework.upload_status,
        "total_controls": total_controls,
        "controls_with_evidence": controls_with_evidence,
        "enhancement_progress": round((controls_with_evidence / total_controls * 100) if total_controls > 0 else 0, 1)
    }


EVIDENCE_GENERATION_PROMPT = """You are a GRC expert specializing in audit evidence and compliance documentation. For the given control requirement, generate SPECIFIC evidence requirements that would satisfy an auditor.

CONTROL INFORMATION:
Control ID: {control_id}
Title: {title}
Description: {description}
Full Text: {full_text}
Domain: {domain}
Is Mandatory: {is_mandatory}

Generate 1-5 SPECIFIC evidence requirements. Be EXACT about what documentation, screenshots, exports, or records are needed.

For each evidence requirement, specify:
1. evidence_title: Clear, specific title
2. evidence_description: Detailed description (2-3 sentences)
3. evidence_type: One of: policy, procedure, configuration, screenshot, log, report, contract, attestation, certificate, training_record
4. evidence_format: e.g., "PDF document", "System screenshot", "CSV export", "Signed PDF"
5. exact_requirements: Array of specific items needed (e.g., ["Rule definitions", "Source/destination IPs", "Date of last update"])
6. acceptance_criteria: Array of criteria (e.g., ["Dated within last 12 months", "Signed by manager", "Shows complete configuration"])
7. sample_evidence: Brief description of an ideal sample
8. collection_guidance: How to collect this evidence (1-2 sentences)
9. collection_frequency: one-time, monthly, quarterly, annually, or on-change
10. retention_period: e.g., "3 years"
11. priority: high (critical controls), medium, or low
12. is_mandatory: true if required for compliance, false if supporting

Return JSON with "evidence_requirements" array."""


def generate_evidence_requirements_for_controls_batch(
    controls: List[ParsedFrameworkControl],
    framework_name: str
) -> List[dict]:
    """Generate evidence requirements for a batch of controls using AI."""
    if not check_ai_available():
        return []
    
    client = get_openai_client()
    
    results = []
    
    for control in controls:
        prompt = EVIDENCE_GENERATION_PROMPT.format(
            control_id=control.control_id,
            title=control.title,
            description=control.description or "",
            full_text=control.full_text or "",
            domain=control.domain or "General",
            is_mandatory=control.is_mandatory
        )
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": GRC_SME_SYSTEM_PROMPT},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_tokens=4096,
                temperature=0
            )
            
            result_text = response.choices[0].message.content or "{}"
            result = json.loads(result_text)
            evidence_reqs = result.get("evidence_requirements", [])
            
            for req in evidence_reqs:
                req["parsed_control_id"] = control.id
            
            results.extend(evidence_reqs)
            
        except Exception as e:
            print(f"[EVIDENCE] Error generating for control {control.id}: {e}", flush=True)
            continue
    
    return results


def _generate_evidence_reqs_body(db, framework_id: int, framework_name: str):
    """Body of the generate-evidence-reqs job. Takes an open tenant-scoped session."""
    try:
        framework = db.query(UploadedFramework).filter(
            UploadedFramework.id == framework_id
        ).first()
        
        if not framework:
            return
        
        framework.parse_error = "Generating evidence requirements..."
        db.commit()
        
        controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework_id
        ).all()
        
        if not controls:
            framework.parse_error = "No controls found"
            db.commit()
            return
        
        total_controls = len(controls)
        batch_size = 5
        total_batches = (total_controls + batch_size - 1) // batch_size
        total_requirements = 0
        
        for batch_num, i in enumerate(range(0, total_controls, batch_size), start=1):
            batch = controls[i:i + batch_size]
            
            progress = round((batch_num / total_batches) * 100)
            framework.parse_error = f"Generating evidence requirements... {progress}% ({batch_num}/{total_batches} batches)"
            db.commit()
            
            print(f"[EVIDENCE] Processing batch {batch_num}/{total_batches} ({len(batch)} controls)", flush=True)
            
            evidence_reqs = generate_evidence_requirements_for_controls_batch(batch, framework_name)
            
            for req in evidence_reqs:
                evidence_record = ControlEvidenceRequirement(
                    framework_id=framework_id,
                    parsed_control_id=req.get("parsed_control_id"),
                    evidence_title=req.get("evidence_title", "Evidence Requirement")[:500],
                    evidence_description=req.get("evidence_description", ""),
                    evidence_type=req.get("evidence_type", "policy")[:100],
                    evidence_format=req.get("evidence_format", "PDF")[:100] if req.get("evidence_format") else None,
                    exact_requirements=req.get("exact_requirements", []),
                    acceptance_criteria=req.get("acceptance_criteria", []),
                    sample_evidence=req.get("sample_evidence"),
                    collection_guidance=req.get("collection_guidance"),
                    collection_frequency=req.get("collection_frequency", "annually")[:50] if req.get("collection_frequency") else None,
                    retention_period=req.get("retention_period", "3 years")[:100] if req.get("retention_period") else None,
                    priority=normalize_priority(req.get("priority", "medium")),
                    is_mandatory=req.get("is_mandatory", True),
                    status="draft",
                    ai_confidence=0.85,
                    ai_reasoning=f"AI-generated evidence requirement for control {req.get('parsed_control_id')}"
                )
                db.add(evidence_record)
                total_requirements += 1
            
            db.commit()
        
        framework.parse_error = None
        db.commit()

        print(f"[EVIDENCE] Completed! Generated {total_requirements} evidence requirements for {total_controls} controls", flush=True)
        return {"status": "completed", "framework_id": framework_id, "total_requirements": total_requirements, "total_controls": total_controls}
    except Exception as e:
        print(f"[EVIDENCE] Error: {str(e)}", flush=True)
        if framework:
            framework.parse_error = f"Error generating evidence: {str(e)[:200]}"
            db.commit()
        raise


def generate_evidence_requirements_background(framework_id: int, framework_name: str, tenant_slug: str):
    """Legacy in-process entry; new dispatches use the Celery task."""
    from ....db import open_tenant_session
    db = open_tenant_session(tenant_slug)
    try:
        return _generate_evidence_reqs_body(db, framework_id, framework_name)
    finally:
        db.close()


@router.post("/{framework_id}/generate-evidence-requirements")
def generate_evidence_requirements(
    framework_id: int,
    background_tasks: BackgroundTasks,
    http_request: Request,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Generate AI-powered evidence requirements for all controls in a framework."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    if not check_ai_available():
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI features unavailable. OpenAI API key not configured."
        )
    
    control_count = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).count()
    
    if control_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No controls found in this framework. Please parse the framework first."
        )
    
    existing_requirements = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.framework_id == framework_id
    ).count()
    
    tenant_slug = getattr(http_request.state, "tenant_slug", None)
    if not tenant_slug:
        raise HTTPException(status_code=400, detail="Tenant context required")

    from ....tasks.base import tenant_rate_limit, RateLimitExceeded
    try:
        tenant_rate_limit(tenant_slug, bucket="framework_evidence_reqs")
    except RateLimitExceeded:
        raise HTTPException(status_code=429, detail="Too many evidence-requirement jobs queued; try again shortly")
    from ....tasks.frameworks import generate_evidence_requirements as _evid_task
    _evid_task.delay(tenant_slug, framework_id, framework.name)
    
    return {
        "message": "Evidence requirement generation started",
        "framework_id": framework_id,
        "framework_name": framework.name,
        "total_controls": control_count,
        "existing_requirements": existing_requirements,
        "estimated_time_minutes": max(1, (control_count // 5) * 0.5)
    }


@router.get("/{framework_id}/evidence-requirements")
def list_evidence_requirements(
    framework_id: int,
    status_filter: Optional[str] = Query(None, alias="status"),
    control_id: Optional[int] = Query(None),
    evidence_type: Optional[str] = Query(None),
    priority: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=500),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """List all evidence requirements for a framework with filtering options."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    query = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.framework_id == framework_id,
        ControlEvidenceRequirement.is_active == True
    )
    
    if status_filter:
        query = query.filter(ControlEvidenceRequirement.status == status_filter)
    
    if control_id:
        query = query.filter(ControlEvidenceRequirement.parsed_control_id == control_id)
    
    if evidence_type:
        query = query.filter(ControlEvidenceRequirement.evidence_type == evidence_type)
    
    if priority:
        query = query.filter(ControlEvidenceRequirement.priority == priority)
    
    total = query.count()
    
    requirements = query.order_by(
        ControlEvidenceRequirement.parsed_control_id,
        ControlEvidenceRequirement.display_order,
        ControlEvidenceRequirement.id
    ).offset(skip).limit(limit).all()
    
    status_counts = db.query(
        ControlEvidenceRequirement.status,
        func.count(ControlEvidenceRequirement.id)
    ).filter(
        ControlEvidenceRequirement.framework_id == framework_id,
        ControlEvidenceRequirement.is_active == True
    ).group_by(ControlEvidenceRequirement.status).all()
    
    return {
        "framework_id": framework_id,
        "framework_name": framework.name,
        "total": total,
        "status_counts": {s: c for s, c in status_counts},
        "requirements": [
            {
                "id": req.id,
                "parsed_control_id": req.parsed_control_id,
                "control_id": req.parsed_control.control_id if req.parsed_control else None,
                "control_title": req.parsed_control.title if req.parsed_control else None,
                "evidence_title": req.evidence_title,
                "evidence_description": req.evidence_description,
                "evidence_type": req.evidence_type,
                "evidence_format": req.evidence_format,
                "exact_requirements": req.exact_requirements,
                "acceptance_criteria": req.acceptance_criteria,
                "sample_evidence": req.sample_evidence,
                "collection_guidance": req.collection_guidance,
                "collection_frequency": req.collection_frequency,
                "retention_period": req.retention_period,
                "priority": req.priority,
                "is_mandatory": req.is_mandatory,
                "status": req.status,
                "ai_confidence": req.ai_confidence,
                "ai_reasoning": req.ai_reasoning,
                "rejection_reason": req.rejection_reason,
                "created_at": req.created_at.isoformat() if req.created_at else None,
                "submitted_at": req.submitted_at.isoformat() if req.submitted_at else None,
                "reviewed_at": req.reviewed_at.isoformat() if req.reviewed_at else None,
                "approved_at": req.approved_at.isoformat() if req.approved_at else None
            }
            for req in requirements
        ]
    }


class WorkflowSubmitRequest(BaseModel):
    notes: Optional[str] = None


class WorkflowReviewRequest(BaseModel):
    notes: Optional[str] = None


class WorkflowApproveRequest(BaseModel):
    notes: Optional[str] = None


class WorkflowRejectRequest(BaseModel):
    reason: str


@router.post("/evidence-requirements/{requirement_id}/submit")
def submit_evidence_requirement(
    requirement_id: int,
    request: WorkflowSubmitRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Submit an evidence requirement for review."""
    requirement = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.id == requirement_id,
        ControlEvidenceRequirement.is_active == True
    ).first()
    
    if not requirement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence requirement not found"
        )
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == requirement.framework_id
    ).first()
    
    if framework:
        validate_framework_access(current_user, framework, db)
    
    if requirement.status != "draft":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot submit requirement with status '{requirement.status}'. Only draft requirements can be submitted."
        )
    
    previous_status = requirement.status
    requirement.status = "submitted"
    requirement.submitted_by = current_user.id
    requirement.submitted_at = datetime.utcnow()
    requirement.submission_notes = request.notes
    
    history = EvidenceRequirementHistory(
        evidence_requirement_id=requirement_id,
        action="submitted",
        previous_status=previous_status,
        new_status="submitted",
        performed_by=current_user.id,
        notes=request.notes
    )
    db.add(history)
    
    db.commit()
    
    return {
        "message": "Evidence requirement submitted for review",
        "requirement_id": requirement_id,
        "status": requirement.status,
        "submitted_at": requirement.submitted_at.isoformat()
    }


@router.post("/evidence-requirements/{requirement_id}/review")
def start_review_evidence_requirement(
    requirement_id: int,
    request: WorkflowReviewRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Start review of an evidence requirement (requires reviewer role)."""
    requirement = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.id == requirement_id,
        ControlEvidenceRequirement.is_active == True
    ).first()
    
    if not requirement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence requirement not found"
        )
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == requirement.framework_id
    ).first()
    
    if framework:
        validate_framework_access(current_user, framework, db)
    
    if requirement.status != "submitted":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot review requirement with status '{requirement.status}'. Only submitted requirements can be reviewed."
        )
    
    previous_status = requirement.status
    requirement.status = "pending_review"
    requirement.reviewer_id = current_user.id
    requirement.reviewed_at = datetime.utcnow()
    requirement.review_notes = request.notes
    
    history = EvidenceRequirementHistory(
        evidence_requirement_id=requirement_id,
        action="review_started",
        previous_status=previous_status,
        new_status="pending_review",
        performed_by=current_user.id,
        notes=request.notes
    )
    db.add(history)
    
    db.commit()
    
    return {
        "message": "Review started for evidence requirement",
        "requirement_id": requirement_id,
        "status": requirement.status,
        "reviewed_at": requirement.reviewed_at.isoformat()
    }


@router.post("/evidence-requirements/{requirement_id}/approve")
def approve_evidence_requirement(
    requirement_id: int,
    request: WorkflowApproveRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Approve an evidence requirement."""
    requirement = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.id == requirement_id,
        ControlEvidenceRequirement.is_active == True
    ).first()
    
    if not requirement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence requirement not found"
        )
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == requirement.framework_id
    ).first()
    
    if framework:
        validate_framework_access(current_user, framework, db)
    
    if requirement.status not in ["submitted", "pending_review"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot approve requirement with status '{requirement.status}'. Only submitted or pending_review requirements can be approved."
        )
    
    previous_status = requirement.status
    requirement.status = "approved"
    requirement.approver_id = current_user.id
    requirement.approved_at = datetime.utcnow()
    requirement.approval_notes = request.notes
    
    history = EvidenceRequirementHistory(
        evidence_requirement_id=requirement_id,
        action="approved",
        previous_status=previous_status,
        new_status="approved",
        performed_by=current_user.id,
        notes=request.notes
    )
    db.add(history)
    
    db.commit()
    
    return {
        "message": "Evidence requirement approved",
        "requirement_id": requirement_id,
        "status": requirement.status,
        "approved_at": requirement.approved_at.isoformat()
    }


@router.post("/evidence-requirements/{requirement_id}/reject")
def reject_evidence_requirement(
    requirement_id: int,
    request: WorkflowRejectRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Reject an evidence requirement with a reason."""
    requirement = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.id == requirement_id,
        ControlEvidenceRequirement.is_active == True
    ).first()
    
    if not requirement:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence requirement not found"
        )
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == requirement.framework_id
    ).first()
    
    if framework:
        validate_framework_access(current_user, framework, db)
    
    if requirement.status not in ["submitted", "pending_review"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot reject requirement with status '{requirement.status}'. Only submitted or pending_review requirements can be rejected."
        )
    
    previous_status = requirement.status
    requirement.status = "rejected"
    requirement.rejection_reason = request.reason
    requirement.approver_id = current_user.id
    requirement.approved_at = datetime.utcnow()
    
    history = EvidenceRequirementHistory(
        evidence_requirement_id=requirement_id,
        action="rejected",
        previous_status=previous_status,
        new_status="rejected",
        performed_by=current_user.id,
        notes=request.reason
    )
    db.add(history)
    
    db.commit()
    
    return {
        "message": "Evidence requirement rejected",
        "requirement_id": requirement_id,
        "status": requirement.status,
        "rejection_reason": requirement.rejection_reason
    }


@router.get("/{framework_id}/evidence-generation-status")
def get_evidence_generation_status(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Get the status of evidence requirement generation for a framework."""
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework not found"
        )
    
    validate_framework_access(current_user, framework, db)
    
    total_controls = db.query(ParsedFrameworkControl).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework_id
    ).count()
    
    total_requirements = db.query(ControlEvidenceRequirement).filter(
        ControlEvidenceRequirement.framework_id == framework_id,
        ControlEvidenceRequirement.is_active == True
    ).count()
    
    controls_with_requirements = db.query(
        func.count(func.distinct(ControlEvidenceRequirement.parsed_control_id))
    ).filter(
        ControlEvidenceRequirement.framework_id == framework_id,
        ControlEvidenceRequirement.is_active == True
    ).scalar() or 0
    
    status_counts = db.query(
        ControlEvidenceRequirement.status,
        func.count(ControlEvidenceRequirement.id)
    ).filter(
        ControlEvidenceRequirement.framework_id == framework_id,
        ControlEvidenceRequirement.is_active == True
    ).group_by(ControlEvidenceRequirement.status).all()
    
    is_generating = framework.parse_error and "Generating evidence requirements" in framework.parse_error
    
    return {
        "framework_id": framework_id,
        "framework_name": framework.name,
        "is_generating": is_generating,
        "progress_message": framework.parse_error if is_generating else None,
        "total_controls": total_controls,
        "controls_with_requirements": controls_with_requirements,
        "total_requirements": total_requirements,
        "average_requirements_per_control": round(total_requirements / controls_with_requirements, 1) if controls_with_requirements > 0 else 0,
        "status_breakdown": {s: c for s, c in status_counts}
    }


parser_router = router
