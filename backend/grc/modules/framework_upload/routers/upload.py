import os
import uuid
from typing import List, Optional
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy import func
from pydantic import BaseModel

from ....models import (
    UploadedFramework, ParsedFrameworkControl, FrameworkAssessment,
    GRCUser, Tenant, get_db, EvidenceControlMapping, Evidence, EvidenceAIAssessment
)
from ....routers.auth_router import require_auth, get_user_tenants, get_user_primary_tenant

router = APIRouter(prefix="/upload", tags=["Framework Upload - Upload"])

UPLOAD_DIR = "uploads/frameworks"

os.makedirs(UPLOAD_DIR, exist_ok=True)


class UploadedFrameworkResponse(BaseModel):
    id: int
    tenant_id: Optional[int]
    name: str
    description: Optional[str]
    file_name: str
    file_path: str
    file_size: Optional[int]
    file_type: str
    upload_status: str
    parse_error: Optional[str]
    parsed_at: Optional[datetime]
    framework_type: Optional[str]
    source_organization: Optional[str]
    version: Optional[str]
    effective_date: Optional[datetime]
    is_shared: bool
    is_active: bool
    uploaded_by: int
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class ExtractTextResponse(BaseModel):
    framework_id: int
    file_name: str
    file_type: str
    text_content: str
    text_length: int
    extraction_status: str


def validate_tenant_access(user: GRCUser, tenant_id: int, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this tenant's data"
        )


def serialize_framework(framework: UploadedFramework, controls_count: int = 0) -> dict:
    return {
        "id": framework.id,
        "tenant_id": framework.tenant_id,
        "name": framework.name,
        "description": framework.description,
        "file_name": framework.file_name,
        "file_path": framework.file_path,
        "file_size": framework.file_size,
        "file_type": framework.file_type,
        "upload_status": framework.upload_status,
        "parse_error": framework.parse_error,
        "parsed_at": framework.parsed_at.isoformat() if framework.parsed_at else None,
        "published_framework_id": framework.published_framework_id,
        "published_at": framework.published_at.isoformat() if framework.published_at else None,
        "framework_type": framework.framework_type,
        "source_organization": framework.source_organization,
        "version": framework.version,
        "effective_date": framework.effective_date.isoformat() if framework.effective_date else None,
        "is_shared": framework.is_shared,
        "is_active": framework.is_active,
        "uploaded_by": framework.uploaded_by,
        "uploader_name": framework.uploader.display_name if framework.uploader else None,
        "created_at": framework.created_at.isoformat() if framework.created_at else None,
        "updated_at": framework.updated_at.isoformat() if framework.updated_at else None,
        "parsed_controls_count": controls_count
    }


@router.post("", status_code=status.HTTP_201_CREATED)
async def upload_framework(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    framework_type: Optional[str] = Form(None),
    source_organization: Optional[str] = Form(None),
    version: Optional[str] = Form(None),
    tenant_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    file_ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if file_ext not in ['pdf', 'docx']:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX files are supported"
        )
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
    else:
        tenant_id = get_user_primary_tenant(current_user, db)
    
    unique_id = str(uuid.uuid4())
    safe_filename = f"{unique_id}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    content = await file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    db_framework = UploadedFramework(
        tenant_id=tenant_id,
        name=name,
        description=description,
        file_name=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_ext,
        upload_status="uploaded",
        framework_type=framework_type,
        source_organization=source_organization,
        version=version,
        uploaded_by=current_user.id,
        is_shared=False,
        is_active=True
    )
    db.add(db_framework)
    db.commit()
    db.refresh(db_framework)
    
    return serialize_framework(db_framework)


@router.get("")
def list_uploaded_frameworks(
    tenant_id: Optional[int] = None,
    status_filter: Optional[str] = Query(None, alias="status"),
    framework_type: Optional[str] = None,
    is_shared: Optional[bool] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    query = db.query(UploadedFramework)
    
    if user_tenants:
        from sqlalchemy import or_
        query = query.filter(
            or_(
                UploadedFramework.tenant_id.in_(user_tenants),
                UploadedFramework.is_shared == True
            )
        )
    else:
        query = query.filter(UploadedFramework.is_shared == True)
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
        query = query.filter(UploadedFramework.tenant_id == tenant_id)
    
    if status_filter:
        query = query.filter(UploadedFramework.upload_status == status_filter)
    
    if framework_type:
        query = query.filter(UploadedFramework.framework_type == framework_type)
    
    if is_shared is not None:
        query = query.filter(UploadedFramework.is_shared == is_shared)
    
    total = query.count()
    
    frameworks = query.order_by(UploadedFramework.created_at.desc()).offset(skip).limit(limit).all()
    
    result = []
    for fw in frameworks:
        controls_count = db.query(func.count(ParsedFrameworkControl.id)).filter(
            ParsedFrameworkControl.uploaded_framework_id == fw.id
        ).scalar() or 0
        result.append(serialize_framework(fw, controls_count))
    
    return {
        "items": result,
        "total": total,
        "skip": skip,
        "limit": limit
    }


@router.get("/{framework_id}")
def get_uploaded_framework(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    if framework.tenant_id and framework.tenant_id not in user_tenants and not framework.is_shared:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this framework"
        )
    
    controls_count = db.query(func.count(ParsedFrameworkControl.id)).filter(
        ParsedFrameworkControl.uploaded_framework_id == framework.id
    ).scalar() or 0
    
    return serialize_framework(framework, controls_count)


@router.delete("/{framework_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_uploaded_framework(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    # Enforce tenant access - framework must belong to user's tenant
    if framework.tenant_id and framework.tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to delete this framework"
        )
    
    # Prevent deletion of shared frameworks - they may be used across tenants
    if getattr(framework, 'is_shared', False):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot delete shared frameworks. Please unshare it first."
        )
    
    # Store file path for cleanup after successful commit
    file_path = framework.file_path
    
    try:
        # Get framework name for invalidating AI assessments
        framework_name = framework.name
        
        # Get all parsed control IDs for this framework
        parsed_control_ids = db.query(ParsedFrameworkControl.id).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework_id
        ).all()
        parsed_control_ids = [pc[0] for pc in parsed_control_ids]
        
        # Delete ALL EvidenceControlMapping records that reference these parsed controls
        # SECURITY NOTE: Framework ownership is verified above, so these controls belong
        # to the user's tenant. Any mappings to these controls are derived from this
        # framework and should be removed when the framework is deleted.
        if parsed_control_ids:
            db.query(EvidenceControlMapping).filter(
                EvidenceControlMapping.parsed_control_id.in_(parsed_control_ids)
            ).delete(synchronize_session=False)
        
        # Delete any mappings by uploaded_framework_id directly
        db.query(EvidenceControlMapping).filter(
            EvidenceControlMapping.uploaded_framework_id == framework_id
        ).delete(synchronize_session=False)
        
        # Invalidate AI assessments that reference this framework
        # Get all assessments that have clause_mappings referencing this framework
        assessments_to_update = db.query(EvidenceAIAssessment).filter(
            EvidenceAIAssessment.clause_mappings.isnot(None)
        ).all()
        
        for assessment in assessments_to_update:
            if assessment.clause_mappings:
                # Filter out mappings that reference the deleted framework
                updated_mappings = [
                    mapping for mapping in assessment.clause_mappings
                    if mapping.get('framework_name') != framework_name
                ]
                # Only update if mappings changed
                if len(updated_mappings) != len(assessment.clause_mappings):
                    assessment.clause_mappings = updated_mappings
                    # Unlock assessment so it can be re-run
                    assessment.is_locked = False
                    assessment.locked_at = None
                    assessment.locked_by = None
                    assessment.lock_reason = None
        
        db.delete(framework)
        db.commit()
        
        # Remove the file from disk AFTER successful commit
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except OSError:
                pass
        
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete framework: {str(e)}"
        )
    
    return None


@router.post("/{framework_id}/extract-text")
def extract_text_from_framework(
    framework_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    framework = db.query(UploadedFramework).filter(
        UploadedFramework.id == framework_id
    ).first()
    
    if not framework:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Uploaded framework not found"
        )
    
    if framework.tenant_id and framework.tenant_id not in user_tenants and not framework.is_shared:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this framework"
        )
    
    if not os.path.exists(framework.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Framework file not found on disk"
        )
    
    extracted_text = ""
    
    try:
        if framework.file_type == "pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(framework.file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted_text = "\n\n".join(text_parts)
        
        elif framework.file_type == "docx":
            from docx import Document
            doc = Document(framework.file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            extracted_text = "\n".join(text_parts)
        
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {framework.file_type}"
            )
        
        framework.upload_status = "text_extracted"
        framework.updated_at = datetime.utcnow()
        db.commit()
        
        return {
            "framework_id": framework.id,
            "file_name": framework.file_name,
            "file_type": framework.file_type,
            "text_content": extracted_text,
            "text_length": len(extracted_text),
            "extraction_status": "success"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        framework.upload_status = "extraction_failed"
        framework.parse_error = str(e)
        framework.updated_at = datetime.utcnow()
        db.commit()
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to extract text: {str(e)}"
        )


upload_router = router
