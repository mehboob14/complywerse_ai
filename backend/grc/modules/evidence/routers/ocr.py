from ....config import get_openai_model
import os
import base64
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Optional
from datetime import datetime
from io import BytesIO
import logging
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from pydantic import BaseModel
from PyPDF2 import PdfReader
from openai import OpenAI

from ....models import Evidence, GRCUser, get_db
from ....routers.auth_router import require_auth, get_user_tenants

router = APIRouter(prefix="/ocr", tags=["Evidence - OCR"])
logger = logging.getLogger(__name__)

AI_INTEGRATIONS_OPENAI_API_KEY = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
AI_INTEGRATIONS_OPENAI_BASE_URL = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")

CURRENT_EVIDENCE_UPLOAD_DIR = Path(__file__).resolve().parents[3] / "uploads" / "evidence"
LEGACY_EVIDENCE_UPLOAD_DIR = Path(__file__).resolve().parents[4] / "uploads" / "evidence"
BACKEND_DIR = Path(__file__).resolve().parents[4]
WORKSPACE_DIR = Path(__file__).resolve().parents[5]

IMAGE_FILE_TYPES = {"png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff", "webp"}
DOCX_FILE_TYPES = {"docx"}
EXCEL_FILE_TYPES = {"xls", "xlsx"}
TEXT_FILE_TYPES = {"txt", "text", "log", "md", "csv", "json", "xml", "html", "htm", "rtf"}
PROCESSABLE_FILE_TYPES = {"pdf"} | IMAGE_FILE_TYPES | DOCX_FILE_TYPES | TEXT_FILE_TYPES | EXCEL_FILE_TYPES

IMAGE_MEDIA_TYPES = {
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "tif": "image/tiff",
    "tiff": "image/tiff",
    "webp": "image/webp",
}

OCR_PROMPT = "Extract all text from this document/image. Return only the extracted text, preserving the original structure and formatting as much as possible."


class BatchProcessRequest(BaseModel):
    evidence_ids: List[int]


class OCRContentResponse(BaseModel):
    evidence_id: int
    ocr_content: Optional[str]
    ocr_status: str
    ocr_processed_at: Optional[str]


class OCRProcessResponse(BaseModel):
    evidence_id: int
    status: str
    extracted_text: Optional[str]
    message: Optional[str]


class BatchProcessResponse(BaseModel):
    total: int
    processed: int
    failed: int
    results: List[OCRProcessResponse]


def validate_evidence_access(user: GRCUser, evidence: Evidence, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if evidence.tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this evidence"
        )


def _is_placeholder_api_key(value: Optional[str]) -> bool:
    if not value:
        return True
    trimmed = value.strip()
    return trimmed.startswith("_DUMMY") or trimmed == "your-api-key-here"


def get_openai_client() -> OpenAI:
    integration_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
    fallback_key = os.environ.get("OPENAI_API_KEY")
    if integration_key and not _is_placeholder_api_key(integration_key):
        api_key = integration_key
    elif fallback_key and not _is_placeholder_api_key(fallback_key):
        api_key = fallback_key
    else:
        api_key = integration_key or fallback_key

    base_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL") or os.environ.get("OPENAI_BASE_URL")
    is_modelfarm = bool(base_url and "modelfarm" in base_url.lower())
    if not api_key and not is_modelfarm:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI features unavailable. OpenAI API key not configured."
        )
    if api_key and not is_modelfarm and _is_placeholder_api_key(api_key):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI features unavailable. OpenAI API key not configured."
        )
    return OpenAI(
        api_key=api_key,
        base_url=base_url
    )


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted_text = "\n\n".join(text_parts)
        return extracted_text.strip() if extracted_text.strip() else None
    except Exception:
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    try:
        from docx import Document

        document = Document(file_path)
        text_parts: List[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        extracted_text = "\n".join(text_parts).strip()
        if extracted_text:
            return extracted_text
    except Exception as exc:
        logger.warning("Primary DOCX extraction failed for %s: %s", file_path, exc)

    # Fallback parser for DOCX XML in case python-docx parsing fails.
    try:
        with zipfile.ZipFile(file_path) as archive:
            xml_bytes = archive.read("word/document.xml")

        root = ET.fromstring(xml_bytes)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        text_nodes = [node.text for node in root.findall(".//w:t", namespace) if node.text]
        extracted_text = " ".join(text_nodes).strip()
        return extracted_text if extracted_text else None
    except Exception as exc:
        logger.warning("Fallback DOCX extraction failed for %s: %s", file_path, exc)
        return None


def extract_text_from_excel(file_path: str) -> Optional[str]:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, data_only=True)
        text_parts: List[str] = []

        for sheet in workbook:
            text_parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        extracted_text = "\n".join(text_parts).strip()
        return extracted_text if extracted_text else None
    except Exception as exc:
        logger.warning("Excel extraction failed for %s: %s", file_path, exc)
        return None


def extract_text_from_plain_file(file_path: str) -> Optional[str]:
    raw_text: Optional[str] = None
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding, errors="ignore") as file_handle:
                candidate = file_handle.read()
                if candidate and candidate.strip():
                    raw_text = candidate
                    break
        except Exception:
            continue

    if not raw_text:
        return None

    if "<html" in raw_text.lower() or "<body" in raw_text.lower():
        try:
            from bs4 import BeautifulSoup

            raw_text = BeautifulSoup(raw_text, "html.parser").get_text("\n")
        except Exception:
            pass

    # Light cleanup for RTF-like content.
    if raw_text.lstrip().startswith("{\\rtf"):
        raw_text = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw_text)
        raw_text = re.sub(r"\\[a-z]+\d* ?", " ", raw_text)
        raw_text = raw_text.replace("{", " ").replace("}", " ")

    extracted_text = raw_text.strip()
    return extracted_text if extracted_text else None


def extract_text_with_tesseract(file_path: str, file_type: str) -> Optional[str]:
    if file_type not in IMAGE_FILE_TYPES:
        return None

    try:
        from PIL import Image, ImageOps
        import pytesseract

        with Image.open(file_path) as image:
            normalized = ImageOps.grayscale(image)
            text = pytesseract.image_to_string(normalized)
            cleaned = text.strip() if text else ""
            return cleaned if len(cleaned) >= 10 else None
    except Exception as exc:
        logger.warning("Local OCR with Tesseract failed for %s: %s", file_path, exc)
        return None


def build_vision_data_url(file_path: str, file_type: str) -> str:
    if file_type in IMAGE_FILE_TYPES:
        try:
            from PIL import Image

            with Image.open(file_path) as image:
                normalized = image.convert("RGB")
                normalized.thumbnail((2000, 2000))

                output = BytesIO()
                normalized.save(output, format="JPEG", quality=85, optimize=True)
                base64_data = base64.standard_b64encode(output.getvalue()).decode("utf-8")
                return f"data:image/jpeg;base64,{base64_data}"
        except Exception as exc:
            logger.warning("Image normalization failed for %s: %s", file_path, exc)

    with open(file_path, "rb") as file_handle:
        base64_data = base64.standard_b64encode(file_handle.read()).decode("utf-8")

    media_type = "application/pdf" if file_type == "pdf" else IMAGE_MEDIA_TYPES.get(file_type, f"image/{file_type}")
    return f"data:{media_type};base64,{base64_data}"


def extract_text_with_vision(file_path: str, file_type: str) -> str:
    client = get_openai_client()

    data_url = build_vision_data_url(file_path, file_type)
    
    try:
        response = client.chat.completions.create(
            model=get_openai_model(),
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": OCR_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url}
                        }
                    ]
                }
            ],
            max_completion_tokens=4096
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"OpenAI API error: {error_msg}"
        )


def get_file_extension(evidence: Evidence) -> str:
    """Extract file extension from file_name or file_path."""
    if evidence.file_name:
        ext = os.path.splitext(evidence.file_name)[1].lower().strip(".")
        if ext:
            return ext
    if evidence.file_path:
        ext = os.path.splitext(evidence.file_path)[1].lower().strip(".")
        if ext:
            return ext

    if evidence.file_type:
        mime = evidence.file_type.lower().strip()
        mime_to_ext = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "text/plain": "txt",
            "text/markdown": "md",
            "text/csv": "csv",
            "application/json": "json",
            "application/xml": "xml",
            "text/xml": "xml",
            "text/html": "html",
            "application/rtf": "rtf",
            "text/rtf": "rtf",
            "image/jpeg": "jpg",
            "image/png": "png",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/webp": "webp",
        }
        if mime in mime_to_ext:
            return mime_to_ext[mime]

        if "/" in mime:
            subtype = mime.split("/")[-1]
            if subtype == "plain":
                return "txt"
            if subtype == "jpeg":
                return "jpg"
            return subtype

    return ""


def _iter_candidate_paths(evidence: Evidence) -> List[Path]:
    candidates: List[Path] = []
    raw_path = (evidence.file_path or "").strip()
    tenant_dir_name = str(evidence.tenant_id) if evidence.tenant_id else ""

    def add_candidate(path_value: Path) -> None:
        if path_value not in candidates:
            candidates.append(path_value)

    if raw_path:
        raw = Path(raw_path)
        add_candidate(raw)

        if not raw.is_absolute():
            add_candidate((Path.cwd() / raw).resolve())
            add_candidate((BACKEND_DIR / raw).resolve())
            add_candidate((WORKSPACE_DIR / raw).resolve())

    file_name_candidate = ""
    if raw_path:
        file_name_candidate = Path(raw_path).name

    upload_roots = [CURRENT_EVIDENCE_UPLOAD_DIR, LEGACY_EVIDENCE_UPLOAD_DIR]
    if file_name_candidate:
        for upload_root in upload_roots:
            if tenant_dir_name:
                add_candidate(upload_root / tenant_dir_name / file_name_candidate)
            add_candidate(upload_root / file_name_candidate)

    return candidates


def resolve_evidence_file_path(evidence: Evidence) -> Optional[str]:
    for candidate in _iter_candidate_paths(evidence):
        try:
            if candidate.exists() and candidate.is_file():
                return str(candidate)
        except OSError:
            continue
    return None


def supported_types_text() -> str:
    sorted_types = sorted(PROCESSABLE_FILE_TYPES)
    return ", ".join(ext.upper() for ext in sorted_types)


def process_evidence_ocr(evidence: Evidence, db: Session) -> OCRProcessResponse:
    resolved_file_path = resolve_evidence_file_path(evidence)
    if not resolved_file_path:
        evidence.ocr_status = "failed"
        evidence.ocr_processed_at = datetime.utcnow()
        db.commit()
        return OCRProcessResponse(
            evidence_id=evidence.id,
            status="failed",
            extracted_text=None,
            message="File not found on disk. Please re-upload this evidence file."
        )

    if evidence.file_path != resolved_file_path:
        evidence.file_path = resolved_file_path
    
    file_type = get_file_extension(evidence)
    if file_type not in PROCESSABLE_FILE_TYPES:
        evidence.ocr_status = "failed"
        evidence.ocr_processed_at = datetime.utcnow()
        db.commit()
        return OCRProcessResponse(
            evidence_id=evidence.id,
            status="failed",
            extracted_text=None,
            message=f"Unsupported file type: {file_type or 'unknown'}. Supported types: {supported_types_text()}"
        )
    
    evidence.ocr_status = "processing"
    db.commit()
    
    try:
        extracted_text = None
        
        if file_type == "pdf":
            extracted_text = extract_text_from_pdf(resolved_file_path)
            if not extracted_text:
                extracted_text = extract_text_with_vision(resolved_file_path, file_type)
        elif file_type in DOCX_FILE_TYPES:
            extracted_text = extract_text_from_docx(resolved_file_path)
        elif file_type in EXCEL_FILE_TYPES:
            extracted_text = extract_text_from_excel(resolved_file_path)
        elif file_type in TEXT_FILE_TYPES:
            extracted_text = extract_text_from_plain_file(resolved_file_path)
        elif file_type in IMAGE_FILE_TYPES:
            extracted_text = extract_text_with_tesseract(resolved_file_path, file_type)
            if not extracted_text:
                extracted_text = extract_text_with_vision(resolved_file_path, file_type)

        if not extracted_text or not extracted_text.strip():
            evidence.ocr_content = None
            evidence.ocr_status = "failed"
            evidence.ocr_processed_at = datetime.utcnow()
            db.commit()
            return OCRProcessResponse(
                evidence_id=evidence.id,
                status="failed",
                extracted_text=None,
                message="OCR did not extract readable text from this file"
            )
        
        evidence.ocr_content = extracted_text.strip()
        evidence.ocr_status = "completed"
        evidence.ocr_processed_at = datetime.utcnow()
        db.commit()
        
        return OCRProcessResponse(
            evidence_id=evidence.id,
            status="completed",
            extracted_text=extracted_text,
            message="OCR processing completed successfully"
        )
    
    except HTTPException as exc:
        evidence.ocr_status = "failed"
        evidence.ocr_processed_at = datetime.utcnow()
        db.commit()
        return OCRProcessResponse(
            evidence_id=evidence.id,
            status="failed",
            extracted_text=None,
            message=str(exc.detail)
        )
    
    except Exception as e:
        evidence.ocr_status = "failed"
        evidence.ocr_processed_at = datetime.utcnow()
        db.commit()
        logger.exception("OCR processing failed for evidence %s", evidence.id)
        return OCRProcessResponse(
            evidence_id=evidence.id,
            status="failed",
            extracted_text=None,
            message=f"OCR processing failed: {str(e)}"
        )


@router.post("/{evidence_id}/process-ocr", response_model=OCRProcessResponse)
def process_ocr(
    evidence_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    evidence = db.query(Evidence).filter(Evidence.id == evidence_id).first()
    
    if not evidence:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence not found"
        )
    
    validate_evidence_access(current_user, evidence, db)
    
    return process_evidence_ocr(evidence, db)


@router.post("/batch-process", response_model=BatchProcessResponse)
def batch_process_ocr(
    request: BatchProcessRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    results = []
    processed_count = 0
    failed_count = 0
    
    for evidence_id in request.evidence_ids:
        evidence = db.query(Evidence).filter(Evidence.id == evidence_id).first()
        
        if not evidence:
            results.append(OCRProcessResponse(
                evidence_id=evidence_id,
                status="failed",
                extracted_text=None,
                message="Evidence not found"
            ))
            failed_count += 1
            continue
        
        if evidence.tenant_id not in user_tenants:
            results.append(OCRProcessResponse(
                evidence_id=evidence_id,
                status="failed",
                extracted_text=None,
                message="Access denied"
            ))
            failed_count += 1
            continue
        
        try:
            result = process_evidence_ocr(evidence, db)
            results.append(result)
            if result.status == "completed":
                processed_count += 1
            else:
                failed_count += 1
        except HTTPException as e:
            results.append(OCRProcessResponse(
                evidence_id=evidence_id,
                status="failed",
                extracted_text=None,
                message=e.detail
            ))
            failed_count += 1
        except Exception as e:
            results.append(OCRProcessResponse(
                evidence_id=evidence_id,
                status="failed",
                extracted_text=None,
                message=str(e)
            ))
            failed_count += 1
    
    return BatchProcessResponse(
        total=len(request.evidence_ids),
        processed=processed_count,
        failed=failed_count,
        results=results
    )


@router.get("/{evidence_id}/ocr-content", response_model=OCRContentResponse)
def get_ocr_content(
    evidence_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    evidence = db.query(Evidence).filter(Evidence.id == evidence_id).first()
    
    if not evidence:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Evidence not found"
        )
    
    validate_evidence_access(current_user, evidence, db)
    
    return OCRContentResponse(
        evidence_id=evidence.id,
        ocr_content=evidence.ocr_content,
        ocr_status=evidence.ocr_status or "pending",
        ocr_processed_at=evidence.ocr_processed_at.isoformat() if evidence.ocr_processed_at else None
    )
