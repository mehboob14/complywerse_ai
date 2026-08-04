"""
Qdrant-backed vector retrieval for ComplyChat.

This module is intentionally self-contained so chatbot routing can use:
- deterministic document chunk indexing
- OpenAI embeddings
- Qdrant search with tenant isolation
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

import requests
from openai import OpenAI
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

BACKEND_DIR = Path(__file__).resolve().parents[3]
WORKSPACE_DIR = BACKEND_DIR.parent

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".htm",
    ".log",
    ".ini",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".sql",
}

DEFAULT_MAX_EXTRACT_CHARS = int(os.getenv("COMPLYCHAT_EXTRACT_MAX_CHARS") or "50000")


def _first_env(*keys: str) -> str:
    for key in keys:
        value = os.getenv(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


@dataclass
class IndexedSourceDocument:
    tenant_id: int
    source_type: str
    source_id: str
    title: str
    text: str
    description: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    updated_at: str = ""


def _coerce_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (dict, list)):
        try:
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value).strip()
    return str(value).strip()


def _normalize_whitespace(text: str) -> str:
    normalized = re.sub(r"\r\n?", "\n", text or "")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def resolve_existing_path(raw_path: Optional[str]) -> Optional[Path]:
    path_value = (raw_path or "").strip()
    if not path_value:
        return None

    raw = Path(path_value)
    candidates: List[Path] = [raw]
    if not raw.is_absolute():
        candidates.extend(
            [
                (Path.cwd() / raw).resolve(),
                (BACKEND_DIR / raw).resolve(),
                (WORKSPACE_DIR / raw).resolve(),
            ]
        )
    for candidate in candidates:
        try:
            if candidate.exists() and candidate.is_file():
                return candidate
        except OSError:
            continue
    return None


def _extract_text_from_plain(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            data = path.read_text(encoding=encoding, errors="ignore")
            if data and data.strip():
                return data
        except Exception:
            continue
    return ""


def _extract_text_from_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts: List[str] = []
        for page in reader.pages[:40]:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        return "\n\n".join(parts)
    except Exception as exc:
        logger.warning("PDF text extraction failed for %s: %s", path, exc)
        return ""


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        chunks: List[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                chunks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    chunks.append(" | ".join(cells))
        return "\n".join(chunks)
    except Exception as exc:
        logger.warning("DOCX text extraction failed for %s: %s", path, exc)
        return ""


def _extract_text_from_excel(path: Path) -> str:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
        parts: List[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Excel text extraction failed for %s: %s", path, exc)
        return ""


def _extract_text_from_legacy_excel(path: Path) -> str:
    try:
        import pandas as pd

        frames = pd.read_excel(str(path), sheet_name=None)
        parts: List[str] = []
        for sheet_name, frame in frames.items():
            parts.append(f"# Sheet: {sheet_name}")
            if frame is None:
                continue
            frame = frame.fillna("")
            for row in frame.values.tolist():
                cells = [str(cell).strip() for cell in row if str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Legacy Excel text extraction failed for %s: %s", path, exc)
        return ""


def extract_text_from_path(raw_path: Optional[str], *, max_chars: Optional[int] = None) -> str:
    resolved = resolve_existing_path(raw_path)
    if not resolved:
        return ""

    suffix = resolved.suffix.lower()
    text_content = ""
    if suffix in TEXT_EXTENSIONS:
        text_content = _extract_text_from_plain(resolved)
    elif suffix == ".pdf":
        text_content = _extract_text_from_pdf(resolved)
    elif suffix == ".docx":
        text_content = _extract_text_from_docx(resolved)
    elif suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        text_content = _extract_text_from_excel(resolved)
    elif suffix == ".xls":
        text_content = _extract_text_from_legacy_excel(resolved)
    else:
        try:
            text_content = resolved.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text_content = ""

    cleaned = _normalize_whitespace(text_content)
    effective_max = max_chars if isinstance(max_chars, int) and max_chars > 0 else DEFAULT_MAX_EXTRACT_CHARS
    return cleaned[:effective_max] if cleaned else ""


def chunk_text(text: str, *, chunk_size: int = 1800, overlap: int = 260) -> List[str]:
    normalized = _normalize_whitespace(text)
    if not normalized:
        return []
    if len(normalized) <= chunk_size:
        return [normalized]

    paragraphs = [part.strip() for part in re.split(r"\n{2,}", normalized) if part.strip()]
    chunks: List[str] = []
    current = ""
    for paragraph in paragraphs:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
            continue

        if current:
            chunks.append(current)
        if len(paragraph) <= chunk_size:
            current = paragraph
            continue

        start = 0
        while start < len(paragraph):
            end = min(start + chunk_size, len(paragraph))
            piece = paragraph[start:end].strip()
            if piece:
                chunks.append(piece)
            if end >= len(paragraph):
                break
            start = max(0, end - overlap)
        current = ""

    if current:
        chunks.append(current)

    return [chunk for chunk in chunks if chunk.strip()]


class QdrantComplyChatService:
    def __init__(self) -> None:
        self.qdrant_url = (
            _first_env(
                "QDRANT_URL",
                "QDRANT_CLUSTER_ENDPOINT",
                "CLUSTER_ENDPOINT",
                "cluster_endpoint",
                "_endpoint",
            )
            or "http://127.0.0.1:6333"
        ).rstrip("/")
        self.qdrant_api_key = _first_env("QDRANT_API_KEY", "qdrant_api_key")
        self.collection_name = _first_env("QDRANT_COMPLYCHAT_COLLECTION", "QDRANT_COLLECTION") or "complyverseai"
        self.request_timeout = float(os.getenv("QDRANT_TIMEOUT_SECONDS") or "20")
        self.embedding_model = os.getenv("COMPLYCHAT_EMBEDDING_MODEL") or "text-embedding-3-small"
        self.vector_size = int(os.getenv("COMPLYCHAT_VECTOR_SIZE") or "1536")
        self.max_chunks_per_doc = int(os.getenv("COMPLYCHAT_MAX_CHUNKS_PER_DOC") or "40")
        self.vector_name = _first_env("COMPLYCHAT_QDRANT_VECTOR_NAME", "QDRANT_VECTOR_NAME") or "abstract"
        self.sparse_vector_name = _first_env("COMPLYCHAT_QDRANT_SPARSE_VECTOR_NAME") or "documents_sparse"
        self.document_id_field = _first_env("COMPLYCHAT_QDRANT_DOCUMENT_ID_FIELD") or "document_id"
        self.legacy_document_id_field = _first_env("COMPLYCHAT_QDRANT_LEGACY_DOC_ID_FIELD") or "docuemnt_id"
        self._named_vector_enabled = True

        openai_api_key = os.getenv("OPENAI_API_KEY") or os.getenv("AI_INTEGRATIONS_OPENAI_API_KEY")
        openai_base_url = os.getenv("OPENAI_BASE_URL") or os.getenv("AI_INTEGRATIONS_OPENAI_BASE_URL")
        if openai_api_key or openai_base_url:
            self.openai_client = OpenAI(api_key=openai_api_key, base_url=openai_base_url)
        else:
            self.openai_client = None

    @property
    def is_available(self) -> bool:
        return bool(self.qdrant_url and self.openai_client is not None)

    def _headers(self) -> Dict[str, str]:
        headers = {"Content-Type": "application/json"}
        if self.qdrant_api_key:
            headers["api-key"] = self.qdrant_api_key
        return headers

    def _request(
        self,
        method: str,
        path: str,
        *,
        payload: Optional[Dict[str, Any]] = None,
        accept_404: bool = False,
    ) -> Dict[str, Any]:
        url = f"{self.qdrant_url}{path}"
        response = requests.request(
            method=method.upper(),
            url=url,
            headers=self._headers(),
            json=payload,
            timeout=self.request_timeout,
        )
        if accept_404 and response.status_code == 404:
            return {}
        if response.status_code >= 400:
            raise RuntimeError(f"Qdrant request failed ({response.status_code}): {response.text[:400]}")
        if not response.text.strip():
            return {}
        return response.json()

    def ensure_collection(self) -> None:
        if not self.is_available:
            raise RuntimeError("Qdrant service is not configured.")

        existing = self._request("GET", f"/collections/{self.collection_name}", accept_404=True)
        if existing:
            self._detect_vector_mode(existing)
            self._ensure_payload_indexes()
            return

        create_payloads: List[Dict[str, Any]] = [
            {
                "vectors": {
                    self.vector_name: {
                        "size": self.vector_size,
                        "distance": "Cosine",
                        "on_disk": True,
                        "datatype": "float32",
                        "hnsw_config": {
                            "m": 0,
                            "payload_m": 24,
                            "ef_construct": 256,
                        },
                    }
                },
                "sparse_vectors": {
                    self.sparse_vector_name: {
                        "index": {"on_disk": True},
                    }
                },
            },
            {
                "vectors": {
                    self.vector_name: {
                        "size": self.vector_size,
                        "distance": "Cosine",
                    }
                }
            },
            {
                "vectors": {
                    "size": self.vector_size,
                    "distance": "Cosine",
                }
            },
        ]

        last_error: Optional[Exception] = None
        for payload in create_payloads:
            try:
                self._request("PUT", f"/collections/{self.collection_name}", payload=payload)
                created = self._request("GET", f"/collections/{self.collection_name}", accept_404=True)
                self._detect_vector_mode(created)
                self._ensure_payload_indexes()
                return
            except Exception as exc:
                last_error = exc
                continue
        raise RuntimeError(f"Unable to create Qdrant collection '{self.collection_name}': {last_error}")

    def _detect_vector_mode(self, collection_info: Dict[str, Any]) -> None:
        params = (((collection_info.get("result") or {}).get("config") or {}).get("params")) or {}
        vectors = params.get("vectors")
        if isinstance(vectors, dict) and "size" in vectors and "distance" in vectors:
            self._named_vector_enabled = False
            return
        self._named_vector_enabled = True

    def _ensure_payload_indexes(self) -> None:
        # Add both correct and legacy-typo fields so existing payload-index requests remain compatible.
        fields = [
            (self.document_id_field, {"type": "keyword", "on_disk": False, "is_tenant": True, "is_principal": True}),
            (self.legacy_document_id_field, {"type": "keyword", "on_disk": False, "is_tenant": True, "is_principal": True}),
            ("tenant_id", "integer"),
            ("source_type", "keyword"),
            ("source_id", "keyword"),
        ]
        for field_name, field_schema in fields:
            self._ensure_payload_index(field_name=field_name, field_schema=field_schema)

    def _ensure_payload_index(self, *, field_name: str, field_schema: Any) -> None:
        payload = {"field_name": field_name, "field_schema": field_schema}
        fallback_payload = {
            "field_name": field_name,
            "field_schema": field_schema.get("type") if isinstance(field_schema, dict) else field_schema,
        }
        paths = [
            f"/collections/{self.collection_name}/payload_indexes",
            f"/collections/{self.collection_name}/index",
        ]
        for path in paths:
            try:
                self._request("PUT", path, payload=payload)
                return
            except Exception:
                try:
                    self._request("PUT", path, payload=fallback_payload)
                    return
                except Exception:
                    continue
        logger.debug("Skipping payload index for '%s' (not supported by current Qdrant endpoint).", field_name)

    def embed_texts(self, texts: Sequence[str]) -> List[List[float]]:
        if not self.openai_client:
            raise RuntimeError("OpenAI client is not configured for embeddings.")
        if not texts:
            return []

        import time
        from ...services.ai_usage import record_provider_attempt, usage_scope
        started_at = time.perf_counter()
        try:
            with usage_scope(module_key="complychat", feature_key="embedding"):
                response = self.openai_client.embeddings.create(
                    model=self.embedding_model,
                    input=list(texts),
                )
                record_provider_attempt(
                    response=response,
                    requested_model=self.embedding_model,
                    provider="openai",
                    api_family="embeddings",
                    started_at=started_at,
                )
        except Exception as exc:
            record_provider_attempt(
                error=exc,
                requested_model=self.embedding_model,
                provider="openai",
                api_family="embeddings",
                started_at=started_at,
            )
            raise
        return [list(item.embedding) for item in response.data]

    def _point_id(self, tenant_id: int, source_type: str, source_id: str, chunk_index: int) -> str:
        raw = f"{tenant_id}|{source_type}|{source_id}|{chunk_index}"
        return hashlib.sha1(raw.encode("utf-8")).hexdigest()

    def delete_source_points(self, tenant_id: int, source_type: str, source_id: str) -> None:
        self._request(
            "POST",
            f"/collections/{self.collection_name}/points/delete?wait=true",
            payload={
                "filter": {
                    "must": [
                        {"key": "tenant_id", "match": {"value": tenant_id}},
                        {"key": "source_type", "match": {"value": source_type}},
                        {"key": "source_id", "match": {"value": source_id}},
                    ]
                }
            },
        )

    def upsert_documents(self, documents: Sequence[IndexedSourceDocument]) -> int:
        if not documents:
            return 0
        self.ensure_collection()

        payload_rows: List[Dict[str, Any]] = []
        chunk_texts: List[str] = []

        for doc in documents:
            content = _normalize_whitespace(doc.text)
            if not content:
                continue
            chunks = chunk_text(content, chunk_size=1800, overlap=260)[: self.max_chunks_per_doc]
            for index, chunk in enumerate(chunks):
                point_id = self._point_id(doc.tenant_id, doc.source_type, doc.source_id, index)
                payload_rows.append(
                    {
                        "id": point_id,
                        self.document_id_field: f"{doc.source_type}:{doc.source_id}",
                        self.legacy_document_id_field: f"{doc.source_type}:{doc.source_id}",
                        "tenant_id": doc.tenant_id,
                        "source_type": doc.source_type,
                        "source_id": doc.source_id,
                        "title": doc.title[:400],
                        "description": doc.description[:800],
                        "snippet": chunk[:700],
                        "text": chunk,
                        "metadata": doc.metadata or {},
                        "updated_at": doc.updated_at or "",
                        "chunk_index": index,
                    }
                )
                chunk_texts.append(chunk)

        if not chunk_texts:
            return 0

        vectors = self.embed_texts(chunk_texts)
        points: List[Dict[str, Any]] = []
        for row, vector in zip(payload_rows, vectors):
            point_vector: Any
            if self._named_vector_enabled:
                point_vector = {self.vector_name: vector}
            else:
                point_vector = vector
            points.append({"id": row["id"], "vector": point_vector, "payload": row})

        batch_size = 64
        for start in range(0, len(points), batch_size):
            batch = points[start : start + batch_size]
            self._request(
                "PUT",
                f"/collections/{self.collection_name}/points?wait=true",
                payload={"points": batch},
            )
        return len(points)

    def search(
        self,
        *,
        tenant_id: int,
        query: str,
        limit: int = 8,
        source_types: Optional[Sequence[str]] = None,
    ) -> List[Dict[str, Any]]:
        self.ensure_collection()
        embeddings = self.embed_texts([query])
        if not embeddings:
            return []
        search_vector: Any
        if self._named_vector_enabled:
            search_vector = {"name": self.vector_name, "vector": embeddings[0]}
        else:
            search_vector = embeddings[0]

        must_filters: List[Dict[str, Any]] = [{"key": "tenant_id", "match": {"value": tenant_id}}]
        if source_types:
            must_filters.append({"key": "source_type", "match": {"any": list(source_types)}})

        try:
            response = self._request(
                "POST",
                f"/collections/{self.collection_name}/points/search",
                payload={
                    "vector": search_vector,
                    "limit": max(1, min(limit, 20)),
                    "with_payload": True,
                    "filter": {"must": must_filters},
                },
            )
        except Exception:
            # Older Qdrant versions may not support match.any; retry without source type filter.
            response = self._request(
                "POST",
                f"/collections/{self.collection_name}/points/search",
                payload={
                    "vector": search_vector,
                    "limit": max(1, min(limit, 20)),
                    "with_payload": True,
                    "filter": {"must": [{"key": "tenant_id", "match": {"value": tenant_id}}]},
                },
            )

        return response.get("result") or []

    def health(self) -> Dict[str, Any]:
        if not self.is_available:
            return {"available": False, "reason": "missing_qdrant_or_openai_configuration"}
        try:
            self.ensure_collection()
            info = self._request("GET", f"/collections/{self.collection_name}")
            status_value = (((info.get("result") or {}).get("status")) or "ok")
            return {"available": True, "collection": self.collection_name, "status": status_value}
        except Exception as exc:
            return {"available": False, "reason": str(exc)}

    def count_points(self, *, tenant_id: Optional[int] = None) -> int:
        self.ensure_collection()
        payload: Dict[str, Any] = {"exact": False}
        if tenant_id is not None:
            payload["filter"] = {"must": [{"key": "tenant_id", "match": {"value": int(tenant_id)}}]}
        result = self._request(
            "POST",
            f"/collections/{self.collection_name}/points/count",
            payload=payload,
        )
        return int((result.get("result") or {}).get("count") or 0)
