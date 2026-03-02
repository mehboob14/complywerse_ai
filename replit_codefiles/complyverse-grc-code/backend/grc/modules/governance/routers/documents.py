from typing import List, Optional
from datetime import datetime, timedelta
import os
import uuid
import json
import html
from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import or_, and_
from pydantic import BaseModel
from openai import OpenAI

from ....models import (
    GovernanceDocument, GovernanceDocumentVersion, DocumentReviewer,
    DocumentApprovalStep, DocumentAuditLog, GRCUser, Tenant, PolicyStatement, 
    InternalControl, ParsedFrameworkControl, UploadedFramework, PolicyReviewHistory, get_db
)
from ....routers.auth_router import require_auth, get_user_tenants, get_user_primary_tenant

AI_INTEGRATIONS_OPENAI_API_KEY = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
AI_INTEGRATIONS_OPENAI_BASE_URL = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")

GOVERNANCE_UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "uploads", "governance")
os.makedirs(GOVERNANCE_UPLOAD_DIR, exist_ok=True)

ALLOWED_FILE_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'doc': 'application/msword',
    'xls': 'application/vnd.ms-excel'
}

router = APIRouter(prefix="/documents", tags=["Governance - Documents"])


class GovernanceDocumentCreate(BaseModel):
    document_code: Optional[str] = None
    title: str
    description: Optional[str] = None
    content: Optional[str] = None
    doc_type: str
    doc_sub_type: Optional[str] = None
    classification: str = "internal"
    parent_document_id: Optional[int] = None
    owner_id: Optional[int] = None
    author_id: Optional[int] = None
    department_id: Optional[int] = None
    effective_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None
    review_cycle_months: int = 12
    next_review_date: Optional[datetime] = None
    regulatory_scope: Optional[List[str]] = []
    framework_ids: Optional[List[int]] = []
    tags: Optional[List[str]] = []


class GovernanceDocumentUpdate(BaseModel):
    document_code: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    content: Optional[str] = None
    doc_type: Optional[str] = None
    doc_sub_type: Optional[str] = None
    classification: Optional[str] = None
    parent_document_id: Optional[int] = None
    status: Optional[str] = None
    owner_id: Optional[int] = None
    author_id: Optional[int] = None
    department_id: Optional[int] = None
    effective_date: Optional[datetime] = None
    expiry_date: Optional[datetime] = None
    review_cycle_months: Optional[int] = None
    next_review_date: Optional[datetime] = None
    regulatory_scope: Optional[List[str]] = None
    framework_ids: Optional[List[int]] = None
    tags: Optional[List[str]] = None


class BulkStatusUpdate(BaseModel):
    document_ids: List[int]
    status: str


class BulkArchive(BaseModel):
    document_ids: List[int]


def validate_tenant_access(user: GRCUser, tenant_id: int, db: Session) -> None:
    user_tenants = get_user_tenants(user, db)
    if tenant_id not in user_tenants:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access denied to this tenant's data"
        )


def create_audit_log(
    db: Session,
    document_id: int,
    tenant_id: int,
    user_id: int,
    action: str,
    action_details: Optional[str] = None,
    field_changed: Optional[str] = None,
    old_value: Optional[str] = None,
    new_value: Optional[str] = None
) -> DocumentAuditLog:
    audit_log = DocumentAuditLog(
        document_id=document_id,
        tenant_id=tenant_id,
        action=action,
        action_details=action_details,
        field_changed=field_changed,
        old_value=old_value,
        new_value=new_value,
        performed_by=user_id,
        performed_at=datetime.utcnow()
    )
    db.add(audit_log)
    return audit_log


def serialize_document(doc: GovernanceDocument, db: Session = None) -> dict:
    result = {
        "id": doc.id,
        "tenant_id": doc.tenant_id,
        "document_code": doc.document_code,
        "title": doc.title,
        "description": doc.description,
        "content": doc.content,
        "file_name": doc.file_name,
        "file_size": doc.file_size,
        "file_type": doc.file_type,
        "has_file": doc.file_path is not None,
        "doc_type": doc.doc_type,
        "doc_sub_type": doc.doc_sub_type,
        "classification": doc.classification,
        "parent_document_id": doc.parent_document_id,
        "current_version": doc.current_version,
        "status": doc.status,
        "owner_id": doc.owner_id,
        "owner_name": doc.owner.display_name if doc.owner else None,
        "author_id": doc.author_id,
        "author_name": doc.author.display_name if doc.author else None,
        "department_id": doc.department_id,
        "effective_date": doc.effective_date.isoformat() if doc.effective_date else None,
        "expiry_date": doc.expiry_date.isoformat() if doc.expiry_date else None,
        "review_cycle_months": doc.review_cycle_months,
        "next_review_date": doc.next_review_date.isoformat() if doc.next_review_date else None,
        "last_reviewed_at": doc.last_reviewed_at.isoformat() if doc.last_reviewed_at else None,
        "last_reviewed_by": doc.last_reviewed_by,
        "regulatory_scope": doc.regulatory_scope or [],
        "framework_ids": doc.framework_ids or [],
        "tags": doc.tags or [],
        "approved_by": doc.approved_by,
        "approved_at": doc.approved_at.isoformat() if doc.approved_at else None,
        "published_by": doc.published_by,
        "published_at": doc.published_at.isoformat() if doc.published_at else None,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
    }
    
    if db:
        policy_count = db.query(PolicyStatement).filter(
            PolicyStatement.document_id == doc.id
        ).count()
        result["policy_statement_count"] = policy_count
    else:
        result["policy_statement_count"] = len(doc.policy_statements) if hasattr(doc, 'policy_statements') and doc.policy_statements else 0
    
    if hasattr(doc, 'workflow_instance') and doc.workflow_instance:
        wf = doc.workflow_instance
        result["workflow_instance"] = {
            "id": wf.id,
            "template_id": wf.template_id,
            "template_name": wf.template.name if wf.template else None,
            "current_step_id": wf.current_step_id,
            "current_step_sequence": wf.current_step_sequence,
            "current_step_name": wf.current_step.name if wf.current_step else None,
            "status": wf.status,
            "started_at": wf.started_at.isoformat() if wf.started_at else None,
            "completed_at": wf.completed_at.isoformat() if wf.completed_at else None,
        }
    else:
        result["workflow_instance"] = None
    
    return result


@router.get("")
def list_documents(
    tenant_id: Optional[int] = None,
    doc_type: Optional[str] = None,
    status_filter: Optional[str] = Query(None, alias="status"),
    owner_id: Optional[int] = None,
    classification: Optional[str] = None,
    framework_id: Optional[int] = None,
    search: Optional[str] = None,
    date_from: Optional[datetime] = None,
    date_to: Optional[datetime] = None,
    sort_by: str = "created_at",
    sort_order: str = "desc",
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return {"items": [], "total": 0, "skip": skip, "limit": limit}
    
    query = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner)
    ).filter(GovernanceDocument.tenant_id.in_(user_tenants))
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
        query = query.filter(GovernanceDocument.tenant_id == tenant_id)
    if doc_type:
        query = query.filter(GovernanceDocument.doc_type == doc_type)
    if status_filter:
        query = query.filter(GovernanceDocument.status == status_filter)
    if owner_id:
        query = query.filter(GovernanceDocument.owner_id == owner_id)
    if classification:
        query = query.filter(GovernanceDocument.classification == classification)
    if framework_id:
        query = query.filter(GovernanceDocument.framework_ids.contains([framework_id]))
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                GovernanceDocument.title.ilike(search_term),
                GovernanceDocument.description.ilike(search_term),
                GovernanceDocument.document_code.ilike(search_term)
            )
        )
    if date_from:
        query = query.filter(GovernanceDocument.created_at >= date_from)
    if date_to:
        query = query.filter(GovernanceDocument.created_at <= date_to)
    
    total = query.count()
    
    sort_column = getattr(GovernanceDocument, sort_by, GovernanceDocument.created_at)
    if sort_order == "asc":
        query = query.order_by(sort_column.asc())
    else:
        query = query.order_by(sort_column.desc())
    
    documents = query.offset(skip).limit(limit).all()
    
    return {
        "items": [serialize_document(doc, db) for doc in documents],
        "total": total,
        "skip": skip,
        "limit": limit
    }


@router.post("", status_code=status.HTTP_201_CREATED)
def create_document(
    document: GovernanceDocumentCreate,
    tenant_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
    else:
        tenant_id = get_user_primary_tenant(current_user, db)
        if not tenant_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User is not assigned to any tenant"
            )
    
    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()
    if not tenant:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Tenant not found"
        )
    
    if document.parent_document_id:
        parent = db.query(GovernanceDocument).filter(
            GovernanceDocument.id == document.parent_document_id,
            GovernanceDocument.tenant_id == tenant_id
        ).first()
        if not parent:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Parent document not found"
            )
    
    db_document = GovernanceDocument(
        tenant_id=tenant_id,
        document_code=document.document_code,
        title=document.title,
        description=document.description,
        content=document.content,
        doc_type=document.doc_type,
        doc_sub_type=document.doc_sub_type,
        classification=document.classification,
        parent_document_id=document.parent_document_id,
        owner_id=document.owner_id or current_user.id,
        author_id=document.author_id or current_user.id,
        department_id=document.department_id,
        effective_date=document.effective_date,
        expiry_date=document.expiry_date,
        review_cycle_months=document.review_cycle_months,
        next_review_date=document.next_review_date,
        regulatory_scope=document.regulatory_scope or [],
        framework_ids=document.framework_ids or [],
        tags=document.tags or [],
        status="draft",
        current_version="1.0"
    )
    db.add(db_document)
    db.flush()
    
    create_audit_log(
        db=db,
        document_id=db_document.id,
        tenant_id=tenant_id,
        user_id=current_user.id,
        action="created",
        action_details=f"Document '{document.title}' created"
    )
    
    db.commit()
    db.refresh(db_document)
    
    return serialize_document(db_document, db)


@router.get("/policies")
def get_policies(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="policy", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/standards")
def get_standards(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="standard", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/procedures")
def get_procedures(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="procedure", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/guidelines")
def get_guidelines(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="guideline", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/charters")
def get_charters(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="charter", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/frameworks")
def get_framework_documents(
    tenant_id: Optional[int] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    return list_documents(
        tenant_id=tenant_id, doc_type="framework", status_filter=None,
        owner_id=None, classification=None, framework_id=None,
        search=None, date_from=None, date_to=None,
        sort_by="created_at", sort_order="desc",
        skip=skip, limit=limit, db=db, current_user=current_user
    )


@router.get("/hierarchy")
def get_document_hierarchy(
    tenant_id: Optional[int] = None,
    parent_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    if not user_tenants:
        return []
    
    query = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner)
    ).filter(GovernanceDocument.tenant_id.in_(user_tenants))
    
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
        query = query.filter(GovernanceDocument.tenant_id == tenant_id)
    
    if parent_id is None:
        query = query.filter(GovernanceDocument.parent_document_id.is_(None))
    else:
        query = query.filter(GovernanceDocument.parent_document_id == parent_id)
    
    documents = query.order_by(GovernanceDocument.doc_type, GovernanceDocument.title).all()
    
    def build_hierarchy(doc):
        result = serialize_document(doc, db)
        children = db.query(GovernanceDocument).filter(
            GovernanceDocument.parent_document_id == doc.id
        ).all()
        result["children"] = [build_hierarchy(child) for child in children]
        return result
    
    return [build_hierarchy(doc) for doc in documents]


@router.get("/reviews/upcoming")
def get_upcoming_reviews(
    days: int = 90,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    now = datetime.utcnow()
    cutoff = now + timedelta(days=days)
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.tenant_id.in_(user_tenants),
        GovernanceDocument.next_review_date != None,
        GovernanceDocument.next_review_date <= cutoff,
        GovernanceDocument.status != "retired"
    ).order_by(GovernanceDocument.next_review_date.asc()).all()
    
    results = []
    for doc in documents:
        is_overdue = doc.next_review_date and doc.next_review_date < now
        days_until = (doc.next_review_date - now).days if doc.next_review_date else None
        
        owner_name = None
        if doc.owner_id:
            owner = db.query(GRCUser).filter(GRCUser.id == doc.owner_id).first()
            if owner:
                owner_name = owner.display_name
        
        results.append({
            "document_id": doc.id,
            "title": doc.title,
            "doc_type": doc.doc_type,
            "status": doc.status,
            "review_cycle_months": doc.review_cycle_months,
            "next_review_date": doc.next_review_date.isoformat() if doc.next_review_date else None,
            "is_overdue": is_overdue,
            "days_until_review": days_until,
            "owner_name": owner_name
        })
    
    overdue_count = sum(1 for r in results if r["is_overdue"])
    upcoming_count = len(results) - overdue_count
    
    return {
        "total": len(results),
        "overdue_count": overdue_count,
        "upcoming_count": upcoming_count,
        "reviews": results
    }


@router.get("/{document_id}")
def get_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).options(
        joinedload(GovernanceDocument.owner),
        joinedload(GovernanceDocument.author),
        joinedload(GovernanceDocument.approver),
        joinedload(GovernanceDocument.versions),
        joinedload(GovernanceDocument.reviewers).joinedload(DocumentReviewer.user),
        joinedload(GovernanceDocument.approval_steps).joinedload(DocumentApprovalStep.approver),
        joinedload(GovernanceDocument.parent_document)
    ).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    result = serialize_document(document, db)
    
    result["versions"] = [
        {
            "id": v.id,
            "version_number": v.version_number,
            "change_type": v.change_type,
            "title": v.title,
            "change_summary": v.change_summary,
            "change_reason": v.change_reason,
            "status": v.status,
            "created_at": v.created_at.isoformat() if v.created_at else None,
            "created_by": v.created_by,
            "creator_name": v.creator.display_name if v.creator else None,
            "approved_by": v.approved_by,
            "approved_at": v.approved_at.isoformat() if v.approved_at else None,
        }
        for v in document.versions
    ]
    
    result["reviewers"] = [
        {
            "id": r.id,
            "user_id": r.user_id,
            "user_name": r.user.display_name if r.user else None,
            "role_type": r.role_type,
            "sequence": r.sequence,
            "is_required": r.is_required,
            "notify_on_update": r.notify_on_update,
            "notify_on_expiry": r.notify_on_expiry,
            "assigned_at": r.assigned_at.isoformat() if r.assigned_at else None,
        }
        for r in document.reviewers
    ]
    
    result["approval_steps"] = [
        {
            "id": s.id,
            "step_sequence": s.step_sequence,
            "step_name": s.step_name,
            "approval_type": s.approval_type,
            "approver_id": s.approver_id,
            "approver_name": s.approver.display_name if s.approver else None,
            "approver_role": s.approver_role,
            "status": s.status,
            "requested_at": s.requested_at.isoformat() if s.requested_at else None,
            "due_date": s.due_date.isoformat() if s.due_date else None,
            "completed_at": s.completed_at.isoformat() if s.completed_at else None,
            "comments": s.comments,
        }
        for s in document.approval_steps
    ]
    
    result["parent_document"] = None
    if document.parent_document:
        result["parent_document"] = {
            "id": document.parent_document.id,
            "title": document.parent_document.title,
            "doc_type": document.parent_document.doc_type,
        }
    
    child_documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.parent_document_id == document_id
    ).all()
    result["child_documents"] = [
        {
            "id": c.id,
            "title": c.title,
            "doc_type": c.doc_type,
            "status": c.status,
        }
        for c in child_documents
    ]
    
    return result


@router.put("/{document_id}")
def update_document(
    document_id: int,
    document_update: GovernanceDocumentUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    update_data = document_update.model_dump(exclude_unset=True)
    changed_fields = []
    
    for field, new_value in update_data.items():
        old_value = getattr(document, field)
        if old_value != new_value:
            changed_fields.append({
                "field": field,
                "old_value": str(old_value) if old_value else None,
                "new_value": str(new_value) if new_value else None
            })
            setattr(document, field, new_value)
    
    document.updated_at = datetime.utcnow()
    
    if changed_fields:
        create_audit_log(
            db=db,
            document_id=document_id,
            tenant_id=document.tenant_id,
            user_id=current_user.id,
            action="updated",
            action_details=f"Updated fields: {', '.join([c['field'] for c in changed_fields])}"
        )
    
    db.commit()
    db.refresh(document)
    
    return serialize_document(document, db)


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    child_count = db.query(GovernanceDocument).filter(
        GovernanceDocument.parent_document_id == document_id
    ).count()
    
    if child_count > 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot delete document with {child_count} child documents. Delete or reassign child documents first."
        )
    
    db.query(InternalControl).filter(InternalControl.source_document_id == document_id).delete()
    
    db.delete(document)
    db.commit()
    
    return None


@router.get("/{document_id}/review-history")
def get_review_history(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    reviews = db.query(PolicyReviewHistory).filter(
        PolicyReviewHistory.document_id == document_id,
        PolicyReviewHistory.tenant_id.in_(user_tenants)
    ).order_by(PolicyReviewHistory.created_at.desc()).all()
    
    return [{
        "id": r.id,
        "review_type": r.review_type,
        "review_status": r.review_status,
        "scheduled_date": r.scheduled_date.isoformat() if r.scheduled_date else None,
        "started_at": r.started_at.isoformat() if r.started_at else None,
        "completed_at": r.completed_at.isoformat() if r.completed_at else None,
        "reviewer_id": r.reviewer_id,
        "review_notes": r.review_notes,
        "changes_made": r.changes_made,
        "outcome": r.outcome,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    } for r in reviews]


class StartReviewRequest(BaseModel):
    review_type: str = "periodic"


class CompleteReviewRequest(BaseModel):
    review_notes: str = ""
    changes_made: str = ""
    outcome: str = "no_changes"


@router.post("/{document_id}/start-review")
def start_review(
    document_id: int,
    request_body: Optional[StartReviewRequest] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    review_type = request_body.review_type if request_body else "periodic"
    
    review = PolicyReviewHistory(
        tenant_id=document.tenant_id,
        document_id=document_id,
        review_type=review_type,
        review_status="in_progress",
        scheduled_date=document.next_review_date,
        started_at=datetime.utcnow(),
        reviewer_id=current_user.id,
    )
    db.add(review)
    db.commit()
    db.refresh(review)
    
    return {
        "id": review.id,
        "review_type": review.review_type,
        "review_status": review.review_status,
        "scheduled_date": review.scheduled_date.isoformat() if review.scheduled_date else None,
        "started_at": review.started_at.isoformat() if review.started_at else None,
        "reviewer_id": review.reviewer_id,
        "message": "Review started successfully"
    }


@router.post("/{document_id}/complete-review")
def complete_review(
    document_id: int,
    request_body: CompleteReviewRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    review = db.query(PolicyReviewHistory).filter(
        PolicyReviewHistory.document_id == document_id,
        PolicyReviewHistory.review_status == "in_progress",
        PolicyReviewHistory.tenant_id.in_(user_tenants)
    ).order_by(PolicyReviewHistory.created_at.desc()).first()
    
    if not review:
        review = PolicyReviewHistory(
            tenant_id=document.tenant_id,
            document_id=document_id,
            review_type="periodic",
            review_status="in_progress",
            started_at=datetime.utcnow(),
            reviewer_id=current_user.id,
        )
        db.add(review)
        db.flush()
    
    now = datetime.utcnow()
    review.review_status = "completed"
    review.completed_at = now
    review.review_notes = request_body.review_notes
    review.changes_made = request_body.changes_made
    review.outcome = request_body.outcome
    review.reviewer_id = current_user.id
    
    document.last_reviewed_at = now
    document.last_reviewed_by = current_user.id
    if document.review_cycle_months:
        from dateutil.relativedelta import relativedelta
        document.next_review_date = now + relativedelta(months=document.review_cycle_months)
    
    document.updated_at = now
    
    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="review_completed",
        action_details=f"Review completed with outcome: {request_body.outcome}"
    )
    
    db.commit()
    db.refresh(review)
    db.refresh(document)
    
    return {
        "id": review.id,
        "review_status": review.review_status,
        "completed_at": review.completed_at.isoformat() if review.completed_at else None,
        "outcome": review.outcome,
        "next_review_date": document.next_review_date.isoformat() if document.next_review_date else None,
        "message": "Review completed successfully"
    }


@router.post("/bulk-update-status")
def bulk_update_status(
    bulk_update: BulkStatusUpdate,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    valid_statuses = ["draft", "pending_review", "pending_approval", "approved", "published", "expired", "archived", "exception_applied"]
    if bulk_update.status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}"
        )
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.id.in_(bulk_update.document_ids),
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found"
        )
    
    updated_count = 0
    for doc in documents:
        old_status = doc.status
        doc.status = bulk_update.status
        doc.updated_at = datetime.utcnow()
        
        create_audit_log(
            db=db,
            document_id=doc.id,
            tenant_id=doc.tenant_id,
            user_id=current_user.id,
            action="status_changed",
            action_details=f"Status changed from '{old_status}' to '{bulk_update.status}'",
            field_changed="status",
            old_value=old_status,
            new_value=bulk_update.status
        )
        updated_count += 1
    
    db.commit()
    
    return {
        "message": f"Successfully updated {updated_count} documents",
        "updated_count": updated_count
    }


@router.post("/bulk-archive")
def bulk_archive(
    bulk_archive_request: BulkArchive,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    documents = db.query(GovernanceDocument).filter(
        GovernanceDocument.id.in_(bulk_archive_request.document_ids),
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).all()
    
    if not documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found"
        )
    
    archived_count = 0
    for doc in documents:
        old_status = doc.status
        doc.status = "archived"
        doc.updated_at = datetime.utcnow()
        
        create_audit_log(
            db=db,
            document_id=doc.id,
            tenant_id=doc.tenant_id,
            user_id=current_user.id,
            action="archived",
            action_details=f"Document archived (previous status: '{old_status}')",
            field_changed="status",
            old_value=old_status,
            new_value="archived"
        )
        archived_count += 1
    
    db.commit()
    
    return {
        "message": f"Successfully archived {archived_count} documents",
        "archived_count": archived_count
    }


@router.put("/{document_id}/status")
def update_document_status(
    document_id: int,
    body: dict,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    new_status = body.get("status")
    valid_statuses = ["draft", "pending_review", "pending_approval", "approved", "published", "expired", "archived"]
    if new_status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}"
        )

    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )

    old_status = document.status
    document.status = new_status
    document.updated_at = datetime.utcnow()

    if new_status == "published":
        document.published_by = current_user.id
        document.published_at = datetime.utcnow()

    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="status_changed",
        action_details=f"Status changed from '{old_status}' to '{new_status}' by {current_user.display_name}",
        field_changed="status",
        old_value=old_status,
        new_value=new_status
    )

    db.commit()
    db.refresh(document)

    return serialize_document(document, db)


@router.post("/{document_id}/publish")
def publish_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Publish an approved document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if document.status != "approved":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only approved documents can be published. Current status: {document.status}"
        )
    
    old_status = document.status
    document.status = "published"
    document.published_by = current_user.id
    document.published_at = datetime.utcnow()
    document.updated_at = datetime.utcnow()
    
    create_audit_log(
        db=db,
        document_id=document_id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="published",
        action_details=f"Document published by {current_user.display_name}",
        field_changed="status",
        old_value=old_status,
        new_value="published"
    )
    
    db.commit()
    db.refresh(document)
    
    return serialize_document(document, db)


@router.get("/{document_id}/audit-logs")
def get_document_audit_logs(
    document_id: int,
    skip: int = 0,
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    logs = db.query(DocumentAuditLog).options(
        joinedload(DocumentAuditLog.user)
    ).filter(
        DocumentAuditLog.document_id == document_id
    ).order_by(DocumentAuditLog.performed_at.desc()).offset(skip).limit(limit).all()
    
    return [
        {
            "id": log.id,
            "action": log.action,
            "action_details": log.action_details,
            "field_changed": log.field_changed,
            "old_value": log.old_value,
            "new_value": log.new_value,
            "performed_by": log.performed_by,
            "performer_name": log.user.display_name if log.user else None,
            "performed_at": log.performed_at.isoformat() if log.performed_at else None,
            "ip_address": log.ip_address,
        }
        for log in logs
    ]


@router.post("/{document_id}/upload-file")
async def upload_document_file(
    document_id: int,
    file: UploadFile = File(...),
    change_summary: Optional[str] = Form(None),
    change_reason: Optional[str] = Form(None),
    create_new_version: bool = Form(True),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Upload a file (PDF, Word, Excel) for an existing governance document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    file_ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only PDF, Word (.doc, .docx), and Excel (.xls, .xlsx) files are supported"
        )
    
    unique_id = str(uuid.uuid4())
    safe_filename = f"{unique_id}_{file.filename}"
    file_path = os.path.join(GOVERNANCE_UPLOAD_DIR, safe_filename)
    
    content = await file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    old_file_name = document.file_name
    
    current_ver = document.current_version or "1.0"
    if create_new_version and document.file_name:
        try:
            version_parts = current_ver.split('.')
            major = int(version_parts[0]) if version_parts else 1
            minor = int(version_parts[1]) if len(version_parts) > 1 else 0
            new_version = f"{major}.{minor + 1}"
        except (ValueError, IndexError):
            new_version = "1.1"
        
        existing_versions = db.query(GovernanceDocumentVersion).filter(
            GovernanceDocumentVersion.document_id == document_id,
            GovernanceDocumentVersion.status == "current"
        ).all()
        for v in existing_versions:
            v.status = "superseded"
        
        version = GovernanceDocumentVersion(
            document_id=document_id,
            version_number=new_version,
            change_type="minor",
            title=document.title,
            content=document.content,
            file_name=file.filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_ext,
            change_summary=change_summary or f"File updated: {file.filename}",
            change_reason=change_reason,
            status="current",
            created_by=current_user.id
        )
        db.add(version)
        document.current_version = new_version
    
    document.file_name = file.filename
    document.file_path = file_path
    document.file_size = file_size
    document.file_type = file_ext
    document.updated_at = datetime.utcnow()
    
    create_audit_log(
        db=db,
        document_id=document.id,
        tenant_id=document.tenant_id,
        user_id=current_user.id,
        action="file_uploaded",
        action_details=f"File uploaded: {file.filename} ({file_size} bytes)",
        field_changed="file_name",
        old_value=old_file_name,
        new_value=file.filename
    )
    
    db.commit()
    db.refresh(document)
    
    return {
        "message": "File uploaded successfully",
        "document": serialize_document(document, db)
    }


@router.get("/{document_id}/download-file")
def download_document_file(
    document_id: int,
    version_id: Optional[int] = None,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Download the attached file for a governance document"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    if version_id:
        version = db.query(GovernanceDocumentVersion).filter(
            GovernanceDocumentVersion.id == version_id,
            GovernanceDocumentVersion.document_id == document_id
        ).first()
        
        if not version:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Version not found"
            )
        
        if not version.file_path or not os.path.exists(version.file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="File not found for this version"
            )
        
        return FileResponse(
            path=version.file_path,
            filename=version.file_name or f"document_{document_id}_v{version.version_number}",
            media_type=ALLOWED_FILE_TYPES.get(version.file_type, "application/octet-stream")
        )
    
    if not document.file_path or not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No file attached to this document"
        )
    
    return FileResponse(
        path=document.file_path,
        filename=document.file_name or f"document_{document_id}",
        media_type=ALLOWED_FILE_TYPES.get(document.file_type, "application/octet-stream")
    )


@router.post("/upload-with-file", status_code=status.HTTP_201_CREATED)
async def create_document_with_file(
    file: UploadFile = File(...),
    title: str = Form(...),
    doc_type: str = Form(...),
    description: Optional[str] = Form(None),
    document_code: Optional[str] = Form(None),
    classification: str = Form("internal"),
    owner_id: Optional[int] = Form(None),
    tenant_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Create a new governance document with an attached file"""
    if tenant_id:
        validate_tenant_access(current_user, tenant_id, db)
    else:
        tenant_id = get_user_primary_tenant(current_user, db)
        if not tenant_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User is not assigned to any tenant"
            )
    
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No file provided"
        )
    
    file_ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Only PDF, Word (.doc, .docx), and Excel (.xls, .xlsx) files are supported"
        )
    
    unique_id = str(uuid.uuid4())
    safe_filename = f"{unique_id}_{file.filename}"
    file_path = os.path.join(GOVERNANCE_UPLOAD_DIR, safe_filename)
    
    content = await file.read()
    file_size = len(content)
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    document = GovernanceDocument(
        tenant_id=tenant_id,
        title=title,
        description=description,
        document_code=document_code,
        doc_type=doc_type,
        classification=classification,
        owner_id=owner_id or current_user.id,
        author_id=current_user.id,
        file_name=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_ext,
        status="draft",
        current_version="1.0"
    )
    db.add(document)
    db.flush()
    
    version = GovernanceDocumentVersion(
        document_id=document.id,
        version_number="1.0",
        change_type="major",
        title=title,
        file_name=file.filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_ext,
        change_summary="Initial version",
        status="current",
        created_by=current_user.id
    )
    db.add(version)
    
    create_audit_log(
        db=db,
        document_id=document.id,
        tenant_id=tenant_id,
        user_id=current_user.id,
        action="created",
        action_details=f"Document created with file: {file.filename}"
    )
    
    db.commit()
    db.refresh(document)
    
    return {
        "message": "Document created successfully",
        "document": serialize_document(document, db)
    }


class PolicyAIDraftRequest(BaseModel):
    doc_type: str
    title: str
    framework_ids: Optional[List[int]] = None
    regulatory_scope: Optional[List[str]] = None
    description: Optional[str] = None
    include_sections: Optional[List[str]] = None


class PolicySuggestRequest(BaseModel):
    framework_ids: List[int]


def generate_policy_with_openai(
    doc_type: str,
    title: str,
    controls_context: str,
    regulatory_scope: List[str],
    description: Optional[str],
    include_sections: Optional[List[str]]
) -> dict:
    if not AI_INTEGRATIONS_OPENAI_API_KEY or not AI_INTEGRATIONS_OPENAI_BASE_URL:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI integration not configured"
        )
    
    client = OpenAI(
        api_key=AI_INTEGRATIONS_OPENAI_API_KEY,
        base_url=AI_INTEGRATIONS_OPENAI_BASE_URL,
        timeout=120.0
    )
    
    doc_type_labels = {
        "policy": "Policy",
        "standard": "Standard", 
        "procedure": "Procedure",
        "guideline": "Guideline"
    }
    doc_label = doc_type_labels.get(doc_type, "Policy")
    
    sections_instruction = ""
    if include_sections:
        sections_instruction = f"\n\nThe document MUST include these specific sections: {', '.join(include_sections)}"
    
    regulatory_context = ""
    if regulatory_scope:
        regulatory_context = f"\n\nThis document should align with the following regulatory frameworks: {', '.join(regulatory_scope)}"
    
    description_context = ""
    if description:
        description_context = f"\n\nAdditional requirements: {description}"
    
    doc_type_structures = {
        "policy": {
            "system_prompt": "You are an expert governance and compliance consultant that creates professional policy documents. Policies define high-level intent, principles, and organizational commitments. Always respond with valid JSON.",
            "structure": """Generate a comprehensive policy document with the following structure:
1. Purpose - Why this policy exists, its objectives, and strategic intent
2. Scope - Who and what this policy applies to (departments, personnel, systems, data)
3. Policy Statements - The key principles, directives, and mandatory requirements. Each statement should be clear, actionable, and auditable
4. Roles and Responsibilities - Detailed accountability matrix: who owns, who implements, who monitors, who enforces
5. Definitions - Key terms and acronyms used in this policy
6. Compliance and Enforcement - How compliance will be measured, monitored, and enforced; consequences of non-compliance
7. Exceptions - Process for requesting policy exceptions
8. Related Documents - References to related standards, procedures, and guidelines
9. Review and Updates - Review frequency, triggers for updates, approval authority

For each section, provide detailed, professional content suitable for an enterprise organization. Each policy statement should be specific and measurable."""
        },
        "procedure": {
            "system_prompt": "You are an expert governance and compliance consultant that creates detailed operational procedure documents. Procedures provide step-by-step instructions for carrying out specific activities. They must be actionable, detailed, and include decision points, checklists, and verification steps. Always respond with valid JSON.",
            "structure": """Generate a comprehensive, detailed PROCEDURE document. A procedure is NOT a policy — it must contain specific, actionable step-by-step instructions that someone can follow to complete a task.

The document MUST follow this structure:
1. Purpose - Why this procedure exists and what operational outcome it achieves
2. Scope - Who performs this procedure, when it applies, and what systems/processes are involved
3. Prerequisites - What must be in place before starting (access, tools, approvals, data, prior steps completed)
4. Definitions and Acronyms - Key terms used in this procedure
5. Roles and Responsibilities - Specific roles involved in each phase (executor, reviewer, approver, escalation contacts)
6. Detailed Procedure Steps - This is the CORE section. Provide numbered, detailed step-by-step instructions organized into logical phases:
   - Each step must have a clear action verb (e.g., "Navigate to...", "Verify that...", "Submit the...", "Record the...")
   - Include sub-steps (e.g., 6.1.1, 6.1.2) for complex actions
   - Include decision points: "IF [condition], THEN [action A]; OTHERWISE [action B]"
   - Include expected outcomes and verification checks after critical steps
   - Include screenshots/form field descriptions where applicable
   - Include wait times, SLAs, or time limits for each phase
   - Group steps into logical phases (e.g., Phase 1: Preparation, Phase 2: Execution, Phase 3: Verification)
7. Escalation Procedures - What to do when issues arise, who to contact, escalation timelines
8. Verification and Quality Checks - How to confirm the procedure was completed correctly; checklist of completion criteria
9. Troubleshooting - Common issues that may arise during execution and how to resolve them
10. Records and Documentation - What records must be kept, where to store them, retention period
11. Flowchart / Process Summary - A text-based process flow summary showing the major steps and decision points
12. References and Related Documents - Related policies, standards, forms, and templates
13. Revision History - Version tracking table

CRITICAL REQUIREMENTS FOR PROCEDURE DOCUMENTS:
- Every step in Section 6 must start with an action verb
- Include at least 15-25 detailed steps across all phases
- Include at least 3 decision points with IF/THEN/OTHERWISE logic
- Include specific form fields, system names, or data points where applicable
- Include timing requirements (e.g., "within 24 hours", "by end of business day")
- Include a verification checklist in Section 8 with at least 5 items"""
        },
        "standard": {
            "system_prompt": "You are an expert governance and compliance consultant that creates technical standards documents. Standards define specific, measurable requirements and technical specifications that must be met. They are more prescriptive than policies and include concrete thresholds, configurations, and acceptance criteria. Always respond with valid JSON.",
            "structure": """Generate a comprehensive technical standard document with the following structure:
1. Purpose - Why this standard exists and what baseline it establishes
2. Scope - Systems, technologies, processes, and personnel this standard applies to
3. Normative References - Other standards and specifications referenced
4. Definitions and Acronyms - Technical terms and abbreviations
5. Standard Requirements - Detailed, specific, measurable requirements organized by category. Each requirement must:
   - Have a unique identifier (e.g., REQ-001, REQ-002)
   - Be testable and verifiable
   - Include specific values, thresholds, or configurations (e.g., "Password minimum length: 12 characters")
   - Specify whether MANDATORY (SHALL/MUST) or RECOMMENDED (SHOULD)
   - Reference applicable framework controls
6. Technical Specifications - Specific configurations, architectures, protocols, or parameters
7. Compliance Criteria - How each requirement will be tested and verified; pass/fail criteria
8. Roles and Responsibilities - Who implements, who verifies, who grants exceptions
9. Exceptions and Waivers - Process for deviating from this standard
10. Implementation Timeline - Phased rollout plan with milestones
11. Related Documents - Linked policies, procedures, and guidelines
12. Review and Updates - Review cycle and update triggers

Each requirement in Section 5 must be specific and measurable, not vague or aspirational."""
        },
        "guideline": {
            "system_prompt": "You are an expert governance and compliance consultant that creates practical guideline documents. Guidelines provide recommended practices and flexible advice. They are advisory rather than mandatory, offering best practices and practical tips. Always respond with valid JSON.",
            "structure": """Generate a comprehensive guideline document with the following structure:
1. Purpose - Why this guideline exists and what best practices it promotes
2. Scope - Who this guideline is intended for and when to apply it
3. Background and Context - Industry trends, common challenges, and why these practices matter
4. Recommended Practices - Detailed best practices organized by topic area. For each practice:
   - Explain what to do and why
   - Provide practical examples and tips
   - Indicate maturity levels (basic, intermediate, advanced)
   - Note common pitfalls to avoid
5. Implementation Guidance - How to adopt these practices; phased approach suggestions
6. Tools and Resources - Recommended tools, templates, and reference materials
7. Roles and Responsibilities - Suggested ownership and accountability
8. Measuring Effectiveness - Key metrics and indicators to track adoption and outcomes
9. Frequently Asked Questions - Common questions and practical answers
10. Related Documents - Linked policies, standards, and procedures
11. Review and Updates - How often this guideline will be refreshed

Use an advisory tone (SHOULD, RECOMMENDED, CONSIDER) rather than mandatory language."""
        }
    }
    
    doc_config = doc_type_structures.get(doc_type, doc_type_structures["policy"])
    system_prompt = doc_config["system_prompt"]
    structure_instructions = doc_config["structure"]
    
    prompt = f"""You are an expert governance and compliance consultant. Generate a professional {doc_label} document titled "{title}".

{controls_context}
{regulatory_context}
{description_context}
{sections_instruction}

{structure_instructions}

Reference the specific control requirements from the frameworks where applicable.

Return a JSON object with these keys:
- "generated_content": The full document content in markdown format with proper headings (##), numbered lists, sub-lists, and formatting. This should be the complete document.
- "suggested_title": The recommended document title as a string.
- "suggested_sections": An array of objects, each with "heading" (string) and "content" (string) keys, one per section from the structure above.
- "framework_alignment": An array of objects with "framework" (string) and "controls" (array of strings) keys.

Example format:
{{
  "generated_content": "# Document Title\\n\\n## 1. Purpose\\n...",
  "suggested_title": "Document Title",
  "suggested_sections": [{{"heading": "1. Purpose", "content": "..."}}, {{"heading": "2. Scope", "content": "..."}}],
  "framework_alignment": [{{"framework": "Framework Name", "controls": ["Control ID"]}}]
}}"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=16000 if doc_type in ("procedure", "standard") else 12000
        )
        
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        return result
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse OpenAI response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"OpenAI API error: {error_msg}"
        )


@router.post("/ai-suggest-policies")
def suggest_policies_for_framework(
    request: PolicySuggestRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Given framework IDs, use AI to suggest what policies, procedures, and standards can be developed"""
    if not request.framework_ids:
        raise HTTPException(status_code=400, detail="At least one framework must be selected")
    
    frameworks = db.query(UploadedFramework).filter(
        UploadedFramework.id.in_(request.framework_ids)
    ).all()
    
    if not frameworks:
        raise HTTPException(status_code=404, detail="No frameworks found")
    
    controls_summary = ""
    for framework in frameworks:
        controls = db.query(ParsedFrameworkControl).filter(
            ParsedFrameworkControl.uploaded_framework_id == framework.id
        ).all()
        
        controls_summary += f"\nFramework: {framework.name}\n"
        controls_summary += f"Total Controls: {len(controls)}\n"
        controls_summary += "Key Control Areas:\n"
        
        for control in controls[:50]:
            ref = control.original_reference or control.control_id
            controls_summary += f"- {ref}: {control.title}"
            if control.description:
                controls_summary += f" - {control.description[:150]}"
            controls_summary += "\n"
    
    if not AI_INTEGRATIONS_OPENAI_API_KEY or not AI_INTEGRATIONS_OPENAI_BASE_URL:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI integration not configured"
        )
    
    client = OpenAI(
        api_key=AI_INTEGRATIONS_OPENAI_API_KEY,
        base_url=AI_INTEGRATIONS_OPENAI_BASE_URL
    )
    
    framework_names = ", ".join([f.name for f in frameworks])
    
    prompt = f"""You are an expert governance and compliance consultant. Based on the following regulatory framework controls, suggest comprehensive lists of policies, procedures, and standards that an organization should develop to achieve compliance.

{controls_summary}

Analyze ALL the control areas and requirements, then suggest documents organized by type. For each suggestion, provide:
- A clear, professional title
- A brief description of what the document should cover
- The priority level (high, medium, low) based on regulatory importance
- Which specific control references from the framework it addresses

Return a JSON object with this structure:
{{
  "framework_names": "{framework_names}",
  "suggestions": [
    {{
      "doc_type": "policy",
      "title": "Information Security Policy",
      "description": "Establishes the organization's approach to information security, defining objectives, principles, and responsibilities for protecting information assets.",
      "priority": "high",
      "relevant_controls": ["Control-1.1", "Control-1.2"],
      "key_sections": ["Purpose", "Scope", "Policy Statements", "Roles and Responsibilities", "Compliance", "Review"]
    }}
  ],
  "total_suggestions": 0
}}

Provide at least 15-25 suggestions covering policies, procedures, standards, and guidelines. Order by priority (high first).
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert governance and compliance consultant. Always respond with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=8192
        )
        
        result_text = response.choices[0].message.content or "{}"
        result = json.loads(result_text)
        return result
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to parse AI response: {str(e)}"
        )
    except Exception as e:
        error_msg = str(e)
        if "FREE_CLOUD_BUDGET_EXCEEDED" in error_msg:
            raise HTTPException(
                status_code=status.HTTP_402_PAYMENT_REQUIRED,
                detail="Cloud budget exceeded. Please upgrade your plan."
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI API error: {error_msg}"
        )


@router.post("/ai-draft")
def generate_policy_ai_draft(
    request: PolicyAIDraftRequest,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Generate a policy/standard/procedure document using AI based on framework requirements"""
    controls_context = ""
    framework_controls = []
    
    if request.framework_ids:
        frameworks = db.query(UploadedFramework).filter(
            UploadedFramework.id.in_(request.framework_ids)
        ).all()
        
        if frameworks:
            controls_context = "The document should align with the following framework controls:\n\n"
            
            for framework in frameworks:
                controls = db.query(ParsedFrameworkControl).filter(
                    ParsedFrameworkControl.uploaded_framework_id == framework.id
                ).limit(50).all()
                
                if controls:
                    framework_name = framework.name
                    control_ids = []
                    controls_context += f"Framework: {framework_name}\n"
                    
                    for control in controls:
                        ref = control.original_reference or control.control_id
                        controls_context += f"- {ref}: {control.title}\n"
                        if control.description:
                            controls_context += f"  Description: {control.description[:200]}...\n"
                        control_ids.append(ref)
                    
                    controls_context += "\n"
                    framework_controls.append({
                        "framework": framework_name,
                        "controls": control_ids[:10]
                    })
    
    if request.regulatory_scope:
        for scope in request.regulatory_scope:
            matching_frameworks = db.query(UploadedFramework).filter(
                UploadedFramework.name.ilike(f"%{scope}%")
            ).all()
            
            for framework in matching_frameworks:
                if not any(fc.get("framework") == framework.name for fc in framework_controls):
                    controls = db.query(ParsedFrameworkControl).filter(
                        ParsedFrameworkControl.uploaded_framework_id == framework.id
                    ).limit(30).all()
                    
                    if controls:
                        control_ids = [c.original_reference or c.control_id for c in controls]
                        framework_controls.append({
                            "framework": framework.name,
                            "controls": control_ids[:10]
                        })
                        
                        controls_context += f"Framework: {framework.name}\n"
                        for control in controls:
                            ref = control.original_reference or control.control_id
                            controls_context += f"- {ref}: {control.title}\n"
                        controls_context += "\n"
    
    result = generate_policy_with_openai(
        doc_type=request.doc_type,
        title=request.title,
        controls_context=controls_context,
        regulatory_scope=request.regulatory_scope or [],
        description=request.description,
        include_sections=request.include_sections
    )
    
    generated_content = result.get("generated_content", "")
    word_count = len(generated_content.split()) if generated_content else 0
    estimated_review_time = f"{max(5, word_count // 100)} minutes"
    
    return {
        "generated_content": generated_content,
        "suggested_title": result.get("suggested_title", request.title),
        "suggested_sections": result.get("suggested_sections", []),
        "framework_alignment": framework_controls if framework_controls else result.get("framework_alignment", []),
        "word_count": word_count,
        "estimated_review_time": estimated_review_time
    }


@router.get("/{document_id}/view-html")
def get_document_html(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: GRCUser = Depends(require_auth)
):
    """Convert uploaded document to HTML for in-platform viewing"""
    user_tenants = get_user_tenants(current_user, db)
    
    document = db.query(GovernanceDocument).filter(
        GovernanceDocument.id == document_id,
        GovernanceDocument.tenant_id.in_(user_tenants)
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # If document has a file, always prefer file-based conversion (better HTML)
    if document.file_path and os.path.exists(document.file_path):
        html_content = ""
        file_type = document.file_type or ""
        
        try:
            if file_type == "pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(document.file_path)
                sections = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        lines = page_text.split('\n')
                        page_html_parts = []
                        for line in lines:
                            stripped = line.strip()
                            if not stripped:
                                continue
                            if stripped.isupper() and len(stripped) < 100:
                                page_html_parts.append(f'<h3 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-weight: 600;">{html.escape(stripped)}</h3>')
                            elif len(stripped) < 80 and not stripped.endswith('.') and not stripped.endswith(','):
                                page_html_parts.append(f'<h4 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(stripped)}</h4>')
                            else:
                                page_html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
                        
                        sections.append(f'<div class="page" style="margin-bottom: 2rem; padding-bottom: 1rem; border-bottom: 1px solid #334155;"><div style="color: #475569; font-size: 0.75rem; margin-bottom: 0.5rem;">Page {i+1}</div>{"".join(page_html_parts)}</div>')
                html_content = "".join(sections)
            
            elif file_type in ("docx", "doc"):
                from docx import Document as DocxDocument
                doc = DocxDocument(document.file_path)
                parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    style_name = (para.style.name or "").lower() if para.style else ""
                    if "heading 1" in style_name:
                        parts.append(f'<h1 style="color: #f1f5f9; margin-top: 2rem; margin-bottom: 0.75rem; font-size: 1.5rem; font-weight: 700;">{html.escape(text)}</h1>')
                    elif "heading 2" in style_name:
                        parts.append(f'<h2 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-size: 1.25rem; font-weight: 600;">{html.escape(text)}</h2>')
                    elif "heading 3" in style_name:
                        parts.append(f'<h3 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.5rem; font-size: 1.1rem; font-weight: 500;">{html.escape(text)}</h3>')
                    elif "heading" in style_name:
                        parts.append(f'<h4 style="color: #cbd5e1; margin-top: 0.75rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(text)}</h4>')
                    elif "title" in style_name:
                        parts.append(f'<h1 style="color: #f1f5f9; margin-top: 1rem; margin-bottom: 1rem; font-size: 1.75rem; font-weight: 700; text-align: center;">{html.escape(text)}</h1>')
                    elif "list" in style_name:
                        parts.append(f'<li style="color: #94a3b8; margin-left: 1.5rem; margin-bottom: 0.25rem; line-height: 1.6;">{html.escape(text)}</li>')
                    else:
                        is_bold = all(run.bold for run in para.runs if run.text.strip()) if para.runs else False
                        if is_bold and len(text) < 100:
                            parts.append(f'<h4 style="color: #cbd5e1; margin-top: 0.75rem; margin-bottom: 0.25rem; font-weight: 600;">{html.escape(text)}</h4>')
                        else:
                            parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(text)}</p>')
                
                for table in doc.tables:
                    table_html = '<table style="width: 100%; border-collapse: collapse; margin: 1rem 0; border: 1px solid #334155;">'
                    for i, row in enumerate(table.rows):
                        table_html += '<tr>'
                        for cell in row.cells:
                            tag = 'th' if i == 0 else 'td'
                            style = 'padding: 0.5rem; border: 1px solid #334155; color: #94a3b8;'
                            if i == 0:
                                style += ' background-color: #1e293b; color: #e2e8f0; font-weight: 600;'
                            table_html += f'<{tag} style="{style}">{html.escape(cell.text.strip())}</{tag}>'
                        table_html += '</tr>'
                    table_html += '</table>'
                    parts.append(table_html)
                
                html_content = "".join(parts)
            
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        except HTTPException:
            raise
        except Exception as e:
            if document.content:
                lines = document.content.split('\n')
                html_parts = []
                for line in lines:
                    stripped = line.strip()
                    if not stripped:
                        continue
                    html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
                return {
                    "document_id": document.id,
                    "title": document.title,
                    "content_type": "html",
                    "html": "".join(html_parts),
                    "file_type": document.file_type
                }
            raise HTTPException(status_code=500, detail=f"Error converting document: {str(e)}")
        
        return {
            "document_id": document.id,
            "title": document.title,
            "content_type": "html",
            "html": html_content,
            "file_type": file_type
        }
    
    # Fall back to content field if no file
    if document.content:
        lines = document.content.split('\n')
        html_parts = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.isupper() and len(stripped) < 100:
                html_parts.append(f'<h3 style="color: #e2e8f0; margin-top: 1.5rem; margin-bottom: 0.5rem; font-weight: 600;">{html.escape(stripped)}</h3>')
            elif len(stripped) < 80 and not stripped.endswith('.') and not stripped.endswith(','):
                html_parts.append(f'<h4 style="color: #cbd5e1; margin-top: 1rem; margin-bottom: 0.25rem; font-weight: 500;">{html.escape(stripped)}</h4>')
            else:
                html_parts.append(f'<p style="color: #94a3b8; margin-bottom: 0.5rem; line-height: 1.6;">{html.escape(stripped)}</p>')
        return {
            "document_id": document.id,
            "title": document.title,
            "content_type": "html",
            "html": "".join(html_parts),
            "file_type": document.file_type
        }
    
    raise HTTPException(status_code=404, detail="Document file not found")
